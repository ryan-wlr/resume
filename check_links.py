#!/usr/bin/env python3
"""Check the latest DOCX for duplicate links and hyperlink status"""

import sys
import os

def check_latest_docx_links():
    """Check the latest generated DOCX for link duplication and hyperlink status"""
    print("🔍 Checking latest generated DOCX for link issues...")
    
    # Find the latest browse_mode_output folder
    folders = [f for f in os.listdir('.') if f.startswith('browse_mode_output_')]
    if not folders:
        print("❌ No browse_mode_output folders found")
        return
    
    latest_folder = sorted(folders)[-1]
    docx_path = os.path.join(latest_folder, 'optimized_resume.docx')
    
    print(f"📂 Checking: {docx_path}")
    
    if not os.path.exists(docx_path):
        print("❌ DOCX file not found")
        return
    
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Get all paragraphs and their text
        content_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        print(f"\n📄 First 10 content lines:")
        for i, line in enumerate(content_lines[:10]):
            print(f"  {i+1:2d}. {line}")
        
        # Look for GitHub/LinkedIn mentions
        github_count = 0
        linkedin_count = 0
        
        for i, line in enumerate(content_lines):
            if 'github' in line.lower():
                print(f"\n🔗 GitHub mention #{github_count + 1} (line {i+1}): {line}")
                github_count += 1
            if 'linkedin' in line.lower():
                print(f"🔗 LinkedIn mention #{linkedin_count + 1} (line {i+1}): {line}")
                linkedin_count += 1
        
        print(f"\n📊 Summary:")
        print(f"   GitHub mentions: {github_count}")
        print(f"   LinkedIn mentions: {linkedin_count}")
        
        if github_count > 1 or linkedin_count > 1:
            print("❌ DUPLICATE LINKS FOUND!")
        else:
            print("✅ No duplicate links")
            
        # Check for hyperlinks in the document
        print(f"\n🔗 Checking for hyperlinks...")
        hyperlink_count = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.color and run.font.color.rgb:
                    print(f"   Found colored text: {run.text}")
                # Check for hyperlink relationships
                if hasattr(run.element, 'hyperlink'):
                    hyperlink_count += 1
                    print(f"   Found hyperlink: {run.text}")
        
        if hyperlink_count == 0:
            print("❌ NO HYPERLINKS FOUND - links are plain text")
        else:
            print(f"✅ Found {hyperlink_count} hyperlinks")
            
    except ImportError:
        print("⚠️  python-docx not available for content verification")
    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")

if __name__ == "__main__":
    check_latest_docx_links()