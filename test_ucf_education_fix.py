#!/usr/bin/env python3
"""Test education fix - verify UCF education appears in DOCX"""

import sys
import os
import tempfile

# Add the current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from resume_windows import ResumeOptimizer

def test_ucf_education_fix():
    """Test that University of Central Florida education appears in DOCX"""
    print("🔍 Testing UCF Education Fix...")
    
    # Create temp output directory
    temp_dir = tempfile.mkdtemp()
    print(f"Using temp directory: {temp_dir}")
    
    try:
        # Initialize optimizer
        optimizer = ResumeOptimizer()
        
        # Mock detected field
        optimizer.current_detected_field = 'data_scientist'
        
        # Create mock results WITHOUT narrative (to test field data path)
        mock_results = {
            '1_job_analysis_report': 'Job analysis for Data Scientist...',
            '2_resume_analysis_flaws': 'Resume analysis...',
            '3_resume_impact_version': 'Impact version...',
            '4_resume_ats_optimized': 'ATS optimized...',
            '5_enhanced_skills_section': 'Enhanced skills...',
            '6_keyword_enhanced_experience': 'Enhanced experience...',
            '7_enhanced_standard_resume': 'Enhanced resume with University of Central Florida education...',
            '8_optimization_executive_summary': 'Executive summary...'
            # NO narrative_resume - should use field data path
        }
        
        # Create the DOCX using standard path (not narrative)
        print("   Creating DOCX using standard field data path...")
        optimizer.create_formatted_docx_resume_specific(mock_results, temp_dir, 'test_ucf_education.docx')
        
        # Check if DOCX was created
        docx_path = os.path.join(temp_dir, 'test_ucf_education.docx')
        if os.path.exists(docx_path):
            print("✅ DOCX created successfully")
            
            # Read and check education
            try:
                from docx import Document
                doc = Document(docx_path)
                
                # Get all text content
                full_text = '\n'.join([p.text for p in doc.paragraphs])
                
                # Check for correct education
                ucf_found = 'University of Central Florida' in full_text
                valencia_found = 'Valencia College' in full_text
                fau_found = 'Florida Atlantic University' in full_text
                
                print(f"\n📄 Education Check:")
                print(f"   ✅ University of Central Florida: {'✅' if ucf_found else '❌'}")
                print(f"   ✅ Valencia College: {'✅' if valencia_found else '❌'}")
                print(f"   ❌ Florida Atlantic University: {'❌' if not fau_found else '✅ (should not be present)'}")
                
                if ucf_found and valencia_found and not fau_found:
                    print("\n🎉 SUCCESS: Correct UCF/Valencia education found, no FAU!")
                else:
                    print(f"\n❌ ISSUE: Wrong education detected")
                    print("   Full education content:")
                    for p in doc.paragraphs:
                        if 'University' in p.text or 'College' in p.text:
                            print(f"     {p.text}")
                
            except Exception as e:
                print(f"   ❌ Error reading DOCX: {e}")
        else:
            print("❌ DOCX not created")
            
    except Exception as e:
        print(f"❌ Error in test: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # Cleanup
        try:
            import shutil
            shutil.rmtree(temp_dir)
            print(f"🗑️  Cleaned up temp directory")
        except:
            pass

if __name__ == "__main__":
    test_ucf_education_fix()