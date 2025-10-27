#!/usr/bin/env python3
"""Verify the latest browse mode DOCX has clickable hyperlinks"""

import sys
import os

def verify_latest_hyperlinks():
    """Verify the latest generated DOCX has proper clickable hyperlinks"""
    print("🔍 Verifying latest browse mode DOCX for clickable hyperlinks...")
    
    # Find the latest browse_mode_output folder
    folders = [f for f in os.listdir('.') if f.startswith('browse_mode_output_')]
    if not folders:
        print("❌ No browse_mode_output folders found")
        return
    
    latest_folder = sorted(folders)[-1]
    docx_path = os.path.join(latest_folder, 'optimized_resume.docx')
    
    print(f"📂 Checking: {docx_path}")
    
    if not os.path.exists(docx_path):
        print("❌ DOCX file not found")
        return
    
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Get all paragraphs and their text
        content_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        print(f"\n📄 Contact section:")
        for i, line in enumerate(content_lines[:6]):
            if any(word in line.lower() for word in ['phone', '📞', 'linkedin', 'github', 'https']):
                print(f"  {i+1:2d}. {line}")
        
        # Check for hyperlinks
        print(f"\n🔗 Hyperlink analysis:")
        hyperlink_count = 0
        hyperlink_urls = []
        
        # Check hyperlink relationships
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                hyperlink_urls.append(rel.target_ref)
                hyperlink_count += 1
                print(f"   ✅ Clickable link: {rel.target_ref}")
        
        # Check XML for hyperlink elements
        xml_hyperlinks = 0
        for paragraph in doc.paragraphs:
            if 'w:hyperlink' in paragraph._p.xml:
                xml_hyperlinks += 1
        
        print(f"\n📊 Summary:")
        print(f"   Total hyperlink relationships: {hyperlink_count}")
        print(f"   XML hyperlink elements: {xml_hyperlinks}")
        print(f"   Duplicate links: {'❌ NO' if hyperlink_count <= 2 else '✅ YES'}")
        
        # Check for expected links
        expected_links = [
            'linkedin.com/in/ryan-weiler',
            'github.com/ryan-wlr'
        ]
        
        found_links = []
        for expected in expected_links:
            for url in hyperlink_urls:
                if expected in url:
                    found_links.append(expected)
                    break
        
        print(f"\n🎯 Expected links found: {len(found_links)}/{len(expected_links)}")
        for link in found_links:
            print(f"   ✅ {link}")
        
        missing_links = [link for link in expected_links if link not in found_links]
        for link in missing_links:
            print(f"   ❌ Missing: {link}")
        
        if hyperlink_count >= 2 and len(found_links) == len(expected_links):
            print(f"\n🎉 SUCCESS: All links are clickable hyperlinks!")
            print(f"📝 Manual test: Open {os.path.abspath(docx_path)} and click the links")
        else:
            print(f"\n⚠️  Issues found with hyperlinks")
            
    except ImportError:
        print("⚠️  python-docx not available for verification")
    except Exception as e:
        print(f"❌ Error verifying hyperlinks: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    verify_latest_hyperlinks()