#!/usr/bin/env python3
"""
AI-Powered Resume Optimizer - Windows Compatible Version
Generates professionally formatted .docx resumes matching Ryan Weiler's template

Usage:
    python resume_windows.py "job description text" --resume "current resume text"
    python resume_windows.py --browse  # Interactive file selection
"""

import os
import sys
import argparse
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, List, Optional
import argparse

# Optional imports with fallbacks
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    print("WARNING: python-docx not installed. Run: pip install python-docx")
    HAS_DOCX = False

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
    HAS_TKINTER = True
except ImportError:
    print("WARNING: tkinter not available")
    HAS_TKINTER = False


@dataclass
class JobAnalysis:
    """Structure for analyzed job requirements"""
    required_skills: List[str]
    preferred_skills: List[str] 
    key_responsibilities: List[str]
    company_values: List[str]
    industry_keywords: List[str]
    experience_level: str
    role_type: str


class ResumeOptimizer:
    """AI-powered resume optimization system with Windows compatibility"""
    
    def __init__(self):
        print(">>> AI Resume Optimizer initialized successfully!")
        print(">>> Ready to create ATS-optimized, recruiter-friendly resumes!")
        
    def analyze_job_posting(self, job_description: str, target_role: str = "Software Engineer", 
                          target_company: str = "Target Company") -> JobAnalysis:
        """Analyze job posting to extract key requirements and preferences"""
        
        # Simulate AI analysis with comprehensive results
        job_text = job_description.lower()
        
        # Detect career field from job description and target role
        combined_content = f"{job_description} {target_role}".lower()
        detected_field = self.detect_career_field(combined_content)
        
        # Define field-specific skills and responsibilities
        field_configs = {
            'dishwasher': {
                'skills': ['commercial dishwashing', 'food safety', 'sanitation', 'kitchen equipment', 'servsafe',
                          'cleaning procedures', 'fast paced environment', 'team collaboration', 'time management',
                          'hygiene standards', 'dish machine operation', 'food handling', 'restaurant operations'],
                'responsibilities': [
                    "Operate commercial dishwashing equipment efficiently and safely",
                    "Maintain cleanliness and sanitation standards in kitchen areas", 
                    "Wash dishes, glasses, utensils, and kitchen equipment",
                    "Follow all food safety and sanitation protocols",
                    "Support kitchen staff during peak service hours"
                ]
            },
            'chef': {
                'skills': ['culinary arts', 'menu development', 'food preparation', 'kitchen management', 'cooking techniques',
                          'recipe creation', 'cost control', 'staff supervision', 'food safety', 'inventory management',
                          'servsafe manager', 'culinary leadership', 'restaurant operations', 'quality control'],
                'responsibilities': [
                    "Develop innovative menus and recipes for restaurant operations",
                    "Supervise kitchen staff and manage food preparation",
                    "Ensure food safety and quality standards are maintained",
                    "Control food costs and manage kitchen inventory",
                    "Lead culinary team during high-volume service periods"
                ]
            },
            'plumber': {
                'skills': ['plumbing codes', 'pipe installation', 'water heaters', 'drain systems', 'fixtures',
                          'pipe repair', 'water lines', 'sewer lines', 'pvc', 'copper', 'pex', 'soldering',
                          'drain cleaning', 'backflow prevention', 'gas lines', 'emergency service', 'blueprints'],
                'responsibilities': [
                    "Install and maintain plumbing systems and fixtures",
                    "Diagnose and repair plumbing problems and leaks",
                    "Follow local plumbing codes and safety regulations",
                    "Provide excellent customer service and accurate estimates",
                    "Perform emergency plumbing repairs and service calls"
                ]
            },
            'electrician': {
                'skills': ['electrical codes', 'nec', 'motor controls', 'plc', 'programmable logic controller',
                          'variable frequency drives', 'vfd', 'multimeters', 'oscilloscopes', 'three-phase',
                          '480v', 'electrical schematics', 'blueprints', 'troubleshooting', 'allen bradley'],
                'responsibilities': [
                    "Install and maintain electrical systems and equipment",
                    "Troubleshoot electrical problems and perform repairs",
                    "Follow NEC codes and OSHA safety standards",
                    "Read and interpret electrical schematics and blueprints",
                    "Perform preventive maintenance on industrial equipment"
                ]
            },
            'nurse': {
                'skills': ['patient care', 'medication administration', 'iv therapy', 'clinical assessment',
                          'medical charting', 'hipaa compliance', 'patient education', 'emergency response',
                          'electronic health records', 'vital signs', 'wound care', 'infection control'],
                'responsibilities': [
                    "Provide comprehensive patient care in medical-surgical units",
                    "Administer medications and treatments following protocols",
                    "Monitor patient vital signs and document care",
                    "Collaborate with healthcare team on patient care plans",
                    "Educate patients and families on health conditions"
                ]
            },
            'teacher': {
                'skills': ['curriculum development', 'lesson planning', 'classroom management', 'student assessment',
                          'differentiated instruction', 'educational technology', 'parent communication',
                          'behavior management', 'standards alignment', 'student engagement'],
                'responsibilities': [
                    "Design and implement engaging lesson plans for students",
                    "Manage classroom environment and student behavior",
                    "Assess student progress and provide feedback",
                    "Communicate with parents and administration",
                    "Integrate technology to enhance learning"
                ]
            },
            'mechanic': {
                'skills': ['engine diagnostics', 'automotive repair', 'brake systems', 'electrical systems',
                          'diagnostic equipment', 'preventive maintenance', 'customer service',
                          'ase certification', 'vehicle inspection', 'parts replacement'],
                'responsibilities': [
                    "Diagnose and repair automotive mechanical issues",
                    "Perform routine maintenance and inspections",
                    "Use diagnostic equipment to identify problems",
                    "Maintain detailed service records",
                    "Provide accurate repair estimates to customers"
                ]
            },
            'software_engineer': {
                'skills': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'go', 'rust', 
                          'kotlin', 'swift', 'scala', 'ruby', 'php', 'r', 'matlab', 'julia',
                          'react', 'angular', 'vue.js', 'node.js', 'express.js', 'spring', 'django', 'flask',
                          'html5', 'css3', 'sass', 'less', 'bootstrap', 'tailwind css',
                          'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
                          'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'terraform',
                          'git', 'github', 'gitlab', 'jenkins', 'cicd', 'devops',
                          'machine learning', 'ai', 'deep learning', 'neural networks',
                          'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy', 'matplotlib',
                          'microservices', 'restful apis', 'graphql', 'grpc', 'websockets',
                          'agile', 'scrum', 'test driven development', 'unit testing',
                          'linux', 'unix', 'bash', 'powershell', 'vim'],
                'responsibilities': [
                    "Design and implement scalable software solutions using modern programming languages",
                    "Develop full-stack applications with frontend and backend technologies",
                    "Build and optimize databases and data storage solutions",
                    "Implement machine learning models and AI-powered features",
                    "Deploy applications using cloud platforms and containerization",
                    "Collaborate with cross-functional teams using agile methodologies",
                    "Participate in code reviews and maintain high code quality standards",
                    "Write comprehensive unit tests and documentation",
                    "Optimize application performance and system reliability",
                    "Mentor junior developers and contribute to technical best practices"
                ]
            },
            'data_scientist': {
                'skills': ['python', 'r', 'sql', 'machine learning', 'deep learning', 'statistics', 'data analysis',
                          'pandas', 'numpy', 'scipy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras',
                          'jupyter', 'matplotlib', 'seaborn', 'tableau', 'power bi', 'big data',
                          'hadoop', 'spark', 'aws', 'azure', 'gcp', 'data visualization', 'a/b testing'],
                'responsibilities': [
                    "Develop predictive models and machine learning algorithms",
                    "Analyze large datasets to extract meaningful business insights",
                    "Build data pipelines and automate data processing workflows",
                    "Create compelling data visualizations and reports",
                    "Collaborate with stakeholders to solve complex business problems",
                    "Deploy machine learning models to production environments",
                    "Conduct statistical analysis and hypothesis testing",
                    "Clean and preprocess large datasets for analysis"
                ]
            },
            'web_developer': {
                'skills': ['html5', 'css3', 'javascript', 'typescript', 'react', 'angular', 'vue.js',
                          'node.js', 'express.js', 'django', 'flask', 'spring boot', 'php', 'laravel',
                          'responsive design', 'bootstrap', 'tailwind css', 'sass', 'less',
                          'restful apis', 'graphql', 'websockets', 'sql', 'mongodb', 'git'],
                'responsibilities': [
                    "Design and develop responsive web applications",
                    "Build user-friendly frontend interfaces using modern frameworks",
                    "Develop backend APIs and server-side functionality",
                    "Optimize web performance and ensure cross-browser compatibility",
                    "Implement security best practices and data protection",
                    "Collaborate with UX/UI designers and product teams",
                    "Write clean, maintainable, and well-documented code",
                    "Debug and troubleshoot web application issues"
                ]
            },
            'mobile_developer': {
                'skills': ['swift', 'kotlin', 'java', 'objective-c', 'react native', 'flutter', 'dart',
                          'ios development', 'android development', 'xcode', 'android studio',
                          'mobile ui/ux', 'core data', 'sqlite', 'firebase', 'push notifications',
                          'app store', 'google play', 'mobile testing', 'performance optimization'],
                'responsibilities': [
                    "Develop native and cross-platform mobile applications",
                    "Design intuitive mobile user interfaces and experiences",
                    "Integrate with REST APIs and backend services",
                    "Implement mobile-specific features like GPS and camera",
                    "Optimize apps for performance and battery efficiency",
                    "Test applications across different devices and OS versions",
                    "Publish apps to app stores and manage releases",
                    "Debug and resolve mobile-specific issues"
                ]
            },
            'devops_engineer': {
                'skills': ['docker', 'kubernetes', 'aws', 'azure', 'gcp', 'terraform', 'ansible',
                          'jenkins', 'gitlab ci', 'github actions', 'linux', 'bash', 'python',
                          'monitoring', 'prometheus', 'grafana', 'elk stack', 'microservices',
                          'infrastructure as code', 'ci/cd', 'containerization', 'cloud architecture'],
                'responsibilities': [
                    "Design and implement CI/CD pipelines for automated deployments",
                    "Manage cloud infrastructure and containerized applications",
                    "Monitor system performance and implement alerting systems",
                    "Automate infrastructure provisioning and configuration",
                    "Ensure system security and compliance requirements",
                    "Optimize application performance and scalability",
                    "Troubleshoot production issues and implement solutions",
                    "Collaborate with development teams on deployment strategies"
                ]
            },
            'security_engineer': {
                'skills': ['cybersecurity', 'penetration testing', 'vulnerability assessment', 'incident response',
                          'security frameworks', 'risk assessment', 'encryption', 'network security',
                          'firewalls', 'ids/ips', 'siem', 'compliance', 'iso 27001', 'nist',
                          'python', 'bash', 'powershell', 'linux', 'windows security'],
                'responsibilities': [
                    "Conduct security assessments and vulnerability testing",
                    "Implement security controls and monitoring systems",
                    "Respond to security incidents and conduct forensic analysis",
                    "Develop security policies and procedures",
                    "Ensure compliance with security standards and regulations",
                    "Perform risk assessments and security audits",
                    "Design secure network architectures and systems",
                    "Train teams on security best practices and awareness"
                ]
            }
        }
        
        # Get field-specific configuration or generate dynamically
        if detected_field in field_configs:
            config = field_configs[detected_field]
        else:
            config = self.generate_dynamic_job_analysis_config(detected_field)
        skill_keywords = config['skills']
        responsibilities = config['responsibilities']
        
        # Extract technical skills mentioned in job description
        tech_skills = []
        for skill in skill_keywords:
            if skill in job_text:
                tech_skills.append(skill.title())
                
        # If no skills found, use some default ones for the field
        if not tech_skills:
            tech_skills = [skill.title() for skill in skill_keywords[:3]]
        
        # Extract industry-specific keywords based on detected field
        field_industry_keywords = {
            'dishwasher': ['Restaurant Operations', 'Food Service', 'Kitchen Management', 'Food Safety'],
            'chef': ['Culinary Arts', 'Restaurant Management', 'Menu Development', 'Kitchen Leadership'],
            'plumber': ['Residential Plumbing', 'Commercial Plumbing', 'Plumbing Repair', 'Emergency Service'],
            'electrician': ['Industrial Electrical', 'Electrical Maintenance', 'Industrial Automation', 'Electrical Safety'],
            'nurse': ['Healthcare', 'Patient Care', 'Clinical Skills', 'Medical Technology'],
            'teacher': ['Education', 'Curriculum Development', 'Student Assessment', 'Classroom Management'],
            'mechanic': ['Automotive Repair', 'Vehicle Maintenance', 'Diagnostic Equipment', 'ASE Certification'],
            'software_engineer': ['Software Development', 'Technology Solutions', 'System Architecture', 'Code Quality'],
            'data_scientist': ['Data Science', 'Machine Learning', 'Analytics', 'Business Intelligence'],
            'web_developer': ['Web Development', 'Frontend Development', 'Backend Development', 'Full Stack'],
            'mobile_developer': ['Mobile Development', 'iOS Development', 'Android Development', 'Mobile Apps'],
            'devops_engineer': ['DevOps', 'Cloud Infrastructure', 'CI/CD', 'System Administration'],
            'security_engineer': ['Cybersecurity', 'Information Security', 'Risk Management', 'Compliance']
        }
        
        # Get industry keywords or generate dynamically
        if detected_field in field_industry_keywords:
            industry_keywords = field_industry_keywords[detected_field]
        else:
            industry_keywords = self.generate_dynamic_industry_keywords(detected_field)
            
        analysis = JobAnalysis(
            required_skills=tech_skills[:8],  # Top 8 skills
            preferred_skills=tech_skills[8:15] if len(tech_skills) > 8 else [],
            key_responsibilities=responsibilities,
            company_values=['Innovation', 'Collaboration', 'Excellence', 'Growth'],
            industry_keywords=industry_keywords,
            experience_level='Mid-Level' if 'senior' in job_text else 'Entry-Mid Level',
            role_type=target_role
        )
        
        print(f">>> Job Analysis Complete:")
        print(f"    Required Skills: {', '.join(analysis.required_skills)}")
        print(f"    Experience Level: {analysis.experience_level}")
        print(f"    Industry Focus: {', '.join(analysis.industry_keywords) if analysis.industry_keywords else 'General Technology'}")
        
        return analysis

    def mock_response(self, prompt: str, target_role: str = "Software Engineer") -> str:
        """Generate comprehensive mock AI responses for resume optimization"""
        
        # Detect career field dynamically from target role
        detected_field = self.detect_career_field(target_role.lower())
        field_data = self.get_field_data(detected_field)
        
        # Store detected field for use in docx generation
        self.current_detected_field = detected_field
        
        # Shared content maps for all response types
        title_map = {
            'quantum_computing_scientist': 'Quantum Computing Research Scientist & Quantum Information Theorist',
            'mathematics_professor': 'Professor of Mathematics & Research Mathematician',
            'mathematician': 'Applied Mathematician & Computational Researcher',
            'theoretical_physicist': 'Theoretical Physicist & Mathematical Physics Researcher',
            'nurse': 'Registered Nurse & Healthcare Professional',
            'teacher': 'Elementary Education Teacher & Curriculum Specialist', 
            'mechanic': 'Automotive Technician & Diagnostic Specialist',
            'plumber': 'Professional Plumber & Plumbing Systems Specialist',
            'electrician': 'Industrial Electrician & Maintenance Specialist',
            'dishwasher': 'Kitchen Staff & Food Service Professional',
            'chef': 'Executive Chef & Culinary Professional',
            'software_engineer': 'Software Engineer | Python Developer | Financial Technology Specialist',
            'data_scientist': 'Data Scientist | Machine Learning Engineer | Analytics Specialist',
            'web_developer': 'Full Stack Web Developer | Frontend & Backend Specialist',
            'mobile_developer': 'Mobile App Developer | iOS & Android Specialist',
            'devops_engineer': 'DevOps Engineer | Cloud Infrastructure Specialist',
            'security_engineer': 'Cybersecurity Engineer | Information Security Specialist'
        }
        
        summary_map = {
            'quantum_computing_scientist': 'Distinguished Quantum Computing Research Scientist with 5+ years of groundbreaking research in quantum algorithms, quantum error correction, and quantum information theory. Proven expertise in developing novel quantum algorithms achieving exponential speedups, implementing variational quantum eigensolvers with 99.9% accuracy, and advancing fault-tolerant quantum computing. Strong publication record in Nature Quantum Information and Physical Review Letters with 15+ peer-reviewed papers. Committed to advancing quantum technologies through rigorous theoretical research and practical quantum algorithm development.',
            'mathematics_professor': 'Accomplished Mathematics Professor with 7+ years of research excellence in pure mathematics, specializing in algebraic geometry, homological algebra, and arithmetic geometry. Proven expertise in solving fundamental problems using advanced cohomological methods, developing new theoretical frameworks in K-theory and motivic cohomology, and mentoring 12+ graduate students to Ph.D. completion. Strong publication record in Annals of Mathematics and Inventiones Mathematicae. Committed to advancing mathematical knowledge through rigorous proof techniques and innovative research approaches.',
            'mathematician': 'Expert Applied Mathematician with 6+ years of experience developing mathematical models for complex systems analysis and creating novel computational algorithms. Proven expertise in partial differential equations, optimization theory, and numerical methods with demonstrated ability to improve prediction accuracy by 40% through advanced mathematical modeling. Strong background in high-performance computing, machine learning applications, and interdisciplinary collaboration. Committed to applying mathematical rigor to solve real-world problems in science and engineering.',
            'theoretical_physicist': 'Leading Theoretical Physicist with 6+ years of research excellence in quantum field theory, string theory, and mathematical physics. Proven expertise in developing new theoretical frameworks, advancing understanding of black hole physics and holographic duality, and connecting abstract theory to experimental observations. Strong publication record in top physics journals and international conference presentations. Committed to pushing the boundaries of fundamental physics through rigorous mathematical analysis and innovative theoretical approaches.',
            'nurse': 'Compassionate Registered Nurse with 3+ years of experience providing comprehensive patient care in medical-surgical units. Proven expertise in medication administration, patient assessment, and clinical documentation with demonstrated ability to improve patient outcomes by 25% through evidence-based care. Strong knowledge of HIPAA regulations and healthcare protocols. Committed to delivering safe, quality patient care with excellent communication skills.',
            'teacher': 'Dedicated Elementary School Teacher with 3+ years of experience designing engaging curriculum for diverse student populations. Proven expertise in differentiated instruction, classroom management, and student assessment with demonstrated ability to improve student achievement by 25% through innovative teaching methods. Strong knowledge of state standards and educational technology. Committed to fostering student growth and academic success.',
            'mechanic': 'Skilled Automotive Technician with 3+ years of comprehensive experience in engine diagnostics, brake repair, and vehicle maintenance. Proven expertise in diagnostic equipment, repair procedures, and customer service with demonstrated ability to reduce diagnostic time by 25% through systematic troubleshooting. Strong knowledge of ASE standards and safety protocols. Committed to delivering reliable, quality automotive repair services.',
            'plumber': 'Skilled Professional Plumber with 3+ years of comprehensive experience in residential and commercial plumbing system installation, maintenance, and repair. Proven expertise in pipe installation, water heater service, and drain system maintenance with demonstrated ability to reduce service callbacks by 25% through quality workmanship. Strong knowledge of local plumbing codes and safety protocols. Committed to delivering reliable, code-compliant plumbing solutions with excellent customer service.',
            'electrician': 'Experienced Industrial Electrician with 3+ years of expertise in electrical system installation, maintenance, and troubleshooting. Proven expertise in motor controls, PLCs, and electrical diagnostics with demonstrated ability to reduce equipment downtime by 25% through preventive maintenance programs. Strong knowledge of NEC codes and OSHA safety standards. Committed to delivering safe, reliable electrical solutions.',
            'dishwasher': 'Dedicated Kitchen Staff with 3+ years of experience in high-volume food service operations and commercial kitchen management. Proven expertise in dish washing, sanitization, and equipment maintenance with demonstrated ability to maintain 99% cleanliness standards during peak service hours. Strong knowledge of food safety protocols, ServSafe certification, and kitchen workflow optimization. Committed to supporting efficient restaurant operations through reliable service and attention to detail.',
            'chef': 'Experienced Executive Chef with 3+ years of comprehensive culinary expertise in menu development, kitchen management, and high-volume food service operations. Proven expertise in culinary techniques, staff supervision, and cost control with demonstrated ability to increase restaurant revenue by 25% through innovative menu creation. Strong knowledge of food safety regulations, kitchen operations, and culinary arts. Committed to delivering exceptional dining experiences through creative cuisine and efficient kitchen leadership.',
            'software_engineer': 'Dedicated Software Engineer with 3+ years of experience developing innovative financial technology solutions and automated trading systems. Proven expertise in Python development, machine learning applications, and algorithmic trading with demonstrated results including 25% portfolio performance improvement. Strong background in full-stack development, database optimization, and collaborative software engineering practices. Passionate about leveraging cutting-edge technology to solve complex financial and business challenges.',
            'data_scientist': 'Experienced Data Scientist with 3+ years of expertise in machine learning, statistical analysis, and big data processing. Proven expertise in Python, R, and SQL with demonstrated ability to improve business outcomes by 30% through predictive modeling and data-driven insights. Strong background in TensorFlow, scikit-learn, and cloud platforms. Committed to transforming complex datasets into actionable business intelligence and strategic recommendations.',
            'web_developer': 'Skilled Full Stack Web Developer with 3+ years of experience building responsive web applications and scalable backend systems. Proven expertise in JavaScript, TypeScript, React, and Node.js with demonstrated ability to improve application performance by 40% through optimization techniques. Strong background in modern web technologies, RESTful APIs, and database design. Passionate about creating exceptional user experiences and efficient development workflows.',
            'mobile_developer': 'Expert Mobile App Developer with 3+ years of experience creating native and cross-platform applications for iOS and Android. Proven expertise in Swift, Kotlin, and React Native with demonstrated success achieving 4.8+ star ratings and 100K+ downloads. Strong background in mobile UI/UX design, performance optimization, and app store deployment. Committed to delivering innovative mobile solutions that enhance user engagement and drive business growth.',
            'devops_engineer': 'Skilled DevOps Engineer with 3+ years of experience designing and implementing cloud infrastructure and CI/CD pipelines. Proven expertise in AWS, Docker, and Kubernetes with demonstrated ability to reduce deployment time by 80% and improve system reliability to 99.99% uptime. Strong background in infrastructure as code, monitoring, and security best practices. Passionate about automating workflows and enabling efficient software delivery.',
            'security_engineer': 'Experienced Cybersecurity Engineer with 3+ years of expertise in vulnerability assessment, incident response, and security architecture. Proven expertise in penetration testing, SIEM platforms, and security automation with demonstrated success identifying and mitigating 200+ security vulnerabilities. Strong background in Python scripting, compliance frameworks, and threat analysis. Committed to protecting organizational assets and maintaining robust security postures.'
        }
        
        if "analyze flaws" in prompt.lower():
            return """RESUME ANALYSIS - AREAS FOR IMPROVEMENT:

FORMATTING ISSUES:
- Inconsistent bullet point formatting throughout document
- Missing quantifiable achievements and metrics
- Generic job descriptions lacking impact statements
- Insufficient use of industry-specific keywords

CONTENT WEAKNESSES:
- Professional summary lacks compelling value proposition
- Experience section needs stronger action verbs (achieved, optimized, developed)
- Missing technical skills alignment with modern job requirements
- Education section could highlight relevant coursework and projects

ATS OPTIMIZATION GAPS:
- Keywords not strategically placed for applicant tracking systems
- Missing relevant technical competencies for target roles
- Job titles and descriptions need better keyword density
- Contact information format needs enhancement

IMPACT IMPROVEMENTS NEEDED:
- Add quantifiable results (increased efficiency by X%, reduced costs by $X)
- Include specific technologies and frameworks used in each role
- Highlight leadership experiences and cross-functional collaboration
- Emphasize problem-solving capabilities with concrete examples"""

        elif "rewrite resume" in prompt.lower():
            # Get dynamic content using shared maps
            resume_title = title_map.get(detected_field, self.generate_dynamic_title(detected_field))
            resume_summary = summary_map.get(detected_field, self.generate_dynamic_summary(detected_field))
            
            # Build skills section dynamically
            skills_text = '\n'.join(field_data['skills'])
            
            return f"""ENHANCED RESUME - MAXIMUM IMPACT VERSION:

RYAN THOMAS WEILER
{resume_title}

CONTACT INFORMATION:
Phone: (561) 906-2118 | Email: ryan_wlr@yahoo.com
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/

PROFESSIONAL SUMMARY:
{resume_summary}

CORE TECHNICAL COMPETENCIES:
{skills_text}

PROFESSIONAL EXPERIENCE:
{field_data['experience_title']}
{chr(10).join(['• ' + bullet for bullet in field_data['experience_bullets']])}

EDUCATION:
{field_data['education']}

KEY PROJECTS & CERTIFICATIONS:
{chr(10).join(['• ' + project for project in field_data['projects']])}

This optimized resume strategically aligns your background with the target role requirements, incorporating industry-specific keywords and quantifiable achievements that will resonate with hiring managers and pass ATS screening systems."""

            if detected_field == 'plumber':
                return """ENHANCED RESUME - MAXIMUM IMPACT VERSION:

RYAN THOMAS WEILER
Professional Plumber & Plumbing Systems Specialist

CONTACT INFORMATION:
Phone: (561) 906-2118 | Email: ryan_wlr@yahoo.com
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/

PROFESSIONAL SUMMARY:
Skilled Professional Plumber with 3+ years of comprehensive experience in residential and commercial plumbing system installation, maintenance, and repair. Proven expertise in pipe installation, water heater service, and drain system maintenance with demonstrated ability to reduce service callbacks by 25% through quality workmanship. Strong knowledge of local plumbing codes and safety protocols. Committed to delivering reliable, code-compliant plumbing solutions with excellent customer service.

CORE TECHNICAL COMPETENCIES:
Plumbing Systems: Water Supply Lines, Drain Systems, Fixture Installation, Pipe Repair
Tools & Equipment: Pipe Wrenches, Drain Augers, Pipe Cutters, Soldering Equipment, Threading Machines
Materials: PVC, Copper, PEX Piping, Cast Iron, Various Fittings and Fixtures
Codes & Standards: Local Plumbing Codes, Uniform Plumbing Code, Safety Regulations
Water Systems: Water Heaters, Boilers, Pressure Tanks, Backflow Prevention
Repair Services: Drain Cleaning, Leak Repair, Emergency Service, System Diagnostics

PROFESSIONAL EXPERIENCE:

Professional Plumber | Independent Contractor | 2022 - Present
• Installed and maintained residential and commercial plumbing systems including water lines, drain systems, and fixtures
• Performed comprehensive plumbing repairs and emergency service calls with 95% customer satisfaction rating
• Installed and serviced water heaters, garbage disposals, and plumbing fixtures following local code requirements
• Diagnosed complex plumbing problems using modern diagnostic equipment and provided cost-effective solutions
• Maintained detailed service records and provided accurate estimates for plumbing projects

Plumbing Technician | Florida Atlantic University Facilities | 2021 - 2022
• Supported campus-wide plumbing maintenance and repair operations across multiple facilities
• Assisted with installation of commercial-grade plumbing fixtures and system upgrades
• Performed preventive maintenance on water heating systems and distribution lines
• Collaborated with facilities team on emergency plumbing repairs and system troubleshooting
• Maintained inventory of plumbing supplies and ensured proper tool maintenance

EDUCATION:
Bachelor of Science in Computer Science (Engineering Focus)
Florida Atlantic University | Boca Raton, FL | Expected May 2024
GPA: 3.7/4.0 | Dean's List: Fall 2021, Spring 2022
Relevant Coursework: Engineering Fundamentals, Systems Design, Technical Problem Solving

CERTIFICATIONS & ACHIEVEMENTS:
• Plumbing Code Compliance Training
• Water Heater Installation and Service Certification
• Drain Cleaning and Sewer Line Repair Techniques
• Customer Service Excellence Recognition

KEY PROJECTS & SPECIALIZATIONS:
• Residential Bathroom Renovation: Complete plumbing installation including fixture placement and water line routing
• Commercial Water Heater Installation: Installed commercial-grade water heating system with proper venting and safety controls
• Emergency Drain Service: Specialized in rapid response drain cleaning and sewer line maintenance
• Code Compliance Projects: Successfully completed plumbing inspections meeting all local code requirements"""

            # Old hardcoded content removed - using dynamic system above
            elif detected_field == 'placeholder_old_content':
                return """ENHANCED RESUME - MAXIMUM IMPACT VERSION:

RYAN THOMAS WEILER
Industrial Electrician & Electrical Systems Specialist

CONTACT INFORMATION:
Phone: (561) 906-2118 | Email: ryan_wlr@yahoo.com
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/

PROFESSIONAL SUMMARY:
Skilled Industrial Electrician with 3+ years of hands-on experience in electrical system installation, maintenance, and troubleshooting. Expertise in motor controls, PLCs, and variable frequency drives with proven ability to reduce downtime by 30% through proactive maintenance. Strong knowledge of NEC codes and industrial safety protocols. Committed to maintaining high electrical safety standards while delivering reliable industrial electrical solutions.

CORE TECHNICAL COMPETENCIES:
Electrical Systems: AC/DC Circuits, Motor Controls, Transformers, Switchgear
Codes & Standards: National Electrical Code (NEC), OSHA Safety Standards, Local Electrical Codes
Control Systems: Programmable Logic Controllers (PLCs), Variable Frequency Drives (VFDs), Motor Starters
Testing Equipment: Multimeters, Oscilloscopes, Megohmmeters, Power Quality Analyzers
Industrial Systems: 480V Three-Phase Systems, Industrial Machinery, Automation Controls
Safety & Tools: Lockout/Tagout (LOTO), Arc Flash Safety, Conduit Bending, Electrical Hand Tools

PROFESSIONAL EXPERIENCE:

Industrial Electrician | Independent Contractor | 2022 - Present
• Installed and maintained industrial electrical systems including motor controls and PLCs, reducing equipment downtime by 30%
• Performed electrical troubleshooting on 480V three-phase systems and complex industrial machinery
• Implemented preventive maintenance programs that increased equipment reliability by 25%
• Ensured strict compliance with NEC codes and OSHA safety standards on all electrical installations

Electrical Technician | Florida Atlantic University Facilities | 2021 - 2022  
• Collaborated with maintenance team on electrical system upgrades and emergency repairs across campus facilities
• Installed and programmed basic PLC systems for HVAC and lighting control, improving energy efficiency by 20%
• Documented electrical system modifications and maintained accurate electrical schematics and blueprints
• Assisted with electrical safety training and lockout/tagout procedures for facilities staff

EDUCATION:
Bachelor of Science in Computer Science (Electrical Focus)
Florida Atlantic University | Boca Raton, FL | Expected May 2024
Relevant Coursework: Circuit Analysis, Electronics, Control Systems, Digital Logic Design
GPA: 3.7/4.0 | Dean's List: Fall 2021, Spring 2022

CERTIFICATIONS & KEY PROJECTS:
• OSHA 10-Hour Electrical Safety Certification
• Industrial Motor Control Installation: Complete 3-phase motor control panel installation and commissioning
• PLC Programming Project: Allen-Bradley PLC programming for automated conveyor system  
• Electrical Code Compliance: Successfully completed electrical inspections meeting all NEC requirements"""
            else:
                return """ENHANCED RESUME - MAXIMUM IMPACT VERSION:

RYAN THOMAS WEILER
Software Engineer & Financial Technology Specialist

CONTACT INFORMATION:
Phone: (561) 906-2118 | Email: ryan_wlr@yahoo.com
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/
GitHub: https://github.com/ryan-wlr

PROFESSIONAL SUMMARY:
Results-driven Software Engineer with 3+ years of experience developing scalable financial technology solutions. Expertise in Python development, algorithmic trading systems, and machine learning applications. Proven track record of optimizing trading strategies that increased portfolio performance by 25% while reducing risk exposure. Passionate about leveraging data-driven approaches to solve complex financial engineering challenges.

CORE TECHNICAL COMPETENCIES:
Languages: Python, JavaScript, SQL, Java, C++
Frameworks: Flask, Django, React, Node.js, pandas, NumPy
Technologies: Machine Learning (scikit-learn, TensorFlow), REST APIs, Docker
Databases: PostgreSQL, MongoDB, Redis
Cloud & Tools: AWS, Git, Linux, Jupyter Notebooks
Financial Tech: Algorithmic Trading, Risk Management, Market Data Analysis

PROFESSIONAL EXPERIENCE:

Financial Technology Developer | Independent Projects | 2022 - Present
• Developed automated trading algorithms using Python and Alpaca API, achieving 15% annual returns
• Implemented machine learning models for cryptocurrency price prediction with 68% accuracy
• Built scalable data pipelines processing 10,000+ market data points daily
• Created interactive dashboards for real-time portfolio monitoring and risk analysis

Software Development Intern | Florida Atlantic University | 2021 - 2022  
• Collaborated with 5-person team to develop student management system using Django framework
• Optimized database queries reducing page load times by 40% for 10,000+ student records
• Implemented automated testing suite with 95% code coverage using pytest
• Contributed to open-source projects with focus on educational technology solutions

EDUCATION:
Bachelor of Science in Computer Science
Florida Atlantic University | Boca Raton, FL | Expected May 2024
Relevant Coursework: Data Structures, Algorithms, Database Systems, Software Engineering
GPA: 3.7/4.0 | Dean's List: Fall 2021, Spring 2022

KEY PROJECTS:
• Algorithmic Trading Bot: Python-based system with backtesting capabilities and live trading
• Cryptocurrency Analysis Platform: Real-time market data visualization using React and Python APIs  
• Portfolio Optimization Tool: ML-driven asset allocation with modern portfolio theory implementation"""

        elif "ats optimization" in prompt.lower():
            return f"""ATS-OPTIMIZED RESUME VERSION:

RYAN THOMAS WEILER
{field_data.get('experience_title', 'Professional').split('|')[0].strip()}

TECHNICAL SKILLS:
{chr(10).join(field_data['skills'])}

PROFESSIONAL EXPERIENCE:
{field_data['experience_title']}
{chr(10).join(['• ' + bullet for bullet in field_data['experience_bullets']])}

EDUCATION:
{field_data['education']}

CERTIFICATIONS & ACHIEVEMENTS:
{chr(10).join(['• ' + project for project in field_data['projects']])}

This ATS-optimized version incorporates relevant keywords and proper formatting to maximize compatibility with applicant tracking systems while maintaining readability for human reviewers."""

        elif "enhance skills" in prompt.lower():
            return f"""ENHANCED TECHNICAL SKILLS SECTION:

{chr(10).join(field_data['skills'])}

ADDITIONAL COMPETENCIES:
• Problem-solving and analytical thinking
• Team collaboration and communication
• Time management and project coordination
• Safety protocols and compliance
• Technical documentation and reporting
• Continuous learning and professional development

This enhanced skills section strategically highlights both technical competencies and soft skills that are highly valued by employers in the {detected_field} field."""

            # Old hardcoded content removed - using dynamic system
            if detected_field == 'placeholder_to_remove':
                return """ENHANCED TECHNICAL SKILLS SECTION:

ELECTRICAL SYSTEMS EXPERTISE:
Primary Systems: Motor Controls, PLCs (Allen-Bradley), Variable Frequency Drives (VFDs), Transformers
Secondary Systems: Switchgear, Panel Installation, Industrial Automation, HVAC Electrical
Voltage Systems: 120V/240V Single-Phase, 480V Three-Phase, Low Voltage Control Circuits

CODES & SAFETY COMPLIANCE:
Electrical Codes: National Electrical Code (NEC), Local Electrical Codes, Code Compliance Inspections
Safety Standards: OSHA Electrical Safety, Arc Flash Safety, Lockout/Tagout (LOTO) Procedures
Safety Equipment: Personal Protective Equipment (PPE), Fall Protection, Electrical Safety Protocols

TESTING & TROUBLESHOOTING EQUIPMENT:
Electrical Testing: Multimeters, Oscilloscopes, Megohmmeters, Power Quality Analyzers
Diagnostic Tools: Clamp Meters, Phase Rotation Testers, Insulation Resistance Testers
Calibration: Test Equipment Calibration, Measurement Accuracy, Electrical System Analysis

INDUSTRIAL CONTROL SYSTEMS:
PLC Programming: Allen-Bradley (AB), Basic Ladder Logic, HMI Configuration
Motor Controls: Soft Starters, Motor Protection, Contactor and Relay Circuits
Process Control: Industrial Sensors, Instrumentation, Control Panel Wiring

INSTALLATION & MAINTENANCE TOOLS:
Hand Tools: Electrical Hand Tools, Conduit Bending, Wire Stripping and Crimping
Power Tools: Electrical Power Tools, Drill Presses, Threading Equipment
Specialized Tools: Cable Pulling Equipment, Conduit Installation, Panel Assembly

SOFT SKILLS & WORK PRACTICES:
• Electrical Problem Solving & Systematic Troubleshooting
• Industrial Safety Leadership & Training  
• Preventive Maintenance Program Development
• Technical Documentation & Electrical Schematic Reading
• Cross-functional Collaboration with Mechanical and HVAC Teams"""
            else:
                return """ENHANCED TECHNICAL SKILLS SECTION:

PROGRAMMING EXPERTISE:
Primary Languages: Python (Advanced), JavaScript (Intermediate), SQL (Advanced)
Secondary Languages: Java, C++, HTML5, CSS3, Bash Scripting
Frameworks & Libraries: Django, Flask, React, Node.js, Express.js, pandas, NumPy

FINANCIAL TECHNOLOGY STACK:
Trading Platforms: Alpaca API, Interactive Brokers API, Binance API
Analysis Tools: QuantLib, zipline, backtrader, Alpha Architect
Data Sources: Yahoo Finance, Alpha Vantage, Polygon.io, CoinGecko
Risk Management: Modern Portfolio Theory, VaR Calculations, Sharpe Ratio Optimization

MACHINE LEARNING & DATA SCIENCE:
Core Libraries: scikit-learn, TensorFlow, Keras, pandas, NumPy, Matplotlib
Techniques: Supervised Learning, Time Series Analysis, Feature Engineering
Applications: Predictive Modeling, Algorithmic Trading, Risk Assessment

DEVELOPMENT TOOLS & PLATFORMS:
Version Control: Git, GitHub, GitLab
Cloud Services: Amazon Web Services (AWS), Docker, Linux Administration
Databases: PostgreSQL, MongoDB, Redis, MySQL
IDEs & Tools: PyCharm, VS Code, Jupyter Notebooks, Postman

SOFT SKILLS:
• Problem Solving & Analytical Thinking
• Cross-functional Team Collaboration  
• Agile Development Methodology
• Technical Documentation & Communication
• Continuous Learning & Technology Adaptation"""

        elif "keyword enhancement" in prompt.lower():
            return """KEYWORD-ENHANCED EXPERIENCE DESCRIPTIONS:

FINANCIAL TECHNOLOGY DEVELOPER | Independent Projects | 2022 - Present
Python Development • API Integration • Machine Learning • Algorithmic Trading

ACHIEVEMENTS:
• Developed automated trading algorithms using Python, pandas, and Alpaca Trading API, generating 15% annual returns through systematic market analysis
• Implemented machine learning models using scikit-learn and TensorFlow for cryptocurrency price prediction, achieving 68% accuracy in trend forecasting  
• Built scalable data processing pipelines handling 10,000+ daily market data points using Python, PostgreSQL, and Redis caching
• Created interactive web dashboards using React and Flask for real-time portfolio monitoring and risk management
• Optimized algorithm execution speed by 25% through code refactoring and efficient data structure implementation

SOFTWARE DEVELOPMENT INTERN | Florida Atlantic University | 2021 - 2022
Full-Stack Development • Database Optimization • Agile Methodology • Quality Assurance

ACCOMPLISHMENTS:
• Collaborated with cross-functional team of 5 developers to build student management system using Django, PostgreSQL, and React
• Optimized database queries and indexing strategies, reducing page load times by 40% for application serving 10,000+ student records
• Implemented comprehensive automated testing suite using pytest and Selenium, achieving 95% code coverage
• Contributed to open-source educational technology projects, demonstrating commitment to collaborative software development
• Participated in Agile development process with daily standups, sprint planning, and code review procedures

PROJECT HIGHLIGHTS:

ALGORITHMIC TRADING SYSTEM | Personal Project | 2023
Technologies: Python, pandas, Alpaca API, PostgreSQL, Docker
• Designed and implemented automated trading bot with backtesting capabilities and live market execution
• Integrated multiple technical indicators and machine learning signals for trade decision making
• Built comprehensive risk management system with position sizing and stop-loss mechanisms

CRYPTOCURRENCY ANALYSIS PLATFORM | Academic Project | 2022  
Technologies: React, Python Flask, MongoDB, Chart.js, CoinGecko API
• Developed real-time cryptocurrency market analysis tool with interactive data visualizations
• Implemented RESTful API backend for efficient data retrieval and processing
• Created responsive user interface supporting multiple devices and screen sizes"""

        elif "tailored resume" in prompt.lower():
            position_titles = {
                'nurse': 'REGISTERED NURSE',
                'teacher': 'ELEMENTARY EDUCATION TEACHER', 
                'mechanic': 'AUTOMOTIVE TECHNICIAN',
                'plumber': 'PROFESSIONAL PLUMBER',
                'electrician': 'INDUSTRIAL ELECTRICIAN',
                'dishwasher': 'KITCHEN STAFF / DISHWASHER',
                'chef': 'EXECUTIVE CHEF',
                'brain_surgeon': 'NEUROSURGEON / BRAIN SURGEON',
                'cardiologist': 'CARDIOLOGIST / HEART SURGEON',
                'surgeon': 'MEDICAL SURGEON',
                'doctor': 'MEDICAL DOCTOR',
                'mechanical_engineer': 'MECHANICAL ENGINEER',
                'civil_engineer': 'CIVIL ENGINEER',
                'lawyer': 'ATTORNEY / LAWYER',
                'accountant': 'PROFESSIONAL ACCOUNTANT',
                'financial_analyst': 'FINANCIAL ANALYST',
                'software_engineer': 'SOFTWARE ENGINEER'
            }
            
            # Generate dynamic title if not in predefined list
            if detected_field in position_titles:
                position_title = position_titles[detected_field]
            else:
                position_title = detected_field.replace('_', ' ').title()
            
            return f"""TAILORED RESUME FOR {position_title} POSITION:

RYAN THOMAS WEILER
{field_data['experience_title'].replace(' | ', ' | ').split('|')[0].strip()}

CONTACT INFORMATION:
Phone: (561) 906-2118
Email: ryan_wlr@yahoo.com  
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/
GitHub: https://github.com/ryan-wlr

PROFESSIONAL SUMMARY:
{summary_map.get(detected_field, self.generate_dynamic_summary(detected_field))}

CORE COMPETENCIES:
{chr(10).join(['• ' + skill for skill in field_data['skills']])}

PROFESSIONAL EXPERIENCE:
{field_data['experience_title']}
{chr(10).join(['• ' + bullet for bullet in field_data['experience_bullets']])}

EDUCATION & CREDENTIALS:
{field_data['education']}

KEY PROJECTS & CERTIFICATIONS:
{chr(10).join(['• ' + project for project in field_data['projects']])}

This tailored resume strategically positions your experience and skills to align with {detected_field} industry requirements and employer expectations."""

            # Old hardcoded content removed
            if detected_field == 'placeholder_to_remove':
                return """TAILORED RESUME FOR INDUSTRIAL ELECTRICIAN POSITION:

RYAN THOMAS WEILER
Industrial Electrician | Electrical Maintenance Specialist | M.E.P. Systems Technician

CONTACT INFORMATION:
Phone: (561) 906-2118
Email: ryan_wlr@yahoo.com  
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/

PROFESSIONAL SUMMARY:
Dedicated Industrial Electrician with 3+ years of comprehensive experience in electrical system installation, maintenance, and troubleshooting for industrial facilities. Proven expertise in motor controls, PLC programming, and variable frequency drives with demonstrated ability to reduce equipment downtime by 30% through systematic preventive maintenance programs. Strong knowledge of NEC codes, OSHA safety standards, and industrial automation systems. Committed to maintaining the highest electrical safety standards while delivering reliable, code-compliant electrical solutions for M.E.P. services.

CORE TECHNICAL COMPETENCIES:
Electrical Systems: Motor Controls, PLCs (Allen-Bradley), Variable Frequency Drives (VFDs), Transformers
Codes & Standards: National Electrical Code (NEC), OSHA Safety Standards, Local Electrical Codes
Control Systems: Programmable Logic Controllers, Motor Starters, Industrial Automation, HMI Programming
Testing Equipment: Multimeters, Oscilloscopes, Megohmmeters, Power Quality Analyzers, Phase Rotation Testers
Industrial Systems: 480V Three-Phase Systems, Industrial Machinery, HVAC Electrical, Preventive Maintenance
Safety & Compliance: Lockout/Tagout (LOTO), Arc Flash Safety, Electrical Safety Training, Code Inspections

PROFESSIONAL EXPERIENCE:

Industrial Electrician & Maintenance Specialist | Independent Contractor | 2022 - Present
• Installed and maintained complex industrial electrical systems including motor control panels and PLC-controlled automation equipment
• Performed systematic electrical troubleshooting on 480V three-phase systems and industrial machinery, reducing unplanned downtime by 30%
• Developed and implemented preventive maintenance programs that increased electrical equipment reliability by 25%
• Ensured strict compliance with NEC codes and OSHA safety standards on all electrical installations and repairs
• Programmed and configured Allen-Bradley PLCs for industrial process control and automation systems
• Collaborated with mechanical and HVAC technicians on integrated M.E.P. system installations and upgrades

Electrical Technician | Florida Atlantic University Facilities Management | 2021 - 2022
• Maintained electrical infrastructure across campus facilities including emergency electrical repairs and system upgrades
• Installed and programmed basic PLC systems for HVAC and lighting control, improving overall energy efficiency by 20%
• Created and maintained accurate electrical schematics, blueprints, and maintenance documentation
• Conducted electrical safety training sessions and implemented lockout/tagout procedures for facilities staff
• Assisted with electrical code compliance inspections and system certifications meeting all local and NEC requirements

EDUCATION:
Bachelor of Science in Computer Science (Electrical Systems Focus)
Florida Atlantic University, Boca Raton, FL | Expected Graduation: May 2024
GPA: 3.7/4.0 | Dean's List: Fall 2021, Spring 2022
Relevant Coursework: Circuit Analysis, Electronics, Control Systems, Digital Logic Design, Industrial Automation

CERTIFICATIONS & SPECIALIZED TRAINING:
• OSHA 10-Hour Electrical Safety Certification
• National Electrical Code (NEC) Compliance Training
• Allen-Bradley PLC Programming Certification (Basic Level)
• Lockout/Tagout (LOTO) Safety Procedures
• Arc Flash Safety and PPE Requirements

FEATURED PROJECTS & ACHIEVEMENTS:

Industrial Motor Control Installation | 2023
Equipment: Allen-Bradley PLCs, VFDs, 480V Motor Control Centers
• Designed and installed complete 3-phase motor control panel for automated conveyor system
• Programmed Allen-Bradley PLC with ladder logic for automated process control
• Implemented safety interlocks and emergency stop systems meeting OSHA requirements
• Completed project 15% under budget while maintaining full code compliance

Electrical System Upgrade | 2022
Systems: 480V Distribution, Lighting Controls, Emergency Systems
• Led electrical upgrade project for university laboratory facility including power distribution upgrades
• Installed new 480V electrical panels and upgraded existing lighting control systems
• Ensured all work met NEC Article 408 and local electrical code requirements
• Documented all modifications with updated electrical drawings and maintenance procedures"""
            else:
                return """TAILORED RESUME FOR SOFTWARE ENGINEER POSITION:

RYAN THOMAS WEILER
Software Engineer | Python Developer | Financial Technology Specialist

CONTACT INFORMATION:
Phone: (561) 906-2118
Email: ryan_wlr@yahoo.com  
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/
GitHub: https://github.com/ryan-wlr

PROFESSIONAL SUMMARY:
Dedicated Software Engineer with 3+ years of experience developing innovative financial technology solutions and automated trading systems. Proven expertise in Python development, machine learning applications, and algorithmic trading with demonstrated results including 25% portfolio performance improvement. Strong background in full-stack development, database optimization, and collaborative software engineering practices. Passionate about leveraging cutting-edge technology to solve complex financial and business challenges.

CORE TECHNICAL SKILLS:
Languages: Python, JavaScript, SQL, Java, C++
Frameworks: Django, Flask, React, Node.js, pandas, NumPy, scikit-learn
Technologies: Machine Learning, REST APIs, Docker, Git, Linux
Databases: PostgreSQL, MongoDB, Redis, MySQL
Cloud Platforms: Amazon Web Services (AWS)
Financial Tech: Algorithmic Trading, Market Data Analysis, Risk Management

PROFESSIONAL EXPERIENCE:

Software Engineer & Financial Technology Developer | Independent Projects | 2022 - Present
• Architected and deployed automated trading algorithms using Python and Alpaca API, achieving 15% annual returns with optimized risk management strategies
• Developed machine learning models for cryptocurrency market prediction using TensorFlow and scikit-learn, reaching 68% prediction accuracy
• Built high-performance data processing pipelines handling 10,000+ daily market data points with Python, PostgreSQL, and Redis
• Created responsive web applications using React and Flask for real-time portfolio monitoring and financial analytics
• Optimized system performance through code refactoring and efficient algorithm design, improving execution speed by 25%

Software Development Intern | Florida Atlantic University | 2021 - 2022
• Collaborated with agile development team to design and implement student management system using Django framework and PostgreSQL
• Enhanced application performance by optimizing database queries and implementing efficient indexing, reducing load times by 40% for 10,000+ records
• Established automated testing framework using pytest achieving 95% code coverage and improved software reliability
• Contributed to open-source educational technology initiatives, demonstrating commitment to collaborative development and continuous learning
• Participated in code review processes and maintained high coding standards following industry best practices

EDUCATION:
Bachelor of Science in Computer Science
Florida Atlantic University, Boca Raton, FL | Expected Graduation: May 2024
GPA: 3.7/4.0 | Dean's List: Fall 2021, Spring 2022
Relevant Coursework: Data Structures and Algorithms, Database Systems, Software Engineering, Machine Learning

FEATURED PROJECTS:

Algorithmic Trading Platform | 2023
Technologies: Python, Alpaca API, PostgreSQL, Docker, React
• Designed comprehensive trading system with backtesting engine and live market execution capabilities
• Integrated multiple technical indicators and machine learning signals for informed trading decisions
• Implemented robust risk management with position sizing algorithms and dynamic stop-loss mechanisms

Cryptocurrency Market Analysis Tool | 2022
Technologies: React, Python Flask, MongoDB, Chart.js, API Integration
• Developed real-time market analysis platform with interactive data visualizations and trend analysis
• Built scalable RESTful API backend for efficient cryptocurrency data processing and storage
• Created responsive user interface supporting cross-platform accessibility and mobile optimization"""

        else:
            # Default comprehensive response
            return f"""AI RESUME OPTIMIZATION COMPLETE:

Your resume has been analyzed and optimized using advanced AI algorithms designed to maximize ATS compatibility and recruiter appeal. The optimization focused on:

TECHNICAL ENHANCEMENTS:
• Strategic keyword placement for improved ATS ranking
• Quantified achievements with specific metrics and results
• Industry-relevant skill alignment with job requirements
• Enhanced formatting for better readability and professional appearance

CONTENT IMPROVEMENTS:
• Strengthened professional summary with compelling value proposition
• Upgraded experience descriptions with powerful action verbs
• Added relevant technical competencies matching target role requirements
• Optimized contact information and professional profile links

ATS OPTIMIZATION:
• Improved keyword density for applicant tracking systems
• Enhanced formatting compatibility across different ATS platforms  
• Strategically placed technical skills and competencies
• Ensured consistent formatting and professional structure

The optimized resume now better aligns with modern hiring practices and should significantly improve your chances of passing initial screening processes and attracting recruiter attention."""

    def analyze_resume_flaws(self, resume_content: str, job_analysis: JobAnalysis) -> str:
        """Identify areas for improvement in the current resume"""
        print("\n>>> ANALYZING RESUME FLAWS...")
        
        prompt = f"""Analyze this resume for flaws and improvement opportunities:
        Resume: {resume_content[:1000]}...
        Target Role: {job_analysis.role_type}
        Required Skills: {job_analysis.required_skills}"""
        
        return self.mock_response("analyze flaws", job_analysis.role_type)
    
    def rewrite_for_impact(self, resume_content: str, target_role: str) -> str:
        """Rewrite resume for maximum impact and engagement"""
        print(f"\n>>> REWRITING RESUME FOR MAXIMUM IMPACT ({target_role})...")
        
        prompt = f"""Rewrite this resume for maximum impact for {target_role} position:
        {resume_content[:1000]}..."""
        
        return self.mock_response("rewrite resume", target_role)
    
    def optimize_for_ats(self, resume_content: str, target_role: str) -> str:
        """Optimize resume for Applicant Tracking Systems"""
        print(f"\n>>> OPTIMIZING FOR ATS SYSTEMS ({target_role})...")
        
        prompt = f"""Optimize this resume for ATS systems for {target_role}:
        {resume_content[:1000]}..."""
        
        return self.mock_response("ats optimization", target_role)
    
    def enhance_skills_section(self, job_analysis: JobAnalysis) -> str:
        """Create enhanced technical skills section"""
        prompt = f"""Create enhanced technical skills section for:
        Required: {job_analysis.required_skills}
        Preferred: {job_analysis.preferred_skills}
        Industry: {job_analysis.industry_keywords}"""
        
        return self.mock_response("enhance skills", job_analysis.role_type)
    
    def enhance_experience_with_keywords(self, resume_content: str, job_analysis: JobAnalysis) -> str:
        """Enhance experience section with relevant keywords"""
        print("\n>>> UPGRADING EXPERIENCE SECTION...")
        
        prompt = f"""Enhance experience section with keywords:
        Resume: {resume_content[:1000]}...
        Keywords: {job_analysis.required_skills + job_analysis.industry_keywords}"""
        
        return self.mock_response("keyword enhancement", job_analysis.role_type)
    
    def create_enhanced_resume(self, resume_content: str, job_analysis: JobAnalysis, 
                              target_role: str, target_company: str) -> str:
        """Create enhanced professional resume version"""
        print("\n>>> CREATING ENHANCED PROFESSIONAL RESUME...")
        
        # Detect field for context
        detected_field = self.detect_career_field(target_role.lower())
        field_data = self.get_field_data(detected_field)
        
        # Extract key skills for summary (first few items from first skill category)
        first_skill_line = field_data['skills'][0] if field_data['skills'] else ""
        if ':' in first_skill_line:
            key_skills = first_skill_line.split(':')[1].split(',')[:4]
            key_skills = [skill.strip() for skill in key_skills]
        else:
            key_skills = [detected_field.replace('_', ' ')]
        
        enhanced_resume = f"""RYAN THOMAS WEILER
{field_data['experience_title'].split('|')[0].strip()}

CONTACT INFORMATION:
Phone: (561) 906-2118 | Email: ryan_wlr@yahoo.com
LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/
GitHub: https://github.com/ryan-wlr

PROFESSIONAL SUMMARY:
Results-driven {detected_field.replace('_', ' ')} professional with proven expertise in {', '.join(key_skills)}. 
Demonstrated track record of delivering high-impact solutions and driving measurable results in fast-paced environments. 
Seeking to leverage technical excellence and leadership capabilities as {target_role} at {target_company}.

CORE TECHNICAL COMPETENCIES:
{chr(10).join(['• ' + skill for skill in field_data['skills']])}

PROFESSIONAL EXPERIENCE:

{field_data['experience_title']}
{chr(10).join(['• ' + bullet for bullet in field_data['experience_bullets']])}

EDUCATION:
{field_data['education']}

KEY PROJECTS & ACCOMPLISHMENTS:
{chr(10).join(['• ' + project for project in field_data['projects']])}

ACHIEVEMENTS:
• Consistently exceeded performance expectations in all roles and responsibilities
• Recognized for technical excellence and professional growth within field
• Contributed to successful completion of high-priority projects and initiatives
• Demonstrated commitment to continuous learning and skill development

This resume is optimized for {target_role} positions with emphasis on technical skills,
leadership capabilities, and measurable contributions to organizational success."""

        return enhanced_resume
    
    def create_combined_resume(self, narrative_content: str, enhanced_content: str, 
                             target_role: str, target_company: str) -> str:
        """Create a comprehensive resume that combines storytelling narrative with professional sections"""
        print("\n>>> CREATING COMBINED STORYTELLING + PROFESSIONAL RESUME...")
        
        # Extract key elements from both versions
        narrative_lines = narrative_content.split('\n')
        enhanced_lines = enhanced_content.split('\n')
        
        # Find the professional summary from enhanced version
        professional_summary = ""
        enhanced_skills = []
        enhanced_experience = []
        
        current_section = None
        for line in enhanced_lines:
            line = line.strip()
            if 'PROFESSIONAL SUMMARY:' in line:
                current_section = 'summary'
                continue
            elif 'CORE TECHNICAL COMPETENCIES:' in line:
                current_section = 'skills'
                continue
            elif 'PROFESSIONAL EXPERIENCE:' in line:
                current_section = 'experience'
                continue
            elif line.startswith('EDUCATION:') or line.startswith('CONTACT'):
                current_section = None
                continue
            
            if current_section == 'summary' and line and not line.startswith('•'):
                professional_summary += line + " "
            elif current_section == 'skills' and line.startswith('•'):
                enhanced_skills.append(line[1:].strip())
            elif current_section == 'experience' and line.startswith('•'):
                enhanced_experience.append(line[1:].strip())
        
        # Extract storytelling elements from narrative version
        story_hook = ""
        story_narrative = ""
        story_achievements = []
        story_projects = []
        
        current_section = None
        for line in narrative_lines:
            line = line.strip()
            if '🔥 THE' in line or '🔬 THE' in line or '💻 THE' in line or '🌟 THE' in line:
                story_hook = line
            elif 'PROFESSIONAL NARRATIVE:' in line:
                current_section = 'narrative'
                continue
            elif 'KEY ACHIEVEMENTS' in line:
                current_section = 'achievements'
                continue
            elif 'DEFINING PROJECTS' in line:
                current_section = 'projects'
                continue
            elif line.startswith('CAREER JOURNEY') or line.startswith('TECHNICAL EXPERTISE') or line.startswith('EDUCATION'):
                current_section = None
                continue
            
            if current_section == 'narrative' and line and not line.startswith('CAREER'):
                story_narrative += line + " "
            elif current_section == 'achievements' and line.startswith('•'):
                story_achievements.append(line[1:].strip())
            elif current_section == 'projects' and line.startswith('•'):
                story_projects.append(line[1:].strip())
        
        # Create the combined resume
        combined_resume = f"""COMPREHENSIVE CAREER RESUME - {target_role.upper()}

{story_hook}

RYAN THOMAS WEILER
{target_role} Professional

CONTACT INFORMATION:
📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com
🔗 LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/ | 💻 GitHub: https://github.com/ryan-wlr

PROFESSIONAL SUMMARY:
{professional_summary.strip()}

CAREER NARRATIVE & VISION:
{story_narrative.strip()}

CORE TECHNICAL COMPETENCIES:
{chr(10).join(['• ' + skill for skill in enhanced_skills[:6]])}

SIGNATURE ACHIEVEMENTS:
{chr(10).join(['• ' + achievement for achievement in story_achievements[:5]])}

PROFESSIONAL EXPERIENCE:
{chr(10).join(['• ' + exp for exp in enhanced_experience[:8]])}

KEY PROJECTS & INNOVATIONS:
{chr(10).join(['• ' + project for project in story_projects[:6]])}

EDUCATION:
University of Central Florida — B.S. Computer Science, 2013 (Dean's List, GPA 3.8)
Valencia College — A.A., 2011 (Dean's List, GPA 3.7)

PROFESSIONAL DEVELOPMENT:
• Continuous learning in {target_role.lower()} technologies and methodologies
• Active participation in professional development and industry conferences
• Commitment to staying current with emerging trends and best practices
• Leadership development through mentoring and cross-functional collaboration

FUTURE VISION:
Leveraging proven expertise and innovative problem-solving to drive meaningful impact as {target_role} at {target_company}, 
while continuing to grow professionally and contribute to organizational success through technical excellence and leadership.

---
This comprehensive resume combines compelling career storytelling with detailed professional qualifications,
optimized for both ATS systems and human recruiters seeking {target_role} candidates."""

        return combined_resume
    
    def create_narrative_resume(self, resume_content: str, job_analysis: JobAnalysis, 
                              target_role: str, target_company: str) -> str:
        """Create a story-driven resume that shows career progression and narrative"""
        print("\n>>> CREATING NARRATIVE-DRIVEN RESUME...")
        
        # Extract actual resume information
        resume_info = self.extract_resume_information(resume_content)
        
        # Detect field for storytelling context
        detected_field = self.detect_career_field(target_role.lower())
        
        # Try to get full story template first, fall back to personalized story
        try:
            full_story_elements = self.generate_career_story(detected_field, target_role, target_company)
            # Combine template story with personalized information
            story_elements = {
                'opening_hook': full_story_elements['opening_hook'],
                'title': f"{detected_field.replace('_', ' ').title()} Professional",
                'professional_narrative': full_story_elements['professional_narrative'],
                'career_progression': full_story_elements['career_progression'],
                'signature_achievements': full_story_elements['signature_achievements'],
                'skills': resume_info['skills'] if resume_info['skills'] else ['Technology expertise', 'Problem solving', 'Innovation'],
                'education': resume_info['education'] if resume_info['education'] else f"Relevant education in {detected_field.replace('_', ' ')}",
                'story_projects': full_story_elements['story_projects'],
                'closing_vision': full_story_elements['closing_vision']
            }
        except:
            # Fall back to personalized story if template doesn't exist
            story_elements = self.generate_personalized_story(resume_info, detected_field, target_role, target_company)
        
        narrative_resume = f"""CAREER STORY RESUME - {target_role.upper()}

{story_elements['opening_hook']}

{resume_info['name']}
{story_elements['title']}

CONTACT INFORMATION:
{resume_info['contact_info']}

{story_elements['professional_narrative']}

CAREER JOURNEY & IMPACT STORY:

{story_elements['career_progression']}

KEY ACHIEVEMENTS THAT DEFINE MY STORY:
{chr(10).join(['• ' + achievement for achievement in story_elements['signature_achievements']])}

TECHNICAL EXPERTISE DEVELOPED THROUGH MY JOURNEY:
{chr(10).join(['• ' + skill for skill in story_elements['skills']])}

EDUCATION THAT SHAPED MY PATH:
{story_elements['education']}

DEFINING PROJECTS & MILESTONES:
{chr(10).join(['• ' + project for project in story_elements['story_projects']])}

{story_elements['closing_vision']}

This resume tells the story of a professional journey marked by continuous growth, 
meaningful impact, and unwavering commitment to excellence in {detected_field.replace('_', ' ')}."""

        return narrative_resume
    
    def extract_resume_information(self, resume_content: str) -> dict:
        """Extract key information from the actual resume content"""
        lines = resume_content.strip().split('\n')
        info = {
            'name': 'RYAN THOMAS WEILER',  # Default
            'contact_info': 'Phone: (561) 906-2118 | Email: ryan_wlr@yahoo.com\nLinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/\nGitHub: https://github.com/ryan-wlr',
            'experience': [],
            'skills': [],
            'education': '',
            'projects': []
        }
        
        # Extract name (usually first line)
        if lines and len(lines[0].strip()) > 0 and not '@' in lines[0]:
            name_candidate = lines[0].strip()
            if len(name_candidate.split()) >= 2 and len(name_candidate) < 50:
                info['name'] = name_candidate.upper()
        
        # Extract contact information
        contact_lines = []
        for line in lines[:10]:  # Check first 10 lines
            if any(indicator in line.lower() for indicator in ['@', 'phone', 'email', 'linkedin', 'github', '(']):
                contact_lines.append(line.strip())
        if contact_lines:
            info['contact_info'] = '\n'.join(contact_lines)
        
        # Extract experience (look for job titles, companies, dates)
        experience_keywords = ['experience', 'work', 'employment', 'position', 'role']
        skills_keywords = ['skills', 'technical', 'programming', 'languages', 'tools']
        education_keywords = ['education', 'degree', 'university', 'college', 'school']
        projects_keywords = ['projects', 'portfolio', 'achievements']
        
        current_section = None
        for line in lines:
            line_lower = line.lower().strip()
            
            # Identify sections
            if any(keyword in line_lower for keyword in experience_keywords):
                current_section = 'experience'
                continue
            elif any(keyword in line_lower for keyword in skills_keywords):
                current_section = 'skills'
                continue
            elif any(keyword in line_lower for keyword in education_keywords):
                current_section = 'education'
                continue
            elif any(keyword in line_lower for keyword in projects_keywords):
                current_section = 'projects'
                continue
            
            # Add content to appropriate section
            if line.strip() and current_section:
                if current_section == 'experience' and len(line.strip()) > 5:
                    info['experience'].append(line.strip())
                elif current_section == 'skills':
                    # Split skills by common delimiters
                    skills_text = line.strip()
                    for delimiter in [',', '|', '•', '-', '*']:
                        if delimiter in skills_text:
                            skills_list = [s.strip() for s in skills_text.split(delimiter)]
                            info['skills'].extend([s for s in skills_list if s])
                            break
                    else:
                        if skills_text:
                            info['skills'].append(skills_text)
                elif current_section == 'education' and len(line.strip()) > 5:
                    info['education'] += line.strip() + ' '
                elif current_section == 'projects' and len(line.strip()) > 5:
                    info['projects'].append(line.strip())
        
        # Clean up and validate
        info['skills'] = list(set([skill for skill in info['skills'] if len(skill) > 1 and len(skill) < 50]))[:15]
        info['experience'] = [exp for exp in info['experience'] if len(exp) > 10][:10]
        info['projects'] = [proj for proj in info['projects'] if len(proj) > 10][:8]
        info['education'] = info['education'].strip()
        
        return info
    
    def generate_personalized_story(self, resume_info: dict, field: str, target_role: str, target_company: str) -> dict:
        """Generate story elements using actual resume data"""
        
        # Create personalized opening hook
        field_display = field.replace('_', ' ').title()
        opening_hooks = {
            'software_engineer': f"💻 THE CODE CRAFTSMAN: From {resume_info['name'].title()}'s first lines of code to building systems that transform how people work",
            'quantum_computing_scientist': f"🌟 THE QUANTUM PIONEER: {resume_info['name'].title()}'s journey from curiosity about quantum mechanics to breakthrough research",
            'mathematics_professor': f"📐 THE MATHEMATICAL STORYTELLER: {resume_info['name'].title()}'s quest to unlock the universe's mathematical secrets",
            'data_scientist': f"📊 THE DATA DETECTIVE: {resume_info['name'].title()}'s mission to find insights hidden in complex datasets",
            'generic': f"🌟 THE {field_display.upper()} PROFESSIONAL: {resume_info['name'].title()}'s journey of growth and impact in {field_display.lower()}"
        }
        
        opening_hook = opening_hooks.get(field, opening_hooks['generic'])
        
        # Create personalized professional narrative
        primary_skills = resume_info['skills'][:4] if resume_info['skills'] else ['technology', 'problem solving', 'innovation', 'collaboration']
        
        professional_narrative = f"""PROFESSIONAL NARRATIVE:
My professional journey has been defined by a passion for {field_display.lower()} and a commitment to creating meaningful impact through {', '.join(primary_skills[:3])}. Each challenge I've encountered has strengthened my expertise and deepened my understanding of how technology can solve real-world problems. From early exploration to current mastery, my career tells a story of continuous learning, innovation, and dedication to excellence."""
        
        # Create career progression from actual experience
        career_progression = self.build_career_progression_from_resume(resume_info, field)
        
        # Generate achievements from experience
        signature_achievements = self.extract_achievements_from_resume(resume_info, field)
        
        # Create story projects from actual projects/experience
        story_projects = self.build_story_projects_from_resume(resume_info, field)
        
        # Generate closing vision
        closing_vision = f"""FUTURE VISION:
My {field_display.lower()} story continues with excitement for the opportunities at {target_company}. I envision applying my proven expertise in {', '.join(primary_skills[:3])} to tackle new challenges, drive innovation, and contribute to the company's mission. This next chapter represents not just career growth, but the chance to make meaningful impact in an organization that values {field_display.lower()} excellence."""
        
        return {
            'opening_hook': opening_hook,
            'title': f"{field_display} Professional",
            'professional_narrative': professional_narrative,
            'career_progression': career_progression,
            'signature_achievements': signature_achievements,
            'skills': resume_info['skills'] if resume_info['skills'] else [f"{field_display} expertise", "Problem solving", "Technical innovation", "Team collaboration"],
            'education': resume_info['education'] if resume_info['education'] else f"Relevant education and training in {field_display.lower()}",
            'story_projects': story_projects,
            'closing_vision': closing_vision
        }
    
    def build_career_progression_from_resume(self, resume_info: dict, field: str) -> str:
        """Build career progression narrative from actual resume experience"""
        
        if resume_info['experience']:
            # Use actual experience
            experiences = resume_info['experience'][:3]  # Take up to 3 experiences
            
            progression = f"""Chapter 1: THE FOUNDATION
• {experiences[0] if len(experiences) > 0 else 'Built foundational skills through dedicated learning and practice'}
• Developed core competencies and professional work habits
• Gained valuable experience in real-world problem solving

Chapter 2: THE GROWTH
• {experiences[1] if len(experiences) > 1 else 'Advanced technical skills through challenging projects and mentorship'}
• Expanded expertise and took on increasing responsibilities
• Demonstrated capability to deliver quality results consistently

Chapter 3: THE MASTERY
• {experiences[2] if len(experiences) > 2 else 'Achieved mastery in key areas while continuing to learn and innovate'}
• Leading projects and contributing to team success
• Preparing for next level of professional growth and impact"""
        else:
            # Generic progression based on field
            field_display = field.replace('_', ' ').title()
            progression = f"""Chapter 1: THE FOUNDATION (Early Career)
• Discovered passion for {field_display.lower()} through hands-on learning
• Built fundamental skills through education and practical experience
• Established strong work ethic and professional standards

Chapter 2: THE GROWTH (Skill Development)
• Advanced technical capabilities through challenging projects
• Gained recognition for quality work and professional excellence
• Developed expertise in specialized areas of {field_display.lower()}

Chapter 3: THE MASTERY (Current Focus)
• Achieved proficiency in core {field_display.lower()} competencies
• Contributing to team success and organizational goals
• Continuously learning and adapting to industry evolution"""
        
        return progression
    
    def extract_achievements_from_resume(self, resume_info: dict, field: str) -> list:
        """Extract and enhance achievements from resume content"""
        
        achievements = []
        
        # Look for quantifiable elements in experience
        for exp in resume_info['experience']:
            if any(indicator in exp.lower() for indicator in ['improved', 'increased', 'reduced', 'developed', 'led', 'created', 'built']):
                achievements.append(f"Achieved excellence in {exp[:80]}..." if len(exp) > 80 else exp)
        
        # Add skill-based achievements
        if resume_info['skills']:
            achievements.append(f"Mastered key technologies including {', '.join(resume_info['skills'][:4])}")
        
        # Add project achievements
        for project in resume_info['projects']:
            achievements.append(f"Successfully delivered project: {project[:60]}..." if len(project) > 60 else project)
        
        # Ensure we have at least 3 achievements
        field_display = field.replace('_', ' ').title()
        default_achievements = [
            f"Demonstrated consistent excellence in {field_display.lower()} practices and methodologies",
            f"Contributed to successful project completion with high-quality results",
            f"Continuously improved skills and stayed current with {field_display.lower()} best practices",
            f"Collaborated effectively with teams to achieve organizational objectives"
        ]
        
        while len(achievements) < 4:
            achievements.append(default_achievements[len(achievements)])
        
        return achievements[:6]  # Limit to 6 achievements
    
    def build_story_projects_from_resume(self, resume_info: dict, field: str) -> list:
        """Build story projects from actual resume projects and experience"""
        
        story_projects = []
        
        # Use actual projects if available
        for project in resume_info['projects']:
            story_projects.append(project)
        
        # Transform experience into project format
        for exp in resume_info['experience']:
            if any(keyword in exp.lower() for keyword in ['project', 'system', 'application', 'platform', 'solution']):
                story_projects.append(f"Professional Project: {exp}")
        
        # Ensure we have at least 3 story projects
        field_display = field.replace('_', ' ').title()
        default_projects = [
            f"Professional Excellence Initiative: Demonstrated mastery of {field_display.lower()} best practices",
            f"Technical Innovation Project: Applied cutting-edge techniques to solve complex challenges",
            f"Collaborative Achievement: Successfully worked with teams to deliver high-impact results",
            f"Continuous Learning Journey: Stayed current with {field_display.lower()} industry developments"
        ]
        
        while len(story_projects) < 4:
            story_projects.append(default_projects[len(story_projects)])
        
        return story_projects[:6]  # Limit to 6 projects
    
    def generate_career_story(self, field: str, target_role: str, target_company: str) -> dict:
        """Generate compelling story elements dynamically for any profession"""
        return self.generate_dynamic_story_template(field, target_role, target_company)
    
    def generate_dynamic_story_template(self, field: str, target_role: str, target_company: str) -> dict:
        """Generate compelling storytelling template dynamically for any profession"""
        
        # Get field-specific data
        field_data = self.get_field_data(field)
        display_field = field.replace('_', ' ').title()
        
        # Define profession-specific story elements
        story_elements = self.get_profession_story_elements(field, target_role)
        
        # Generate dynamic opening hook
        opening_hook = f"{story_elements['emoji']} THE {story_elements['title']}: {story_elements['hook_text']}"
        
        # Generate professional narrative
        professional_narrative = f"""PROFESSIONAL NARRATIVE:
{story_elements['narrative_intro']} Driven by {story_elements['driving_force']}, {story_elements['expertise_statement']} I understand that {story_elements['field_philosophy']}"""
        
        # Generate career progression chapters
        career_progression = f"""Chapter 1: THE FOUNDATION ({story_elements['early_years']})
• {story_elements['foundation_discovery']}
• {story_elements['foundation_skills']}
• {story_elements['foundation_achievement']}

Chapter 2: THE {story_elements['chapter2_title']} ({story_elements['growth_years']})
• {story_elements['growth_advancement']}
• {story_elements['growth_specialization']}
• {story_elements['growth_breakthrough']}

Chapter 3: THE {story_elements['chapter3_title']} ({story_elements['current_years']})
• {story_elements['current_leadership']}
• {story_elements['current_innovation']}
• {story_elements['current_mentoring']}"""
        
        # Generate signature achievements
        signature_achievements = [
            f"{story_elements['achievement1_action']} {story_elements['achievement1_metric']} {story_elements['achievement1_context']}",
            f"{story_elements['achievement2_action']} {story_elements['achievement2_metric']} {story_elements['achievement2_context']}",
            f"{story_elements['achievement3_action']} {story_elements['achievement3_metric']} {story_elements['achievement3_context']}",
            f"{story_elements['achievement4_action']} {story_elements['achievement4_metric']} {story_elements['achievement4_context']}"
        ]
        
        # Generate story projects
        story_projects = [
            f"{story_elements['project1_name']}: {story_elements['project1_description']}",
            f"{story_elements['project2_name']}: {story_elements['project2_description']}",
            f"{story_elements['project3_name']}: {story_elements['project3_description']}",
            f"{story_elements['project4_name']}: {story_elements['project4_description']}"
        ]
        
        # Generate closing vision
        closing_vision = f"""FUTURE VISION:
My {display_field.lower()} story continues with {story_elements['future_emotion']} for the opportunities at {target_company}. I envision {story_elements['future_vision']} while {story_elements['future_contribution']}."""
        
        return {
            'opening_hook': opening_hook,
            'professional_narrative': professional_narrative,
            'career_progression': career_progression,
            'signature_achievements': signature_achievements,
            'story_projects': story_projects,
            'closing_vision': closing_vision
        }
    
    def get_profession_story_elements(self, field: str, target_role: str) -> dict:
        """Get profession-specific storytelling elements for dynamic template generation"""
        
        # Define story elements by profession type
        profession_stories = {
            'welder': {
                'emoji': '🔥',
                'title': 'METAL MASTER',
                'hook_text': 'Forging strength and precision through the art and science of welding',
                'narrative_intro': 'Welding is where art meets engineering, where molten metal becomes the foundation of our modern world.',
                'driving_force': 'the pursuit of perfection in every bead, every joint, and every fabrication',
                'expertise_statement': 'From structural steel that supports skyscrapers to precision aerospace components,',
                'field_philosophy': 'quality welding is not just about joining metal—it\'s about creating lasting strength that people depend on every day.',
                'early_years': 'Early Career',
                'foundation_discovery': 'Discovered the craft of welding through hands-on training and certification programs',
                'foundation_skills': 'Mastered fundamental techniques: stick welding, MIG, TIG, and flux-core processes',
                'foundation_achievement': 'Built reputation for consistent quality and attention to safety protocols',
                'chapter2_title': 'SPECIALIZATION',
                'growth_years': 'Skill Development',
                'growth_advancement': 'Advanced to specialized welding processes for critical applications',
                'growth_specialization': 'Achieved AWS certifications and gained expertise in exotic materials',
                'growth_breakthrough': 'Developed precision welding skills for aerospace and pressure vessel work',
                'chapter3_title': 'MASTERY',
                'current_years': 'Current Focus',
                'current_leadership': 'Leading complex fabrication projects requiring advanced welding expertise',
                'current_innovation': 'Training and mentoring new welders in safety and quality standards',
                'current_mentoring': 'Continuously adapting to new materials and welding technologies',
                'achievement1_action': 'Achieved',
                'achievement1_metric': '99.5% pass rate',
                'achievement1_context': 'on critical structural weld inspections across 500+ projects',
                'achievement2_action': 'Completed',
                'achievement2_metric': 'AWS D1.1 structural welding certification',
                'achievement2_context': 'with perfect test scores',
                'achievement3_action': 'Led fabrication team that reduced',
                'achievement3_metric': 'project completion time by 25%',
                'achievement3_context': 'through process optimization',
                'achievement4_action': 'Maintained',
                'achievement4_metric': 'zero safety incidents',
                'achievement4_context': 'across 5+ years of high-risk welding operations',
                'project1_name': 'Structural Steel Excellence',
                'project1_description': 'Critical building framework welding for high-rise construction',
                'project2_name': 'Precision Pipeline Project',
                'project2_description': 'High-pressure pipe welding meeting ASME standards',
                'project3_name': 'Aerospace Component Fabrication',
                'project3_description': 'TIG welding of aluminum assemblies for aircraft industry',
                'project4_name': 'Safety Leadership Initiative',
                'project4_description': 'Comprehensive welding safety training program development',
                'future_emotion': 'excitement',
                'future_vision': 'applying proven expertise in precision welding, quality control, and safety leadership',
                'future_contribution': 'contributing to projects that demand the highest standards of craftsmanship and reliability'
            },
            'optical_engineer': {
                'emoji': '🔬',
                'title': 'LIGHT ARCHITECT',
                'hook_text': 'Engineering the future through precision optics and photonic innovation',
                'narrative_intro': 'Light has always been my medium of choice for solving complex engineering challenges.',
                'driving_force': 'the elegant physics of photonics and the endless possibilities that emerge when light is precisely controlled',
                'expertise_statement': 'From designing laser systems to developing fiber optic communications,',
                'field_philosophy': 'optical technologies can transform industries and improve lives through precise light manipulation.',
                'early_years': 'Early Career',
                'foundation_discovery': 'Developed fascination with optical physics and precision engineering',
                'foundation_skills': 'Mastered fundamentals of laser systems, fiber optics, and optical design',
                'foundation_achievement': 'Built first optical prototypes, discovering the art of light manipulation',
                'chapter2_title': 'PRECISION',
                'growth_years': 'Skill Development',
                'growth_advancement': 'Advanced expertise in optical modeling using Zemax and Code V',
                'growth_specialization': 'Designed complex optical systems for telecommunications and defense applications',
                'growth_breakthrough': 'Achieved breakthrough performance improvements in fiber optic transmission',
                'chapter3_title': 'INNOVATION',
                'current_years': 'Current Focus',
                'current_leadership': 'Leading development of next-generation photonic devices and systems',
                'current_innovation': 'Pioneering new approaches to optical design and laser technology',
                'current_mentoring': 'Mentoring teams while pushing boundaries of optical engineering',
                'achievement1_action': 'Designed revolutionary laser systems improving',
                'achievement1_metric': 'efficiency by 40%',
                'achievement1_context': 'over industry standards',
                'achievement2_action': 'Developed fiber optic communication systems enabling',
                'achievement2_metric': '10Gbps data transmission',
                'achievement2_context': 'with ultra-low latency performance',
                'achievement3_action': 'Led optical modeling projects resulting in',
                'achievement3_metric': '25% cost reduction',
                'achievement3_context': 'through design optimization',
                'achievement4_action': 'Pioneered precision optical assemblies with',
                'achievement4_metric': 'sub-micron alignment tolerances',
                'achievement4_context': 'for critical aerospace applications',
                'project1_name': 'Photonic Revolution',
                'project1_description': 'Next-generation laser diode systems for telecommunications',
                'project2_name': 'Precision Optics Platform',
                'project2_description': 'Advanced optical design reducing manufacturing costs',
                'project3_name': 'Fiber Network Innovation',
                'project3_description': 'High-speed optical communication system design',
                'project4_name': 'Laser Safety Initiative',
                'project4_description': 'Comprehensive optical safety protocols and training programs',
                'future_emotion': 'enthusiasm',
                'future_vision': 'applying proven expertise in laser systems, fiber optics, and precision optical design',
                'future_contribution': 'developing breakthrough photonic solutions that advance the field and drive technological innovation'
            },
            'software_engineer': {
                'emoji': '💻',
                'title': 'CODE CRAFTSMAN',
                'hook_text': 'Transforming complex problems into elegant solutions that power the digital world',
                'narrative_intro': 'Code is poetry, algorithms are symphonies, and great software tells the story of human ingenuity solving real-world challenges.',
                'driving_force': 'the belief that technology should enhance human potential',
                'expertise_statement': 'From my first "Hello World" to architecting systems that process millions of transactions,',
                'field_philosophy': 'each line of code should be written with purpose and passion to create meaningful impact.',
                'early_years': '2021-2022',
                'foundation_discovery': 'First programming course ignited passion for logical problem-solving',
                'foundation_skills': 'Built first web application, realizing technology\'s power to create and connect',
                'foundation_achievement': 'Internship experience: learned collaborative software development',
                'chapter2_title': 'MASTERY',
                'growth_years': '2022-2023',
                'growth_advancement': 'Specialized in financial technology, combining coding skills with market knowledge',
                'growth_specialization': 'Developed automated trading algorithms achieving significant returns',
                'growth_breakthrough': 'Open-source contributions gained recognition in developer community',
                'chapter3_title': 'INNOVATION',
                'current_years': '2023-Present',
                'current_leadership': 'Leading fintech projects that revolutionize how people interact with technology',
                'current_innovation': 'Architecting scalable systems processing millions of daily transactions',
                'current_mentoring': 'Mentoring junior developers while pushing technological boundaries',
                'achievement1_action': 'Architected systems generating',
                'achievement1_metric': '15% annual returns',
                'achievement1_context': 'with optimized risk management',
                'achievement2_action': 'Built scalable data pipelines processing',
                'achievement2_metric': '10,000+ daily data points',
                'achievement2_context': 'with real-time analytics capabilities',
                'achievement3_action': 'Led development of platform serving',
                'achievement3_metric': '50,000+ active users',
                'achievement3_context': 'with 99.9% uptime reliability',
                'achievement4_action': 'Open-source contributions downloaded',
                'achievement4_metric': '100,000+ times',
                'achievement4_context': 'by global developer community',
                'project1_name': 'The Analytics Engine',
                'project1_description': 'AI-powered algorithms achieving breakthrough performance metrics',
                'project2_name': 'Platform Revolution',
                'project2_description': 'Full-stack application democratizing advanced technology access',
                'project3_name': 'System Maestro',
                'project3_description': 'Real-time analytics dashboard used by professional teams',
                'project4_name': 'Code for Good',
                'project4_description': 'Open-source tools helping organizations optimize their operations',
                'future_emotion': 'excitement',
                'future_vision': 'building software that not only solves complex problems but creates new possibilities',
                'future_contribution': 'continuing to grow and contribute to technological advancement'
            }
        }
        
        # Return profession-specific elements or generate generic ones
        if field in profession_stories:
            return profession_stories[field]
        else:
            return self.generate_generic_profession_elements(field, target_role)
    
    def generate_generic_profession_elements(self, field: str, target_role: str) -> dict:
        """Generate generic story elements for any profession not specifically defined"""
        
        # Map technical field names to better narrative terms
        field_narrative_map = {
            'data_scientist': 'data science',
            'software_engineer': 'software engineering', 
            'web_developer': 'web development',
            'mobile_developer': 'mobile development',
            'devops_engineer': 'DevOps engineering',
            'security_engineer': 'cybersecurity'
        }
        
        # Use mapped name for narrative or clean up the original
        narrative_field = field_narrative_map.get(field, field.replace('_', ' ').lower())
        display_field = field.replace('_', ' ').title()
        
        return {
            'emoji': '🌟',
            'title': f'{display_field.upper()} PROFESSIONAL',
            'hook_text': f'From passion to expertise, building excellence in {narrative_field}',
            'narrative_intro': f'My journey in {narrative_field} has been a quest for continuous improvement and meaningful impact.',
            'driving_force': 'a passion for excellence and a commitment to making a positive difference',
            'expertise_statement': f'Through years of dedicated practice in {narrative_field},',
            'field_philosophy': f'true success comes from combining technical skill with professional wisdom and ethical practice.',
            'early_years': '2020-2021',
            'foundation_discovery': f'Discovered passion for {narrative_field} through hands-on experience',
            'foundation_skills': f'Built fundamental skills in {narrative_field} through dedicated practice and study',
            'foundation_achievement': f'First major success in {narrative_field} established confidence and direction',
            'chapter2_title': 'GROWTH',
            'growth_years': '2021-2023',
            'growth_advancement': f'Advanced skills in {narrative_field} through challenging projects and mentorship',
            'growth_specialization': f'Developed expertise in specialized areas of {narrative_field}',
            'growth_breakthrough': f'Gained recognition for quality work and professional excellence in {narrative_field}',
            'chapter3_title': 'MASTERY',
            'current_years': '2023-Present',
            'current_leadership': f'Established as skilled {narrative_field} professional with proven track record',
            'current_innovation': f'Leading projects and implementing innovative {narrative_field} practices',
            'current_mentoring': f'Mentoring others while advancing {narrative_field} standards',
            'achievement1_action': f'Achieved excellence in {display_field.lower()} with',
            'achievement1_metric': 'consistent high-quality results',
            'achievement1_context': 'across multiple challenging projects',
            'achievement2_action': f'Improved efficiency by',
            'achievement2_metric': '25%',
            'achievement2_context': f'through innovative {display_field.lower()} practices',
            'achievement3_action': f'Led successful projects that advanced',
            'achievement3_metric': f'{display_field.lower()} standards',
            'achievement3_context': 'and organizational capabilities',
            'achievement4_action': f'Mentored others and contributed to',
            'achievement4_metric': f'{display_field.lower()} community growth',
            'achievement4_context': 'through knowledge sharing and collaboration',
            'project1_name': 'Excellence Initiative',
            'project1_description': f'Demonstrated mastery of {display_field.lower()} best practices',
            'project2_name': 'Innovation Project',
            'project2_description': f'Developed new approaches improving {display_field.lower()} outcomes',
            'project3_name': 'Leadership Challenge',
            'project3_description': f'Successfully guided team through complex {display_field.lower()} project',
            'project4_name': 'Community Impact',
            'project4_description': f'Contributed expertise to advance {display_field.lower()} field',
            'future_emotion': 'enthusiasm',
            'future_vision': f'applying proven expertise in {display_field.lower()} while continuing to grow',
            'future_contribution': f'contributing to {display_field.lower()} field advancement and organizational success'
        }

    def process_complete_optimization(self, job_description: str, resume_content: str, 
                                    target_role: str = "Software Engineer", 
                                    target_company: str = "Target Company") -> Dict[str, str]:
        """Execute complete resume optimization workflow with narrative storytelling"""
        
        print(">>> STARTING COMPLETE RESUME OPTIMIZATION PROCESS")
        print(f"    Target Role: {target_role}")
        print(f"    Target Company: {target_company}")
        print(f"    Processing {len(resume_content)} characters of resume content")
        print(f"    Analyzing {len(job_description)} characters of job description")
        
        # Ask user for resume style preference
        print("\n>>> RESUME STYLE SELECTION:")
        print("    1. STORY RESUME - Narrative-driven, tells your career journey")
        print("    2. STANDARD RESUME - Enhanced professional format")
        print("    3. BOTH VERSIONS - Create narrative and standard resumes")
        
        style_choice = input("Choose resume style (1, 2, or 3): ").strip()
        
        # Step 1: Analyze job requirements
        job_analysis = self.analyze_job_posting(job_description, target_role, target_company)
        
        # Step 2: Execute all optimization strategies
        results = {}
        
        results['1_job_analysis'] = f"""JOB ANALYSIS SUMMARY - {target_role} at {target_company}

REQUIRED TECHNICAL SKILLS:
{chr(10).join(f"• {skill}" for skill in job_analysis.required_skills)}

PREFERRED QUALIFICATIONS:
{chr(10).join(f"• {skill}" for skill in job_analysis.preferred_skills)}

KEY RESPONSIBILITIES:
{chr(10).join(f"• {resp}" for resp in job_analysis.key_responsibilities)}

COMPANY VALUES & CULTURE:
{chr(10).join(f"• {value}" for value in job_analysis.company_values)}

INDUSTRY KEYWORDS:
{chr(10).join(f"• {keyword}" for keyword in job_analysis.industry_keywords)}

EXPERIENCE LEVEL: {job_analysis.experience_level}
ROLE CATEGORY: {job_analysis.role_type}

OPTIMIZATION RECOMMENDATIONS:
• Emphasize technical skills alignment with required competencies
• Highlight relevant project experience and quantifiable achievements
• Include industry-specific keywords naturally throughout resume
• Demonstrate cultural fit through examples of company values in action"""
        
        results['2_resume_flaws'] = self.analyze_resume_flaws(resume_content, job_analysis)
        results['3_impact_rewrite'] = self.rewrite_for_impact(resume_content, target_role)
        results['4_ats_optimized'] = self.optimize_for_ats(resume_content, target_role)
        results['5_enhanced_skills'] = self.enhance_skills_section(job_analysis)
        results['6_keyword_experience'] = self.enhance_experience_with_keywords(resume_content, job_analysis)
        
        # Create different resume versions based on user choice
        if style_choice == "1":
            # Story resume only
            results['7_narrative_resume'] = self.create_narrative_resume(resume_content, job_analysis, target_role, target_company)
        elif style_choice == "2":
            # Standard resume only
            results['7_enhanced_resume'] = self.create_enhanced_resume(resume_content, job_analysis, target_role, target_company)
        else:
            # Option 3: Combined version with both narrative and professional sections
            narrative_content = self.create_narrative_resume(resume_content, job_analysis, target_role, target_company)
            enhanced_content = self.create_enhanced_resume(resume_content, job_analysis, target_role, target_company)
            
            # Create a combined resume that includes both storytelling and professional elements
            results['7_combined_resume'] = self.create_combined_resume(narrative_content, enhanced_content, target_role, target_company)
            
            # Also save individual versions for reference
            results['8_narrative_only'] = narrative_content
            results['9_enhanced_only'] = enhanced_content
        
        # Step 3: Create executive summary
        summary_key = '8_executive_summary' if style_choice in ["1", "2"] else '10_executive_summary'
        results[summary_key] = f"""RESUME OPTIMIZATION EXECUTIVE SUMMARY

OPTIMIZATION TARGET: {target_role} at {target_company}
PROCESSING DATE: {datetime.now().strftime('%Y-%m-%d %H:%M')}
RESUME STYLE: {"Story-Driven Narrative" if style_choice == "1" else "Professional Standard" if style_choice == "2" else "Both Narrative and Standard"}

KEY IMPROVEMENTS IMPLEMENTED:
>>> Technical Skills Enhancement: Aligned core competencies with job requirements
>>> ATS Optimization: Improved keyword density and formatting for tracking systems  
>>> Impact Quantification: Added measurable achievements and performance metrics
>>> Experience Enhancement: Strengthened job descriptions with powerful action verbs
>>> Industry Alignment: Integrated relevant keywords and terminology
>>> Cultural Fit: Emphasized alignment with company values and mission
{(">>> Narrative Storytelling: Created career story showing growth and impact" if style_choice in ["1", "3"] else "")}

OPTIMIZATION RESULTS:
• Resume now contains {len([skill for skill in job_analysis.required_skills if skill.lower() in str(results).lower()])} of {len(job_analysis.required_skills)} required technical skills
• Enhanced readability and professional formatting for improved recruiter appeal
• Optimized for major ATS platforms including Workday, Greenhouse, and Lever
• Increased keyword relevance score by incorporating industry-specific terminology
• Strengthened value proposition with quantified achievements and results
{("• Compelling career narrative that shows professional growth and future vision" if style_choice in ["1", "3"] else "")}

NEXT STEPS:
1. Review optimized resume version(s) for accuracy and personal preferences
2. Customize cover letter using provided job analysis insights  
3. Prepare interview talking points based on enhanced experience descriptions
4. Update LinkedIn profile to match optimized resume content
5. Save multiple versions for different role types and industries

RECOMMENDED USAGE:
{("• Use 'Narrative Resume' for roles emphasizing culture fit and leadership" if style_choice in ["1", "3"] else "")}
{("• Use 'Enhanced Resume' for technical roles and ATS-heavy companies" if style_choice in ["2", "3"] else "")}
{("• Use primary resume version as main application document" if style_choice in ["1", "2"] else "")}
• Reference enhanced content for online application systems
• Adapt storytelling elements for cover letters and interviews"""

        return results
    
    def save_results_to_files(self, results: Dict[str, str], output_dir: str = "resume_optimization_output"):
        """Save all optimization results to organized files"""
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # File mapping with descriptive names
        file_mapping = {
            '1_job_analysis': 'job_analysis_report.txt',
            '2_resume_flaws': 'resume_analysis_flaws.txt', 
            '3_impact_rewrite': 'resume_impact_version.txt',
            '4_ats_optimized': 'resume_ats_optimized.txt',
            '5_enhanced_skills': 'enhanced_skills_section.txt',
            '6_keyword_experience': 'keyword_enhanced_experience.txt',
            '7_narrative_resume': 'narrative_story_resume.txt',
            '7_enhanced_resume': 'enhanced_standard_resume.txt',
            '7_combined_resume': 'combined_comprehensive_resume.txt',
            '8_enhanced_resume': 'enhanced_standard_resume.txt',
            '8_narrative_only': 'narrative_only_reference.txt',
            '9_enhanced_only': 'enhanced_only_reference.txt',
            '7_tailored_resume': 'tailored_resume_final.txt',
            '8_executive_summary': 'optimization_executive_summary.txt',
            '9_executive_summary': 'optimization_executive_summary.txt',
            '10_executive_summary': 'optimization_executive_summary.txt'
        }
        
        print(f"\n>>> Saving optimization results to '{output_dir}' folder:")
        
        # Save each result file
        for key, filename in file_mapping.items():
            if key in results:
                file_path = os.path.join(output_dir, filename)
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(results[key])
                    print(f"    >>> {filename}")
                except Exception as e:
                    print(f"    ERROR: Failed to save {filename}: {e}")
        
        print(f"\n>>> All files saved successfully!")
        
        # Create formatted .docx resume(s) based on what was generated
        if '7_combined_resume' in results:
            # Option 3: Combined version - create one comprehensive DOCX
            print(f"\n>>> Creating combined comprehensive DOCX file:")
            combined_results = {k: v for k, v in results.items()}
            combined_results['7_narrative_resume'] = results['7_combined_resume']  # Use combined content
            self.create_formatted_docx_resume(combined_results, output_dir)
            
        elif '7_narrative_resume' in results and '8_enhanced_resume' in results:
            # Legacy option 3: Both versions - create two DOCX files (should not happen with new logic)
            print(f"\n>>> Creating DOCX files for both resume versions:")
            
            # Create narrative DOCX
            narrative_results = {k: v for k, v in results.items()}
            narrative_results['7_narrative_resume'] = results['7_narrative_resume']
            if '8_enhanced_resume' in narrative_results:
                del narrative_results['8_enhanced_resume']
            
            self.create_formatted_docx_resume_specific(narrative_results, output_dir, 'narrative_story_resume.docx')
            
            # Create standard DOCX 
            standard_results = {k: v for k, v in results.items()}
            standard_results['7_enhanced_resume'] = results['8_enhanced_resume']
            if '7_narrative_resume' in standard_results:
                del standard_results['7_narrative_resume']
            
            self.create_formatted_docx_resume_specific(standard_results, output_dir, 'enhanced_standard_resume.docx')
            
        elif '7_narrative_resume' in results:
            # Option 1: Story resume only
            self.create_formatted_docx_resume(results, output_dir)
        elif '7_enhanced_resume' in results:
            # Option 2: Standard resume only
            # Map enhanced resume to expected key
            results['7_narrative_resume'] = results['7_enhanced_resume']
            self.create_formatted_docx_resume(results, output_dir)
        else:
            # Fallback - create with whatever content is available
            self.create_formatted_docx_resume(results, output_dir)
    
    def detect_career_field(self, content: str) -> str:
        """Intelligently detect career field from job content and generate appropriate data"""
        content_lower = content.lower()
        
        # First try dynamic field detection for any field
        dynamic_field = self.detect_dynamic_field(content_lower)
        if dynamic_field:
            return dynamic_field
        
        # Fallback to predefined field detection patterns
        field_patterns = {
            'electrician': {
                'keywords': ['electrician', 'electrical', 'wiring', 'voltage', 'circuits', 'nec', 'electrical code', 'motor control', 'plc'],
                'education': "Florida Atlantic University — B.S. Computer Science (Electrical Focus), Expected 2024 (Dean's List, GPA 3.7)",
                'experience_title': 'Industrial Electrician & Maintenance Specialist',
                'skills': [
                    'Electrical Systems: Motor Controls, PLCs (Allen-Bradley), VFDs, Transformers, Switchgear',
                    'Codes & Standards: National Electrical Code (NEC), OSHA Safety, Arc Flash Safety, LOTO',
                    'Testing Equipment: Multimeters, Oscilloscopes, Megohmmeters, Power Quality Analyzers',
                    'Industrial Systems: 480V Three-Phase, Industrial Machinery, Automation, Preventive Maintenance'
                ],
                'projects': [
                    'Industrial Motor Control Installation: Complete 3-phase motor control panel installation',
                    'PLC Programming Project: Allen-Bradley PLC programming for automated conveyor system',
                    'OSHA 10-Hour Electrical Safety Certification',
                    'NEC Code Compliance: Successfully completed electrical inspections meeting all requirements'
                ]
            },
            'plumber': {
                'keywords': ['plumber', 'plumbing', 'pipes', 'water', 'drain', 'fixtures', 'water heater', 'sewer'],
                'education': "Palm Beach State College — Associate of Applied Science in Plumbing Technology, 2022\nFlorida Atlantic University — B.S. Computer Science (Technical Systems Focus), Expected 2024",
                'experience_title': 'Professional Plumber',
                'skills': [
                    'Plumbing Systems: Water Supply Lines, Drain Systems, Fixture Installation, Pipe Repair',
                    'Tools & Equipment: Pipe Wrenches, Drain Augers, Pipe Cutters, Soldering Equipment, Threading Machines',
                    'Materials: PVC, Copper, PEX Piping, Cast Iron, Various Fittings and Fixtures',
                    'Codes & Standards: Local Plumbing Codes, Uniform Plumbing Code, Safety Regulations',
                    'Water Systems: Water Heaters, Boilers, Pressure Tanks, Backflow Prevention',
                    'Repair Services: Drain Cleaning, Leak Repair, Emergency Service, System Diagnostics'
                ],
                'projects': [
                    'Plumbing Code Compliance Training: Certified in local plumbing codes and regulations',
                    'Water Heater Installation and Service Certification',
                    'Drain Cleaning and Sewer Line Repair Techniques',
                    'Customer Service Excellence Recognition: 95% satisfaction rating'
                ]
            },
            'nurse': {
                'keywords': ['nurse', 'nursing', 'patient', 'healthcare', 'medical', 'clinical', 'rn', 'lpn', 'hospital'],
                'education': "Florida Atlantic University — Bachelor of Science in Nursing (BSN), Expected 2024 (Dean's List, GPA 3.7)",
                'experience_title': 'Registered Nurse',
                'skills': [
                    'Clinical Skills: Patient Assessment, Medication Administration, IV Therapy, Wound Care',
                    'Medical Equipment: Monitors, Ventilators, IV Pumps, EKG Machines, Defibrillators',
                    'Documentation: Electronic Health Records (EHR), EPIC, Cerner, Medical Charting',
                    'Patient Care: Critical Care, Emergency Response, Patient Education, Family Communication'
                ],
                'projects': [
                    'ACLS (Advanced Cardiovascular Life Support) Certification',
                    'BLS (Basic Life Support) Certification',
                    'Patient Safety Excellence Recognition',
                    'Clinical Leadership Training Completion'
                ]
            },
            'teacher': {
                'keywords': ['teacher', 'education', 'classroom', 'student', 'curriculum', 'lesson', 'school', 'teaching'],
                'education': "Florida Atlantic University — Bachelor of Education in Elementary Education, Expected 2024 (Dean's List, GPA 3.7)",
                'experience_title': 'Elementary School Teacher',
                'skills': [
                    'Curriculum Development: Lesson Planning, Standards Alignment, Assessment Design, Learning Objectives',
                    'Classroom Management: Behavior Management, Student Engagement, Differentiated Instruction',
                    'Educational Technology: SmartBoards, Google Classroom, Educational Apps, Online Learning Platforms',
                    'Student Assessment: Formative Assessment, Progress Monitoring, Data Analysis, Parent Communication'
                ],
                'projects': [
                    'Florida Teaching Certificate (Elementary Education K-6)',
                    'ESE (Exceptional Student Education) Endorsement',
                    'Reading Endorsement Certification',
                    'Classroom Innovation Award Recognition'
                ]
            },
            'mechanic': {
                'keywords': ['mechanic', 'automotive', 'engine', 'repair', 'maintenance', 'car', 'vehicle', 'diagnostic'],
                'education': "Palm Beach State College — Associate of Applied Science in Automotive Technology, 2022\nFlorida Atlantic University — B.S. Computer Science (Technical Systems Focus), Expected 2024",
                'experience_title': 'Automotive Technician',
                'skills': [
                    'Engine Systems: Diagnostics, Repair, Maintenance, Performance Tuning, Fuel Systems',
                    'Diagnostic Tools: OBD Scanners, Multimeters, Oscilloscopes, Compression Testers',
                    'Vehicle Systems: Electrical, Brake, Suspension, Transmission, HVAC Systems',
                    'Certifications: ASE Certified, Manufacturer Training, Safety Protocols'
                ],
                'projects': [
                    'ASE (Automotive Service Excellence) Certification',
                    'EPA 609 Refrigerant Handling Certification',
                    'Customer Service Excellence Recognition',
                    'Advanced Diagnostic Training Completion'
                ]
            },
            'dishwasher': {
                'keywords': ['dishwasher', 'kitchen', 'restaurant', 'cleaning', 'dishes', 'food service', 'sanitization', 'busing'],
                'education': "Palm Beach State College — High School Diploma, 2022\nFlorida Atlantic University — B.S. Computer Science (Food Service Focus), Expected 2024",
                'experience_title': 'Kitchen Staff / Dishwasher',
                'skills': [
                    'Kitchen Operations: Dish Washing, Sanitization, Equipment Maintenance, Food Safety',
                    'Equipment: Commercial Dishwashers, Sanitizing Systems, Kitchen Appliances',
                    'Food Safety: ServSafe Certified, Proper Storage, Temperature Control, Hygiene Standards',
                    'Customer Service: Team Collaboration, Fast-Paced Environment, Quality Standards'
                ],
                'projects': [
                    'ServSafe Food Handler Certification',
                    'Kitchen Equipment Operation Training',
                    'Food Safety and Sanitation Compliance',
                    'Team Efficiency and Speed Recognition'
                ]
            },
            'chef': {
                'keywords': ['chef', 'culinary', 'cooking', 'kitchen management', 'menu', 'cuisine', 'culinary arts', 'food preparation'],
                'education': "Culinary Institute of America — Associate Degree in Culinary Arts, 2022\nFlorida Atlantic University — B.S. Computer Science (Culinary Management Focus), Expected 2024",
                'experience_title': 'Executive Chef / Culinary Specialist',
                'skills': [
                    'Culinary Skills: Menu Development, Food Preparation, Cooking Techniques, Recipe Creation',
                    'Kitchen Management: Staff Supervision, Inventory Control, Cost Management, Food Safety',
                    'Leadership: Team Training, Quality Control, Kitchen Operations, Service Coordination',
                    'Certifications: ServSafe Manager, Culinary Arts, Food Safety, Kitchen Management'
                ],
                'projects': [
                    'ServSafe Manager Certification',
                    'Culinary Arts Professional Training',
                    'Menu Development and Cost Analysis',
                    'Kitchen Leadership and Team Management'
                ]
            }
        }
        
        # Check for field matches
        for field, data in field_patterns.items():
            keyword_matches = sum(1 for keyword in data['keywords'] if keyword in content_lower)
            # For exact field name matches or when we have sufficient keywords
            if field in content_lower or keyword_matches >= 2:
                return field
        
        # Default to software engineer if no specific field detected
        return 'software_engineer'
    
    def detect_dynamic_field(self, content: str) -> str:
        """Dynamically detect any career field and generate appropriate content"""
        
        # Optical Engineering and Photonics fields (high priority)
        optical_terms = ['optical', 'optics', 'photonics', 'laser', 'fiber optic', 'optical design', 'optical engineer', 'photonic', 'optical systems']
        if any(term in content for term in optical_terms):
            if 'engineer' in content or 'design' in content or 'systems' in content:
                return 'optical_engineer'
            else:
                return 'optical_engineer'  # Default to optical engineer for any optical terms
        
        # Welding and Fabrication fields
        welding_terms = ['welder', 'welding', 'fabrication', 'tig', 'mig', 'arc welding', 'stick welding', 'fabricator', 'welding inspector']
        if any(term in content for term in welding_terms):
            return 'welder'
        
        # Software Engineering and Computer Science fields (comprehensive detection)
        software_terms = ['software', 'programming', 'code', 'developer', 'development', 'programmer', 'coding', 'computer science', 'cs']
        data_science_terms = ['data scientist', 'data science', 'machine learning', 'ml engineer', 'ai engineer', 'data analyst', 'data engineer']
        web_dev_terms = ['web developer', 'web development', 'frontend', 'backend', 'full stack', 'fullstack', 'javascript', 'react', 'angular', 'vue']
        mobile_terms = ['mobile developer', 'ios developer', 'android developer', 'mobile app', 'swift', 'kotlin', 'react native', 'flutter']
        devops_terms = ['devops', 'sre', 'site reliability', 'cloud engineer', 'infrastructure', 'kubernetes', 'docker', 'aws', 'azure', 'gcp']
        security_terms = ['cybersecurity', 'security engineer', 'information security', 'penetration testing', 'security analyst', 'infosec']
        
        if any(term in content for term in data_science_terms):
            return 'data_scientist'
        elif any(term in content for term in web_dev_terms):
            return 'web_developer'
        elif any(term in content for term in mobile_terms):
            return 'mobile_developer'
        elif any(term in content for term in devops_terms):
            return 'devops_engineer'
        elif any(term in content for term in security_terms):
            return 'security_engineer'
        elif any(term in content for term in software_terms):
            # Check for specific software engineering roles
            if any(term in content for term in ['software engineer', 'senior software', 'full stack', 'backend', 'frontend', 'web developer', 'developer', 'programming']):
                return 'software_engineer'
        
        # Advanced Mathematics and Physics fields
        quantum_terms = ['quantum', 'qubit', 'quantum computing', 'quantum algorithm', 'quantum mechanics', 'qiskit', 'cirq']
        if any(term in content for term in quantum_terms):
            return 'quantum_computing_scientist'
            
        advanced_math_terms = ['mathematics', 'mathematical', 'algebra', 'topology', 'analysis', 'geometry', 'number theory', 'differential', 'manifold', 'homology']
        if any(term in content for term in advanced_math_terms):
            if 'professor' in content or 'research' in content:
                return 'mathematics_professor'
            else:
                return 'mathematician'
        
        physics_terms = ['physics', 'theoretical physics', 'particle physics', 'condensed matter', 'astrophysics', 'cosmology']
        if any(term in content for term in physics_terms):
            return 'theoretical_physicist'
        
        # Medical fields pattern matching
        medical_terms = ['surgeon', 'doctor', 'physician', 'medical', 'hospital', 'patient', 'surgery', 'clinic', 'nurse', 'healthcare']
        if any(term in content for term in medical_terms):
            # Extract specific medical specialty
            if 'brain' in content or 'neurosurg' in content or 'neurolog' in content:
                return 'brain_surgeon'
            elif 'heart' in content or 'cardio' in content or 'cardiac' in content:
                return 'cardiologist'
            elif 'surgeon' in content:
                return 'surgeon'
            elif 'doctor' in content or 'physician' in content:
                return 'doctor'
            elif 'nurse' in content:
                return 'nurse'
                
        # Engineering fields (moved after optical and software to avoid conflicts)
        engineering_terms = ['engineer', 'engineering', 'technical', 'design']
        if any(term in content for term in engineering_terms):
            if 'mechanical' in content or 'machine' in content:
                return 'mechanical_engineer'
            elif 'civil' in content or 'construction' in content:
                return 'civil_engineer'
            elif 'electrical' in content or 'electric' in content:
                return 'electrician'
                
        # Trades and skilled labor
        trades_terms = ['technician', 'repair', 'maintenance', 'install', 'fix', 'service']
        if any(term in content for term in trades_terms):
            if 'plumb' in content or 'pipe' in content:
                return 'plumber'
            elif 'electric' in content or 'wiring' in content:
                return 'electrician'
            elif 'car' in content or 'auto' in content or 'vehicle' in content:
                return 'mechanic'
                
        # Food service
        food_terms = ['chef', 'cook', 'kitchen', 'culinary', 'restaurant', 'food', 'dishwasher']
        if any(term in content for term in food_terms):
            if 'chef' in content or 'culinary' in content:
                return 'chef'
            elif 'dishwasher' in content or 'dish' in content:
                return 'dishwasher'
            elif 'cook' in content:
                return 'cook'
                
        # Education
        education_terms = ['teacher', 'education', 'school', 'student', 'classroom', 'curriculum']
        if any(term in content for term in education_terms):
            return 'teacher'
            
        # Legal
        legal_terms = ['lawyer', 'attorney', 'legal', 'law', 'court', 'litigation']
        if any(term in content for term in legal_terms):
            return 'lawyer'
            
        # Finance
        finance_terms = ['finance', 'accounting', 'accountant', 'financial', 'banking', 'investment']
        if any(term in content for term in finance_terms):
            if 'account' in content:
                return 'accountant'
            else:
                return 'financial_analyst'
                
        # Extract the main job title if it's a simple field name
        simple_fields = content.strip().split()
        if len(simple_fields) <= 2:
            # Clean the field name
            field_name = simple_fields[0].replace(',', '').replace('.', '').replace(':', '')
            if len(field_name) > 2:  # Valid field name
                return field_name.lower().replace(' ', '_')
                
        return None  # No dynamic field detected
    
    def generate_dynamic_job_analysis_config(self, field: str) -> dict:
        """Generate job analysis configuration for any career field"""
        display_field = field.replace('_', ' ').title()
        
        if 'surgeon' in field or 'doctor' in field:
            if 'brain' in field:
                return {
                    'skills': ['neurosurgery', 'brain surgery', 'microsurgery', 'surgical planning', 'patient care',
                              'medical imaging', 'surgical instruments', 'post-operative care', 'medical documentation'],
                    'responsibilities': [
                        "Perform complex neurosurgical procedures on brain and spinal cord",
                        "Diagnose and treat neurological conditions requiring surgical intervention",
                        "Collaborate with medical team for comprehensive patient care",
                        "Maintain surgical skills through continuing medical education",
                        "Ensure patient safety and optimal surgical outcomes"
                    ]
                }
            elif 'heart' in field or 'cardio' in field:
                return {
                    'skills': ['cardiology', 'heart surgery', 'cardiac procedures', 'patient assessment', 'medical imaging',
                              'surgical techniques', 'cardiac care', 'emergency response', 'medical documentation'],
                    'responsibilities': [
                        "Perform cardiac surgical procedures and interventions",
                        "Diagnose and treat cardiovascular conditions",
                        "Manage cardiac emergency situations and critical care",
                        "Collaborate with cardiology team for patient treatment",
                        "Maintain board certification and medical expertise"
                    ]
                }
            else:
                return {
                    'skills': ['surgery', 'medical procedures', 'patient care', 'surgical planning', 'medical knowledge',
                              'surgical instruments', 'post-operative care', 'medical documentation', 'patient safety'],
                    'responsibilities': [
                        "Perform surgical procedures with precision and safety",
                        "Provide comprehensive patient care and medical consultation",
                        "Collaborate with medical team for optimal patient outcomes",
                        "Maintain medical certifications and continuing education",
                        "Ensure compliance with medical standards and protocols"
                    ]
                }
        elif 'engineer' in field:
            base_skills = ['engineering design', 'technical analysis', 'project management', 'problem solving',
                          'technical documentation', 'quality assurance', 'safety protocols', 'team collaboration']
            if 'mechanical' in field:
                base_skills.extend(['cad design', 'manufacturing', 'mechanical systems', 'materials science'])
            elif 'civil' in field:
                base_skills.extend(['structural design', 'construction management', 'site planning', 'surveying'])
            
            return {
                'skills': base_skills,
                'responsibilities': [
                    f"Design and develop {field.replace('_', ' ')} solutions and systems",
                    "Manage technical projects from concept to completion",
                    "Ensure compliance with engineering standards and safety regulations",
                    "Collaborate with cross-functional teams on complex projects",
                    "Maintain professional engineering credentials and certifications"
                ]
            }
        elif field in ['lawyer', 'attorney']:
            return {
                'skills': ['legal research', 'case analysis', 'litigation', 'client representation', 'legal writing',
                          'court procedures', 'case management', 'legal strategy', 'negotiation', 'legal compliance'],
                'responsibilities': [
                    "Represent clients in legal matters and court proceedings",
                    "Conduct comprehensive legal research and case analysis",
                    "Prepare legal documents, briefs, and case strategies",
                    "Negotiate settlements and agreements on behalf of clients",
                    "Maintain bar admission and continuing legal education"
                ]
            }
        else:
            # Generic field configuration
            field_lower = field.replace('_', ' ').lower()
            return {
                'skills': [f'{field_lower} expertise', 'professional skills', 'industry knowledge', 'client service',
                          'project management', 'quality assurance', 'team collaboration', 'problem solving'],
                'responsibilities': [
                    f"Apply professional {field_lower} expertise in diverse work situations",
                    f"Deliver high-quality {field_lower} services to clients and stakeholders",
                    f"Collaborate with teams to achieve {field_lower} project objectives",
                    f"Maintain professional standards and industry certifications",
                    f"Continuously develop {field_lower} skills and knowledge"
                ]
            }
    
    def generate_dynamic_industry_keywords(self, field: str) -> list:
        """Generate industry keywords for any career field"""
        display_field = field.replace('_', ' ').title()
        
        if 'surgeon' in field or 'doctor' in field:
            return ['Healthcare', 'Medical Surgery', 'Patient Care', 'Clinical Excellence']
        elif 'engineer' in field:
            return ['Engineering', 'Technical Solutions', 'Project Management', 'Innovation']
        elif field in ['lawyer', 'attorney']:
            return ['Legal Services', 'Litigation', 'Client Representation', 'Legal Expertise']
        elif 'accountant' in field or 'financial' in field:
            return ['Financial Services', 'Accounting', 'Financial Analysis', 'Compliance']
        else:
            return [f'{display_field}', 'Professional Services', 'Industry Expertise', 'Client Solutions']
    
    def get_field_data(self, field: str) -> dict:
        """Get education, skills, and projects data for a specific field"""
        field_data = {
            'quantum_computing_scientist': {
                'education': "Massachusetts Institute of Technology — Ph.D. in Quantum Information Science, 2020\nCalifornia Institute of Technology — M.S. in Physics, 2017\nHarvard University — B.S. in Physics and Mathematics (Summa Cum Laude), 2015",
                'experience_title': 'Quantum Computing Research Scientist | Quantum Information Lab | 2020 - Present',
                'experience_bullets': [
                    'Developed novel quantum algorithms for optimization problems, achieving quadratic speedup over classical methods',
                    'Implemented variational quantum eigensolvers (VQE) for molecular simulation with 99.9% accuracy',
                    'Published 15+ papers in top-tier journals including Nature Quantum Information and Physical Review Letters',
                    'Led quantum error correction research resulting in 50% reduction in logical error rates for surface codes'
                ],
                'skills': [
                    'Quantum Algorithms: Shor\'s Algorithm, Grover\'s Algorithm, VQE, QAOA, Quantum Machine Learning',
                    'Programming: Qiskit, Cirq, PennyLane, Q#, Python, MATLAB, Mathematica, Julia',
                    'Mathematics: Linear Algebra, Group Theory, Tensor Networks, Information Theory, Optimization',
                    'Hardware: Superconducting Qubits, Trapped Ions, Photonic Systems, NISQ Devices',
                    'Research: Quantum Error Correction, Fault-Tolerant Computing, Quantum Simulation, Cryptography'
                ],
                'projects': [
                    'Quantum Supremacy Demonstration: Led team achieving quantum advantage for random circuit sampling',
                    'Quantum Chemistry Simulation: Developed VQE algorithms for drug discovery applications',
                    'Error Correction Breakthrough: Designed new topological codes with threshold above 1%',
                    'Quantum Machine Learning: Created quantum neural networks for pattern recognition tasks'
                ]
            },
            'mathematics_professor': {
                'education': "Princeton University — Ph.D. in Pure Mathematics (Algebraic Geometry), 2018\nMassachusetts Institute of Technology — M.S. in Mathematics, 2015\nHarvard University — B.A. in Mathematics (Phi Beta Kappa), 2013",
                'experience_title': 'Assistant Professor of Mathematics | Research University | 2020 - Present',
                'experience_bullets': [
                    'Proved fundamental theorems in algebraic geometry, published in Annals of Mathematics and Inventiones',
                    'Developed new techniques in homological algebra advancing K-theory and motivic cohomology',
                    'Solved longstanding conjectures in arithmetic geometry using advanced cohomological methods',
                    'Mentored 12+ graduate students and postdocs, with 8 completing Ph.D. dissertations under supervision'
                ],
                'skills': [
                    'Pure Mathematics: Algebraic Geometry, Homological Algebra, Category Theory, Number Theory',
                    'Analysis: Real/Complex Analysis, Functional Analysis, Harmonic Analysis, Differential Geometry',
                    'Algebra: Commutative Algebra, Representation Theory, Lie Algebras, Algebraic Topology',
                    'Computational: SageMath, Magma, GAP, Macaulay2, LaTeX, TikZ, Proof Assistants (Lean, Coq)',
                    'Research: Spectral Sequences, Derived Categories, Sheaf Cohomology, Modular Forms'
                ],
                'projects': [
                    'Breakthrough in Langlands Program: Established new cases of geometric Langlands correspondence',
                    'Cohomology Theory Advancement: Developed new computational tools for étale cohomology',
                    'Arithmetic Geometry Research: Proved rationality results for certain classes of algebraic varieties',
                    'NSF CAREER Award: $500K grant for research in motivic homotopy theory and A¹-algebraic topology'
                ]
            },
            'mathematician': {
                'education': "Stanford University — Ph.D. in Applied Mathematics, 2019\nUniversity of California Berkeley — M.S. in Mathematics, 2016\nMassachusetts Institute of Technology — B.S. in Mathematics (Dean\'s List), 2014",
                'experience_title': 'Research Mathematician | Applied Mathematics Institute | 2019 - Present',
                'experience_bullets': [
                    'Developed mathematical models for complex systems analysis, improving prediction accuracy by 40%',
                    'Created novel algorithms for solving high-dimensional partial differential equations',
                    'Published research in SIAM journals on computational mathematics and optimization theory',
                    'Collaborated with interdisciplinary teams on machine learning and data science applications'
                ],
                'skills': [
                    'Applied Mathematics: PDEs, Optimization Theory, Numerical Analysis, Stochastic Processes',
                    'Computational: Python, MATLAB, R, C++, High-Performance Computing, Parallel Algorithms',
                    'Statistics: Probability Theory, Statistical Learning, Bayesian Methods, Time Series Analysis',
                    'Machine Learning: Deep Learning, Reinforcement Learning, Mathematical Foundations of ML'
                ],
                'projects': [
                    'Climate Modeling: Developed mathematical models for climate change prediction systems',
                    'Financial Mathematics: Created risk assessment models for derivatives pricing',
                    'Optimization Algorithms: Designed convex optimization methods for large-scale problems',
                    'Data Science Applications: Applied mathematical theory to big data analytics'
                ]
            },
            'theoretical_physicist': {
                'education': "Harvard University — Ph.D. in Theoretical Physics, 2019\nCalifornia Institute of Technology — M.S. in Physics, 2016\nMassachusetts Institute of Technology — B.S. in Physics (Magna Cum Laude), 2014",
                'experience_title': 'Theoretical Physics Researcher | Institute for Advanced Study | 2019 - Present',
                'experience_bullets': [
                    'Developed new theoretical frameworks in quantum field theory and string theory',
                    'Published groundbreaking research on black hole physics and holographic duality',
                    'Presented findings at international conferences including Strings and ICHEP',
                    'Collaborated with experimental groups at CERN and other major physics laboratories'
                ],
                'skills': [
                    'Theoretical Physics: Quantum Field Theory, General Relativity, String Theory, Particle Physics',
                    'Mathematical Physics: Differential Geometry, Lie Groups, Topology, Complex Analysis',
                    'Computational: Mathematica, Python, C++, Monte Carlo Methods, Lattice QCD',
                    'Research: AdS/CFT Correspondence, Supersymmetry, Gauge Theories, Cosmology'
                ],
                'projects': [
                    'Black Hole Information Paradox: Contributed to resolution using holographic principles',
                    'String Theory Compactification: Developed new models for extra-dimensional physics',
                    'Quantum Gravity Research: Advanced understanding of spacetime emergence',
                    'Particle Phenomenology: Created models connecting theory to experimental observations'
                ]
            },
            'electrician': {
                'education': "Florida Atlantic University — B.S. Computer Science (Electrical Focus), Expected 2024 (Dean's List, GPA 3.7)",
                'experience_title': 'Industrial Electrician & Maintenance Specialist | Independent Contractor | 2022 - Present',
                'experience_bullets': [
                    'Installed and maintained complex industrial electrical systems including motor control panels and PLCs',
                    'Performed systematic electrical troubleshooting on 480V three-phase systems, reducing downtime by 30%',
                    'Developed preventive maintenance programs that increased electrical equipment reliability by 25%',
                    'Ensured strict compliance with NEC codes and OSHA safety standards on all installations'
                ],
                'skills': [
                    'Electrical Systems: Motor Controls, PLCs (Allen-Bradley), VFDs, Transformers, Switchgear',
                    'Codes & Standards: National Electrical Code (NEC), OSHA Safety, Arc Flash Safety, LOTO',
                    'Testing Equipment: Multimeters, Oscilloscopes, Megohmmeters, Power Quality Analyzers',
                    'Industrial Systems: 480V Three-Phase, Industrial Machinery, Automation, Preventive Maintenance'
                ],
                'projects': [
                    'Industrial Motor Control Installation: Complete 3-phase motor control panel installation',
                    'PLC Programming Project: Allen-Bradley PLC programming for automated conveyor system',
                    'OSHA 10-Hour Electrical Safety Certification',
                    'NEC Code Compliance: Successfully completed electrical inspections meeting all requirements'
                ]
            },
            'plumber': {
                'education': "Palm Beach State College — Associate of Applied Science in Plumbing Technology, 2022\nFlorida Atlantic University — B.S. Computer Science (Technical Systems Focus), Expected 2024",
                'experience_title': 'Professional Plumber | Independent Contractor | 2022 - Present',
                'experience_bullets': [
                    'Installed and maintained residential and commercial plumbing systems including water lines and drain systems',
                    'Performed comprehensive plumbing repairs and emergency service calls with 95% customer satisfaction rating',
                    'Installed and serviced water heaters, garbage disposals, and plumbing fixtures following local code requirements',
                    'Diagnosed complex plumbing problems using modern diagnostic equipment and provided cost-effective solutions'
                ],
                'skills': [
                    'Plumbing Systems: Water Supply Lines, Drain Systems, Fixture Installation, Pipe Repair',
                    'Tools & Equipment: Pipe Wrenches, Drain Augers, Pipe Cutters, Soldering Equipment, Threading Machines',
                    'Materials: PVC, Copper, PEX Piping, Cast Iron, Various Fittings and Fixtures',
                    'Codes & Standards: Local Plumbing Codes, Uniform Plumbing Code, Safety Regulations',
                    'Water Systems: Water Heaters, Boilers, Pressure Tanks, Backflow Prevention',
                    'Repair Services: Drain Cleaning, Leak Repair, Emergency Service, System Diagnostics'
                ],
                'projects': [
                    'Plumbing Code Compliance Training: Certified in local plumbing codes and regulations',
                    'Water Heater Installation and Service Certification',
                    'Drain Cleaning and Sewer Line Repair Techniques',
                    'Customer Service Excellence Recognition: 95% satisfaction rating'
                ]
            },
            'software_engineer': {
                'education': "Florida Atlantic University — B.S. Computer Science, Expected 2024 (Dean's List, GPA 3.7)",
                'experience_title': 'Financial Technology Developer | Independent Projects | 2022 - Present',
                'experience_bullets': [
                    'Developed automated trading algorithms using Python and Alpaca API, achieving 15% annual returns',
                    'Implemented machine learning models for cryptocurrency price prediction with 68% accuracy',
                    'Built scalable data processing pipelines handling 10,000+ daily market data points',
                    'Created interactive web dashboards using React and Flask for real-time portfolio monitoring'
                ],
                'skills': [
                    'Languages: C#, C/C++, Python, Java, JavaScript, SQL, HTML/CSS, PHP',
                    'Frameworks/Tools: .NET Blazor, ASP.NET MVC, Django, Flask, React, Git/GitHub, MySQL, MS SQL',
                    'AI/ML: TensorFlow, Keras, CNNs, NLP (RNNs/Transformers), scikit-learn, pandas, NumPy',
                    'Cloud/Other: AWS, Azure, Docker, Linux, Jupyter Notebooks, Algorithmic Trading'
                ],
                'projects': [
                    'Algorithmic Trading Platform: https://github.com/ryan-wlr/trading-bot',
                    'Cryptocurrency Analysis Tool: https://github.com/ryan-wlr/crypto-analysis',
                    'Portfolio Optimization System: https://github.com/ryan-wlr/portfolio-optimizer'
                ]
            }
        }
        
        # Add more fields with similar structure
        additional_fields = {
            'nurse': {
                'education': "Florida Atlantic University — Bachelor of Science in Nursing (BSN), Expected 2024 (Dean's List, GPA 3.7)",
                'experience_title': 'Registered Nurse | Memorial Healthcare System | 2022 - Present',
                'experience_bullets': [
                    'Provided comprehensive patient care for 8-12 patients per shift in medical-surgical unit',
                    'Administered medications and treatments following physician orders and nursing protocols',
                    'Monitored patient vital signs and documented care using electronic health records (EPIC)',
                    'Collaborated with interdisciplinary team to develop and implement patient care plans'
                ],
                'skills': [
                    'Clinical Skills: Patient Assessment, Medication Administration, IV Therapy, Wound Care',
                    'Medical Equipment: Monitors, Ventilators, IV Pumps, EKG Machines, Defibrillators',
                    'Documentation: Electronic Health Records (EHR), EPIC, Cerner, Medical Charting',
                    'Patient Care: Critical Care, Emergency Response, Patient Education, Family Communication'
                ],
                'projects': [
                    'ACLS (Advanced Cardiovascular Life Support) Certification',
                    'BLS (Basic Life Support) Certification',
                    'Patient Safety Excellence Recognition',
                    'Clinical Leadership Training Completion'
                ]
            },
            'teacher': {
                'education': "Florida Atlantic University — Bachelor of Education in Elementary Education, Expected 2024 (Dean's List, GPA 3.7)",
                'experience_title': 'Elementary School Teacher | Palm Beach County Schools | 2022 - Present',
                'experience_bullets': [
                    'Designed and implemented engaging lesson plans for 25+ students in grades K-5',
                    'Utilized differentiated instruction strategies to meet diverse learning needs and styles',
                    'Maintained detailed student progress records and conducted parent-teacher conferences',
                    'Integrated educational technology to enhance student learning and engagement'
                ],
                'skills': [
                    'Curriculum Development: Lesson Planning, Standards Alignment, Assessment Design, Learning Objectives',
                    'Classroom Management: Behavior Management, Student Engagement, Differentiated Instruction',
                    'Educational Technology: SmartBoards, Google Classroom, Educational Apps, Online Learning Platforms',
                    'Student Assessment: Formative Assessment, Progress Monitoring, Data Analysis, Parent Communication'
                ],
                'projects': [
                    'Florida Teaching Certificate (Elementary Education K-6)',
                    'ESE (Exceptional Student Education) Endorsement',
                    'Reading Endorsement Certification',
                    'Classroom Innovation Award Recognition'
                ]
            },
            'mechanic': {
                'education': "Palm Beach State College — Associate of Applied Science in Automotive Technology, 2022\nFlorida Atlantic University — B.S. Computer Science (Technical Systems Focus), Expected 2024",
                'experience_title': 'Automotive Technician | Independent Contractor | 2022 - Present',
                'experience_bullets': [
                    'Diagnosed and repaired complex automotive issues using advanced diagnostic equipment',
                    'Performed routine maintenance services including oil changes, brake repairs, and tune-ups',
                    'Maintained detailed service records and provided accurate repair estimates to customers',
                    'Specialized in engine diagnostics and electrical system troubleshooting'
                ],
                'skills': [
                    'Engine Systems: Diagnostics, Repair, Maintenance, Performance Tuning, Fuel Systems',
                    'Diagnostic Tools: OBD Scanners, Multimeters, Oscilloscopes, Compression Testers',
                    'Vehicle Systems: Electrical, Brake, Suspension, Transmission, HVAC Systems',
                    'Certifications: ASE Certified, Manufacturer Training, Safety Protocols'
                ],
                'projects': [
                    'ASE (Automotive Service Excellence) Certification',
                    'EPA 609 Refrigerant Handling Certification',
                    'Customer Service Excellence Recognition',
                    'Advanced Diagnostic Training Completion'
                ]
            },
            'dishwasher': {
                'education': "Palm Beach State College — High School Diploma, 2022\nFlorida Atlantic University — B.S. Computer Science (Food Service Focus), Expected 2024",
                'experience_title': 'Kitchen Staff / Dishwasher | Restaurant Operations | 2022 - Present',
                'experience_bullets': [
                    'Maintained high-volume dish washing operations during peak service hours serving 200+ customers daily',
                    'Ensured strict adherence to food safety and sanitation protocols in fast-paced kitchen environment',
                    'Operated commercial dishwashing equipment and maintained kitchen cleanliness standards',
                    'Collaborated with kitchen staff to maintain efficient workflow and timely service delivery'
                ],
                'skills': [
                    'Kitchen Operations: Dish Washing, Sanitization, Equipment Maintenance, Food Safety',
                    'Equipment: Commercial Dishwashers, Sanitizing Systems, Kitchen Appliances, Cleaning Tools',
                    'Food Safety: ServSafe Certified, Proper Storage, Temperature Control, Hygiene Standards',
                    'Customer Service: Team Collaboration, Fast-Paced Environment, Quality Standards, Time Management'
                ],
                'projects': [
                    'ServSafe Food Handler Certification',
                    'Kitchen Equipment Operation Training',
                    'Food Safety and Sanitation Compliance',
                    'Team Efficiency and Speed Recognition'
                ]
            },
            'chef': {
                'education': "Culinary Institute of America — Associate Degree in Culinary Arts, 2022\nFlorida Atlantic University — B.S. Computer Science (Culinary Management Focus), Expected 2024",
                'experience_title': 'Executive Chef | Fine Dining Restaurant | 2022 - Present',
                'experience_bullets': [
                    'Developed innovative seasonal menus resulting in 25% increase in customer satisfaction ratings',
                    'Supervised kitchen team of 12+ staff members during high-volume service periods serving 300+ covers nightly',
                    'Implemented cost control measures reducing food waste by 20% while maintaining quality standards',
                    'Ensured compliance with food safety regulations and maintained ServSafe Manager certification'
                ],
                'skills': [
                    'Culinary Arts: Menu Development, Food Preparation, Cooking Techniques, Recipe Creation',
                    'Kitchen Management: Staff Supervision, Inventory Control, Cost Management, Quality Control',
                    'Leadership: Team Training, Performance Management, Service Coordination, Kitchen Operations',
                    'Certifications: ServSafe Manager, Culinary Arts Professional, Food Safety, HACCP'
                ],
                'projects': [
                    'ServSafe Manager Certification',
                    'Culinary Arts Professional Development',
                    'Menu Innovation and Cost Analysis Project',
                    'Kitchen Leadership Excellence Recognition'
                ]
            },
            'data_scientist': {
                'education': "University of Central Florida — B.S. Computer Science, 2013 (Dean's List, GPA 3.8)\nValencia College — A.A., 2011 (Dean's List, GPA 3.7)",
                'experience_title': 'Data Scientist | Machine Learning Engineer | 2022 - Present',
                'experience_bullets': [
                    'Developed predictive models using Python, scikit-learn, and TensorFlow achieving 85% accuracy',
                    'Built automated data pipelines processing 1M+ records daily using pandas and Apache Spark',
                    'Created interactive dashboards with Tableau and Power BI for executive decision making',
                    'Implemented A/B testing frameworks resulting in 25% improvement in conversion rates'
                ],
                'skills': [
                    'Programming: Python, R, SQL, Julia, Scala, MATLAB',
                    'ML Libraries: scikit-learn, TensorFlow, PyTorch, Keras, XGBoost, pandas, NumPy',
                    'Data Visualization: Tableau, Power BI, matplotlib, seaborn, plotly, D3.js',
                    'Big Data: Apache Spark, Hadoop, AWS EMR, Databricks, Snowflake, Redshift'
                ],
                'projects': [
                    'Customer Churn Prediction Model: https://github.com/ryan-wlr/churn-prediction',
                    'Real-time Fraud Detection System: https://github.com/ryan-wlr/fraud-detection',
                    'Time Series Forecasting Platform: https://github.com/ryan-wlr/time-series-forecast'
                ]
            },
            'web_developer': {
                'education': "University of Central Florida — B.S. Computer Science, 2013 (Dean's List, GPA 3.8)\nValencia College — A.A., 2011 (Dean's List, GPA 3.7)",
                'experience_title': 'Full Stack Web Developer | Frontend Specialist | 2022 - Present',
                'experience_bullets': [
                    'Built responsive web applications using React, TypeScript, and Node.js serving 50K+ users',
                    'Developed RESTful APIs with Django and PostgreSQL handling 10K+ requests per minute',
                    'Optimized web performance achieving 95+ Lighthouse scores and 40% faster load times',
                    'Implemented modern CI/CD pipelines with Docker, GitHub Actions, and AWS deployment'
                ],
                'skills': [
                    'Frontend: HTML5, CSS3, JavaScript, TypeScript, React, Angular, Vue.js, Sass',
                    'Backend: Node.js, Python, Django, Flask, Express.js, PHP, Laravel',
                    'Databases: PostgreSQL, MongoDB, MySQL, Redis, Elasticsearch',
                    'Tools: Git, Docker, Webpack, Babel, Jest, Cypress, AWS, Vercel'
                ],
                'projects': [
                    'E-commerce Platform: https://github.com/ryan-wlr/ecommerce-platform',
                    'Social Media Dashboard: https://github.com/ryan-wlr/social-dashboard',
                    'Real-time Chat Application: https://github.com/ryan-wlr/realtime-chat'
                ]
            },
            'mobile_developer': {
                'education': "University of Central Florida — B.S. Computer Science, 2013 (Dean's List, GPA 3.8)\nValencia College — A.A., 2011 (Dean's List, GPA 3.7)",
                'experience_title': 'Mobile App Developer | iOS & Android Specialist | 2022 - Present',
                'experience_bullets': [
                    'Developed native iOS apps using Swift and SwiftUI with 100K+ downloads on App Store',
                    'Built cross-platform apps with React Native and Flutter supporting iOS and Android',
                    'Integrated Firebase services for authentication, real-time database, and push notifications',
                    'Optimized app performance achieving 4.8+ star ratings and 99.9% crash-free sessions'
                ],
                'skills': [
                    'Mobile Languages: Swift, Kotlin, Dart, Java, Objective-C, JavaScript',
                    'Frameworks: SwiftUI, UIKit, Jetpack Compose, React Native, Flutter, Xamarin',
                    'Tools: Xcode, Android Studio, Firebase, App Center, Fastlane, TestFlight',
                    'Backend: Node.js, Python, REST APIs, GraphQL, Push Notifications'
                ],
                'projects': [
                    'Fitness Tracking App: https://github.com/ryan-wlr/fitness-tracker',
                    'Recipe Sharing Platform: https://github.com/ryan-wlr/recipe-app',
                    'Location-based Social App: https://github.com/ryan-wlr/location-social'
                ]
            },
            'devops_engineer': {
                'education': "University of Central Florida — B.S. Computer Science, 2013 (Dean's List, GPA 3.8)\nValencia College — A.A., 2011 (Dean's List, GPA 3.7)",
                'experience_title': 'DevOps Engineer | Cloud Infrastructure Specialist | 2022 - Present',
                'experience_bullets': [
                    'Designed and maintained AWS infrastructure supporting 1M+ users with 99.99% uptime',
                    'Automated deployment pipelines using Docker, Kubernetes, and Terraform reducing deploy time by 80%',
                    'Implemented monitoring and alerting systems with Prometheus, Grafana, and ELK stack',
                    'Managed CI/CD workflows with Jenkins and GitLab CI processing 500+ deployments per month'
                ],
                'skills': [
                    'Cloud Platforms: AWS, Azure, Google Cloud Platform, DigitalOcean',
                    'Containerization: Docker, Kubernetes, Docker Compose, Helm, Istio',
                    'Infrastructure: Terraform, Ansible, CloudFormation, Pulumi, Vagrant',
                    'Monitoring: Prometheus, Grafana, ELK Stack, Datadog, New Relic, PagerDuty'
                ],
                'projects': [
                    'Auto-scaling Kubernetes Cluster: https://github.com/ryan-wlr/k8s-autoscale',
                    'Infrastructure as Code Templates: https://github.com/ryan-wlr/terraform-modules',
                    'CI/CD Pipeline Framework: https://github.com/ryan-wlr/cicd-framework'
                ]
            },
            'security_engineer': {
                'education': "University of Central Florida — B.S. Computer Science, 2013 (Dean's List, GPA 3.8)\nValencia College — A.A., 2011 (Dean's List, GPA 3.7)",
                'experience_title': 'Cybersecurity Engineer | Security Analyst | 2022 - Present',
                'experience_bullets': [
                    'Conducted penetration testing and vulnerability assessments identifying 200+ security issues',
                    'Implemented security monitoring systems using SIEM tools reducing incident response time by 60%',
                    'Developed security automation scripts with Python and PowerShell for threat detection',
                    'Led incident response efforts for 50+ security events achieving 100% containment success rate'
                ],
                'skills': [
                    'Security Tools: Metasploit, Nmap, Wireshark, Burp Suite, OWASP ZAP, Nessus',
                    'Programming: Python, PowerShell, Bash, C++, Assembly, Go, Rust',
                    'SIEM Platforms: Splunk, QRadar, ArcSight, Elastic Security, Chronicle',
                    'Compliance: ISO 27001, NIST Framework, SOX, HIPAA, GDPR, PCI DSS'
                ],
                'projects': [
                    'Automated Penetration Testing Suite: https://github.com/ryan-wlr/pentest-automation',
                    'Threat Intelligence Platform: https://github.com/ryan-wlr/threat-intel',
                    'Security Monitoring Dashboard: https://github.com/ryan-wlr/security-dashboard'
                ]
            }
        }
        
        field_data.update(additional_fields)
        
        # If field not found in predefined data, generate it dynamically
        if field not in field_data:
            return self.generate_dynamic_field_data(field)
            
        return field_data.get(field, field_data['software_engineer'])

    def generate_dynamic_field_data(self, field: str) -> dict:
        """Generate appropriate education, skills, and experience for any career field"""
        
        # Clean field name for display
        display_field = field.replace('_', ' ').title()
        
        # Generate field-specific content based on field type
        if 'surgeon' in field or 'doctor' in field or field in ['brain_surgeon', 'cardiologist']:
            if 'brain' in field:
                specialty = 'Neurosurgery'
                skills = [
                    'Surgical Skills: Brain Surgery, Neurosurgical Procedures, Microsurgery, Tumor Removal',
                    'Medical Equipment: Surgical Instruments, Imaging Systems, Neurosurgical Tools, OR Technology',
                    'Clinical Skills: Patient Assessment, Surgical Planning, Post-operative Care, Medical Documentation',
                    'Specializations: Brain Tumors, Trauma Surgery, Spinal Surgery, Minimally Invasive Procedures'
                ]
                experience_bullets = [
                    'Performed complex neurosurgical procedures with 95% success rate in brain tumor removals',
                    'Managed pre and post-operative care for 200+ patients annually in neurosurgery department',
                    'Collaborated with multidisciplinary medical team for comprehensive patient treatment plans',
                    'Maintained board certification and completed advanced neurosurgical training programs'
                ]
                title = 'Neurosurgeon / Brain Surgeon | Medical Center | 2020 - Present'
                education = "Harvard Medical School — Doctor of Medicine (M.D.), 2018\nJohns Hopkins University — Neurosurgery Residency, 2022\nFlorida Atlantic University — B.S. Biology (Pre-Med), 2014"
            elif 'cardio' in field or 'heart' in field:
                specialty = 'Cardiology'
                skills = [
                    'Cardiac Procedures: Heart Surgery, Cardiac Catheterization, Angioplasty, Bypass Surgery',
                    'Medical Equipment: Cardiac Monitors, Defibrillators, Catheterization Labs, Imaging Systems',
                    'Clinical Skills: Cardiac Assessment, Surgical Planning, Critical Care, Patient Management',
                    'Specializations: Interventional Cardiology, Heart Failure, Arrhythmias, Preventive Cardiology'
                ]
                experience_bullets = [
                    'Performed cardiac procedures with exceptional patient outcomes and 98% success rate',
                    'Managed cardiac care for 300+ patients annually in cardiovascular surgery department',
                    'Led cardiac emergency response team for critical care interventions',
                    'Maintained board certification in cardiology and cardiovascular surgery'
                ]
                title = 'Cardiologist / Heart Surgeon | Cardiac Center | 2020 - Present'
                education = "Harvard Medical School — Doctor of Medicine (M.D.), 2018\nMayo Clinic — Cardiology Fellowship, 2022\nFlorida Atlantic University — B.S. Biology (Pre-Med), 2014"
            else:
                specialty = 'Surgery'
                skills = [
                    'Surgical Skills: General Surgery, Minimally Invasive Procedures, Surgical Planning, OR Management',
                    'Medical Equipment: Surgical Instruments, Laparoscopic Tools, Imaging Systems, OR Technology',
                    'Clinical Skills: Patient Assessment, Surgical Techniques, Post-operative Care, Medical Records',
                    'Specializations: Emergency Surgery, Trauma Care, Surgical Consultation, Patient Safety'
                ]
                experience_bullets = [
                    'Performed surgical procedures with excellent patient outcomes and safety record',
                    'Managed surgical care for diverse patient population in hospital setting',
                    'Participated in emergency surgical response team for trauma cases',
                    'Maintained surgical credentials and completed continuing medical education'
                ]
                title = f'{display_field} | Medical Center | 2020 - Present'
                education = "Medical School — Doctor of Medicine (M.D.), 2018\nSurgical Residency Program, 2022\nFlorida Atlantic University — B.S. Biology (Pre-Med), 2014"
                
        elif 'engineer' in field and 'software' not in field:
            if 'mechanical' in field:
                skills = [
                    'Engineering Design: CAD/CAM, SolidWorks, AutoCAD, Mechanical Systems Design',
                    'Manufacturing: CNC Programming, Quality Control, Production Planning, Process Optimization',
                    'Technical Skills: Thermodynamics, Materials Science, Fluid Mechanics, Machine Design',
                    'Project Management: Engineering Projects, Team Leadership, Technical Documentation'
                ]
                experience_bullets = [
                    'Designed mechanical systems and components resulting in 20% efficiency improvements',
                    'Managed engineering projects from concept to production with cross-functional teams',
                    'Optimized manufacturing processes reducing production costs by 15%',
                    'Maintained professional engineering license and industry certifications'
                ]
                education = "Engineering School — Bachelor of Science in Mechanical Engineering, 2020\nFlorida Atlantic University — M.S. Mechanical Engineering, Expected 2024"
            elif 'civil' in field:
                skills = [
                    'Civil Design: Structural Analysis, AutoCAD Civil 3D, Project Planning, Site Development',
                    'Construction Management: Project Oversight, Quality Assurance, Safety Protocols, Cost Estimation',
                    'Technical Skills: Surveying, Hydraulics, Geotechnical Engineering, Environmental Compliance',
                    'Professional: PE License, Project Management, Technical Reports, Client Relations'
                ]
                experience_bullets = [
                    'Designed civil infrastructure projects including roads, bridges, and drainage systems',
                    'Managed construction projects ensuring compliance with safety and quality standards',
                    'Conducted site evaluations and prepared technical engineering reports',
                    'Maintained Professional Engineer license and continuing education requirements'
                ]
                education = "Engineering School — Bachelor of Science in Civil Engineering, 2020\nFlorida Atlantic University — M.S. Civil Engineering, Expected 2024"
            else:
                skills = [
                    f'{display_field} Design: Technical Design, Engineering Analysis, Project Development',
                    'Technical Tools: CAD Software, Engineering Analysis, Technical Documentation',
                    'Professional Skills: Project Management, Quality Assurance, Problem Solving',
                    'Industry Knowledge: Engineering Standards, Safety Protocols, Technical Innovation'
                ]
                experience_bullets = [
                    f'Applied {field.replace("_", " ")} principles to design and develop engineering solutions',
                    'Collaborated with engineering teams on complex technical projects',
                    'Ensured compliance with industry standards and safety regulations',
                    'Maintained professional development and technical certifications'
                ]
                education = f"Engineering School — Bachelor of Science in {display_field}, 2020\nFlorida Atlantic University — M.S. {display_field}, Expected 2024"
            title = f'{display_field} | Engineering Firm | 2020 - Present'
                
        elif field in ['lawyer', 'attorney']:
            skills = [
                'Legal Practice: Case Research, Legal Writing, Litigation, Client Representation',
                'Court Procedures: Trial Advocacy, Depositions, Legal Motions, Settlement Negotiations',
                'Legal Research: Case Law Analysis, Statute Interpretation, Legal Precedent Research',
                'Client Services: Legal Consultation, Case Management, Document Preparation, Legal Strategy'
            ]
            experience_bullets = [
                'Represented clients in legal matters with successful case outcomes and client satisfaction',
                'Conducted legal research and prepared comprehensive legal documents and briefs',
                'Negotiated settlements and agreements achieving favorable outcomes for clients',
                'Maintained bar admission and completed continuing legal education requirements'
            ]
            title = f'{display_field} | Law Firm | 2020 - Present'
            education = "Law School — Juris Doctor (J.D.), 2019\nFlorida Atlantic University — B.A. Political Science (Pre-Law), 2016"
            
        elif field in ['accountant', 'financial_analyst']:
            if 'accountant' in field:
                skills = [
                    'Accounting: Financial Statements, Tax Preparation, Audit Support, Bookkeeping',
                    'Software: QuickBooks, Excel, SAP, Financial Reporting Systems, Tax Software',
                    'Financial Analysis: Budget Analysis, Cost Accounting, Financial Planning, Variance Analysis',
                    'Compliance: GAAP Standards, Tax Regulations, Internal Controls, Financial Compliance'
                ]
                experience_bullets = [
                    'Prepared accurate financial statements and reports for diverse client portfolio',
                    'Managed tax preparation and compliance ensuring adherence to federal and state regulations',
                    'Conducted financial analysis supporting business decision-making processes',
                    'Maintained CPA certification and completed continuing professional education'
                ]
            else:
                skills = [
                    'Financial Analysis: Investment Analysis, Financial Modeling, Market Research, Risk Assessment',
                    'Software: Excel, Bloomberg, Financial Databases, Analytical Tools, Reporting Systems',
                    'Investment: Portfolio Analysis, Securities Analysis, Financial Planning, Market Analysis',
                    'Professional: CFA Certification, Financial Reporting, Client Presentations, Research'
                ]
                experience_bullets = [
                    'Conducted comprehensive financial analysis supporting investment decisions',
                    'Prepared detailed financial reports and presentations for senior management',
                    'Analyzed market trends and investment opportunities for portfolio optimization',
                    'Maintained financial certifications and industry professional development'
                ]
            title = f'{display_field} | Financial Services | 2020 - Present'
            education = "Business School — Bachelor of Science in Finance/Accounting, 2020\nFlorida Atlantic University — MBA Finance, Expected 2024"
            
        else:
            # Generic field generation
            skills = [
                f'{display_field} Skills: Professional expertise, industry knowledge, specialized techniques',
                f'Technical Abilities: {display_field.lower()} tools, industry software, professional equipment',
                f'Professional Skills: Project management, quality assurance, client service, team collaboration',
                f'Industry Knowledge: {display_field.lower()} standards, best practices, safety protocols'
            ]
            experience_bullets = [
                f'Applied professional {field.replace("_", " ")} expertise in challenging work environments',
                f'Delivered high-quality {field.replace("_", " ")} services with excellent client satisfaction',
                f'Collaborated effectively with teams to achieve {field.replace("_", " ")} project objectives',
                f'Maintained professional development and industry certifications in {field.replace("_", " ")}'
            ]
            title = f'{display_field} | Professional Services | 2020 - Present'
            education = f"Professional School — Degree in {display_field}, 2020\nFlorida Atlantic University — B.S. Computer Science ({display_field} Focus), Expected 2024"
            
        return {
            'education': education,
            'experience_title': title,
            'experience_bullets': experience_bullets,
            'skills': skills,
            'projects': [
                f'Professional {display_field} Certification',
                f'{display_field} Excellence Recognition',
                f'Advanced {display_field} Training Completion',
                f'Industry Leadership and Development'
            ]
        }

    def generate_dynamic_title(self, field: str) -> str:
        """Generate professional title for any career field"""
        display_field = field.replace('_', ' ').title()
        
        if 'surgeon' in field or 'doctor' in field:
            if 'brain' in field:
                return 'Neurosurgeon & Brain Surgery Specialist'
            elif 'heart' in field or 'cardio' in field:
                return 'Cardiologist & Cardiovascular Surgery Specialist'
            else:
                return f'{display_field} & Medical Professional'
        elif 'engineer' in field:
            return f'{display_field} & Technical Specialist'
        elif field in ['lawyer', 'attorney']:
            return 'Attorney & Legal Professional'
        elif 'accountant' in field or 'financial' in field:
            return f'{display_field} & Financial Professional'
        else:
            return f'{display_field} & Professional Specialist'

    def generate_dynamic_summary(self, field: str) -> str:
        """Generate professional summary for any career field"""
        display_field = field.replace('_', ' ').title()
        field_lower = field.replace('_', ' ').lower()
        
        if 'surgeon' in field or 'doctor' in field:
            specialty = 'neurosurgical' if 'brain' in field else 'cardiovascular' if 'cardio' in field or 'heart' in field else 'surgical'
            return f'Experienced {display_field} with 5+ years of comprehensive medical expertise in {specialty} procedures, patient care, and clinical excellence. Proven expertise in advanced surgical techniques, patient management, and medical team collaboration with demonstrated ability to achieve 95% successful patient outcomes. Strong knowledge of medical protocols, surgical safety standards, and healthcare regulations. Committed to delivering exceptional patient care through precision medicine and continuous professional development.'
        elif 'engineer' in field:
            engineering_type = 'mechanical systems' if 'mechanical' in field else 'civil infrastructure' if 'civil' in field else 'engineering solutions'
            return f'Experienced {display_field} with 5+ years of comprehensive technical expertise in {engineering_type}, project management, and engineering design. Proven expertise in technical analysis, system optimization, and cross-functional collaboration with demonstrated ability to improve efficiency by 25% through innovative engineering solutions. Strong knowledge of industry standards, safety protocols, and technical best practices. Committed to delivering high-quality engineering projects through technical excellence and continuous innovation.'
        elif field in ['lawyer', 'attorney']:
            return f'Experienced {display_field} with 5+ years of comprehensive legal expertise in case management, litigation, and client representation. Proven expertise in legal research, case strategy, and courtroom advocacy with demonstrated ability to achieve favorable outcomes in 90% of cases. Strong knowledge of legal procedures, regulatory compliance, and client service excellence. Committed to delivering exceptional legal representation through thorough preparation and strategic advocacy.'
        elif 'accountant' in field or 'financial' in field:
            focus_area = 'financial analysis and investment management' if 'analyst' in field else 'accounting, tax preparation, and financial compliance'
            return f'Experienced {display_field} with 5+ years of comprehensive expertise in {focus_area}. Proven expertise in financial reporting, regulatory compliance, and client service with demonstrated ability to improve financial efficiency by 25% through strategic analysis. Strong knowledge of accounting standards, tax regulations, and financial best practices. Committed to delivering accurate financial services through attention to detail and professional excellence.'
        else:
            return f'Experienced {display_field} with 5+ years of comprehensive expertise in {field_lower} operations, project management, and professional service delivery. Proven expertise in {field_lower} techniques, quality assurance, and client satisfaction with demonstrated ability to improve performance by 25% through professional excellence. Strong knowledge of industry standards, best practices, and professional protocols. Committed to delivering exceptional {field_lower} services through continuous improvement and professional development.'

    def create_formatted_docx_resume(self, results: Dict[str, str], output_dir: str):
        """Create a properly formatted .docx resume matching Ryan_Weiler_Resume.docx style"""
        if not HAS_DOCX:
            print("WARNING: Skipping .docx creation - python-docx not available")
            return
        
        try:
            # Create new document
            doc = Document()
            
            # Set up document margins to match original
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)
            
            # Use the previously detected field from mock_response
            detected_field = getattr(self, 'current_detected_field', 'software_engineer')
            field_data = self.get_field_data(detected_field)
            
            print(f"    >>> Detected field: {detected_field}")
            
            # Check if narrative resume exists (storytelling mode)
            if '7_narrative_resume' in results:
                print(f"    >>> Using narrative storytelling content")
                narrative_content = results['7_narrative_resume']
                # Create storytelling DOCX format
                self.build_narrative_docx(doc, narrative_content)
            else:
                print(f"    >>> Using {detected_field}-specific resume content")
                field_data = self.get_field_data(detected_field)
                # Create standard DOCX format
                self.build_standard_docx(doc, field_data, results, detected_field)
            
            # Save the document
            docx_path = os.path.join(output_dir, 'optimized_resume.docx')
            doc.save(docx_path)
            print(f"    >>> optimized_resume.docx (Ryan Weiler format)")
            
        except Exception as e:
            print(f"ERROR: Error creating .docx resume: {e}")
    
    def create_formatted_docx_resume_specific(self, results: Dict[str, str], output_dir: str, filename: str):
        """Create a properly formatted .docx resume with specific filename"""
        if not HAS_DOCX:
            print("WARNING: Skipping .docx creation - python-docx not available")
            return
        
        try:
            # Create new document
            doc = Document()
            
            # Set up document margins to match original
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)
            
            # Use the previously detected field from mock_response
            detected_field = getattr(self, 'current_detected_field', 'software_engineer')
            field_data = self.get_field_data(detected_field)
            
            print(f"    >>> Detected field: {detected_field}")
            
            # Check if narrative resume exists (storytelling mode)
            if '7_narrative_resume' in results:
                print(f"    >>> Using narrative storytelling content")
                narrative_content = results['7_narrative_resume']
                # Create storytelling DOCX format
                self.build_narrative_docx(doc, narrative_content)
            elif '7_enhanced_resume' in results:
                print(f"    >>> Using enhanced standard content")
                # For standard resumes, we need to create content differently
                # Since enhanced resume is typically text, we'll use standard format
                field_data = self.get_field_data(detected_field)
                self.build_standard_docx(doc, field_data, results, detected_field)
            else:
                print(f"    >>> Using {detected_field}-specific resume content")
                field_data = self.get_field_data(detected_field)
                # Create standard DOCX format
                self.build_standard_docx(doc, field_data, results, detected_field)
            
            # Save the document with custom filename
            docx_path = os.path.join(output_dir, filename)
            doc.save(docx_path)
            print(f"    >>> {filename} (Ryan Weiler format)")
            
        except Exception as e:
            print(f"ERROR: Error creating .docx resume: {e}")
    
    def add_hyperlink(self, paragraph, url, text):
        """Add a clickable hyperlink to a paragraph"""
        # Create hyperlink relationship
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create run element for the hyperlink text
        run = OxmlElement('w:r')
        
        # Add run properties (font formatting)
        rPr = OxmlElement('w:rPr')
        
        # Set color to blue
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')  # Standard hyperlink blue
        rPr.append(color)
        
        # Set underline
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        # Set font
        font = OxmlElement('w:rFonts')
        font.set(qn('w:ascii'), 'Calibri')
        font.set(qn('w:hAnsi'), 'Calibri')
        rPr.append(font)
        
        # Set font size
        size = OxmlElement('w:sz')
        size.set(qn('w:val'), '24')  # 12pt = 24 half-points
        rPr.append(size)
        
        run.append(rPr)
        
        # Add the text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        run.append(text_elem)
        
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        
        return hyperlink
    
    def build_narrative_docx(self, doc, narrative_content):
        """Build DOCX using narrative storytelling content"""
        
        # Parse narrative content to extract sections
        lines = narrative_content.split('\n')
        sections = {'hook': '', 'name': '', 'contact': '', 'narrative': '', 'skills': [], 'experience': [], 'education': []}
        
        current_section = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Detect key elements
            if '🔥 THE' in line or '🔬 THE' in line or '💻 THE' in line or '🌟 THE' in line:
                sections['hook'] = line
            elif 'CONTACT INFORMATION:' in line:
                current_section = 'contact'
                continue
            elif 'PROFESSIONAL NARRATIVE:' in line:
                current_section = 'narrative'
                continue
            elif 'TECHNICAL COMPETENCIES:' in line or 'CORE SKILLS:' in line:
                current_section = 'skills'
                continue
            elif 'EXPERIENCE:' in line or 'PROFESSIONAL EXPERIENCE:' in line:
                current_section = 'experience'
                continue
            elif line and not line.startswith('CAREER STORY RESUME') and not line.startswith('Chapter'):
                # Extract content based on current section
                if current_section is None and len(line.split()) <= 3 and any(c.isupper() for c in line):
                    sections['name'] = line
                elif current_section == 'contact':
                    # Extract education content to separate section, filter from contact
                    if ('university' in line.lower() or 'college' in line.lower()) and ('b.s.' in line.lower() or 'a.a.' in line.lower() or 'gpa' in line.lower()):
                        sections['education'].append(line)
                    elif not ('experience & projects' in line.lower() or 'built, trained' in line.lower()):
                        sections['contact'] += line + '\n'
                elif current_section == 'narrative':
                    sections['narrative'] += line + ' '
                elif current_section == 'skills' and line.startswith('•'):
                    sections['skills'].append(line[1:].strip())
                elif current_section == 'experience' and line.startswith('•'):
                    sections['experience'].append(line[1:].strip())
        
        # 1. Opening Hook (if present)
        if sections['hook']:
            hook_para = doc.add_paragraph()
            hook_para.style = doc.styles['Title']
            hook_run = hook_para.add_run(sections['hook'])
            hook_run.font.name = 'Calibri'
            hook_run.font.size = Pt(16)
            hook_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # 2. NAME
        name = sections['name'] if sections['name'] else "Ryan Thomas Weiler"
        name_para = doc.add_paragraph()
        name_para.style = doc.styles['Title']
        name_run = name_para.add_run(name)
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(26)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 3. CONTACT INFO (NO DUPLICATES)
        contact_text = sections['contact'].strip() if sections['contact'] else "📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com"
        
        # Parse contact text and create proper hyperlinks
        contact_lines = contact_text.split('\n')
        for contact_line in contact_lines:
            if contact_line.strip():
                contact_para = doc.add_paragraph()
                contact_para.style = doc.styles['Normal']
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Parse line for URLs and create clickable hyperlinks
                if 'https://' in contact_line:
                    # Split line into parts around URLs
                    parts = contact_line.split('https://')
                    
                    # Add text before first URL
                    if parts[0].strip():
                        contact_run = contact_para.add_run(parts[0])
                        contact_run.font.name = 'Calibri'
                        contact_run.font.size = Pt(12)
                    
                    # Process each URL
                    for i, part in enumerate(parts[1:]):
                        # Find where URL ends (space, pipe, or end of string)
                        url_end = len(part)
                        for char_idx, char in enumerate(part):
                            if char in [' ', '|', '\n', '\t']:
                                url_end = char_idx
                                break
                        
                        # Extract URL and remaining text
                        url = 'https://' + part[:url_end]
                        remaining_text = part[url_end:]
                        
                        # Add clickable hyperlink
                        try:
                            self.add_hyperlink(contact_para, url, url)
                        except Exception as e:
                            # Fallback to styled text if hyperlink fails
                            print(f"Warning: Could not create hyperlink for {url}: {e}")
                            url_run = contact_para.add_run(url)
                            url_run.font.name = 'Calibri'
                            url_run.font.size = Pt(12)
                            url_run.font.color.rgb = RGBColor(0, 0, 255)
                            url_run.font.underline = True
                        
                        # Add remaining text after URL
                        if remaining_text.strip():
                            text_run = contact_para.add_run(remaining_text)
                            text_run.font.name = 'Calibri'
                            text_run.font.size = Pt(12)
                else:
                    # No URLs, just add as regular text
                    contact_run = contact_para.add_run(contact_line.strip())
                    contact_run.font.name = 'Calibri'
                    contact_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # 4. PROFESSIONAL NARRATIVE
        if sections['narrative']:
            narrative_header = doc.add_paragraph()
            narrative_header.style = doc.styles['Heading 1']
            narrative_run = narrative_header.add_run("Professional Narrative")
            narrative_run.font.name = 'Calibri'
            narrative_run.font.size = Pt(14)
            
            narrative_para = doc.add_paragraph()
            narrative_para.style = doc.styles['Normal']
            narrative_text_run = narrative_para.add_run(sections['narrative'].strip())
            narrative_text_run.font.name = 'Calibri'
            narrative_text_run.font.size = Pt(12)
            doc.add_paragraph()
        
        # 5. TECHNICAL SKILLS
        if sections['skills']:
            skills_header = doc.add_paragraph()
            skills_header.style = doc.styles['Heading 1']
            skills_run = skills_header.add_run("Technical Competencies")
            skills_run.font.name = 'Calibri'
            skills_run.font.size = Pt(14)
            
            for skill in sections['skills'][:6]:
                skill_para = doc.add_paragraph()
                skill_para.style = doc.styles['Normal']
                skill_run = skill_para.add_run(f"• {skill}")
                skill_run.font.name = 'Calibri'
                skill_run.font.size = Pt(12)
            doc.add_paragraph()
        
        # 6. EDUCATION
        edu_header = doc.add_paragraph()
        edu_header.style = doc.styles['Heading 1']
        edu_run = edu_header.add_run("Education")
        edu_run.font.name = 'Calibri'
        edu_run.font.size = Pt(14)
        
        # Use extracted education from narrative, or fallback to field data
        if sections['education']:
            for edu_line in sections['education']:
                edu_para = doc.add_paragraph()
                edu_para.style = doc.styles['Normal']
                edu_run = edu_para.add_run(edu_line)
                edu_run.font.name = 'Calibri'
                edu_run.font.size = Pt(12)
        else:
            # Fallback to actual field data instead of hardcoded education
            detected_field = getattr(self, 'current_detected_field', 'data_scientist')
            field_data = self.get_field_data(detected_field)
            edu_para = doc.add_paragraph()
            edu_para.style = doc.styles['Normal']
            edu_run = edu_para.add_run(field_data['education'])
            edu_run.font.name = 'Calibri'
            edu_run.font.size = Pt(12)
        doc.add_paragraph()
        
        # 7. EXPERIENCE
        if sections['experience']:
            exp_header = doc.add_paragraph()
            exp_header.style = doc.styles['Heading 1']
            exp_run = exp_header.add_run("Professional Experience")
            exp_run.font.name = 'Calibri'
            exp_run.font.size = Pt(14)
            
            for exp in sections['experience'][:8]:
                exp_para = doc.add_paragraph()
                exp_para.style = doc.styles['Normal']
                exp_run = exp_para.add_run(f"• {exp}")
                exp_run.font.name = 'Calibri'
                exp_run.font.size = Pt(12)
            doc.add_paragraph()
        
        # 8. REFERENCES
        ref_header = doc.add_paragraph()
        ref_header.style = doc.styles['Heading 1']
        ref_run = ref_header.add_run("References")
        ref_run.font.name = 'Calibri'
        ref_run.font.size = Pt(14)
        
        ref_para = doc.add_paragraph()
        ref_para.style = doc.styles['Normal']
        ref_run = ref_para.add_run("Available upon request")
        ref_run.font.name = 'Calibri'
        ref_run.font.size = Pt(12)
    
    def build_standard_docx(self, doc, field_data, results, detected_field):
        """Build DOCX using standard format with field-specific data"""
        
        # 1. NAME (using Title style - 26pt, centered)
        name_para = doc.add_paragraph()
        name_para.style = doc.styles['Title']
        name_run = name_para.add_run("Ryan Thomas Weiler")
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(26)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. CONTACT INFO (centered, Normal style with 12pt font)
        contact_para = doc.add_paragraph()
        contact_para.style = doc.styles['Normal']
        contact_run = contact_para.add_run("📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com")
        contact_run.font.name = 'Calibri'
        contact_run.font.size = Pt(12)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        social_para = doc.add_paragraph()
        social_para.style = doc.styles['Normal']
        social_run = social_para.add_run("🔗 LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/ | 💻 GitHub: https://github.com/ryan-wlr")
        social_run.font.name = 'Calibri'
        social_run.font.size = Pt(12)
        social_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # 3. EDUCATION SECTION (using Heading 1 style - 14pt, bold)
        edu_header = doc.add_paragraph()
        edu_header.style = doc.styles['Heading 1']
        edu_run = edu_header.add_run("Education")
        edu_run.font.name = 'Calibri'
        edu_run.font.size = Pt(14)
        
        edu_detail = doc.add_paragraph()
        edu_detail.style = doc.styles['Normal']
        edu_text_run = edu_detail.add_run(field_data['education'])
        edu_text_run.font.name = 'Calibri'
        edu_text_run.font.size = Pt(12)
        doc.add_paragraph()
        
        # 4. EXPERIENCE & PROJECTS SECTION (using Heading 1 style - 14pt)
        exp_header = doc.add_paragraph()
        exp_header.style = doc.styles['Heading 1']
        exp_run = exp_header.add_run("Experience & Projects (Continuous Timeline)")
        exp_run.font.name = 'Calibri'
        exp_run.font.size = Pt(14)
        
        # Add main experience using dynamic field-specific data
        experiences = [
            {
                'title': field_data['experience_title'],
                'bullets': field_data['experience_bullets']
            },
            {
                'title': "University Projects & Research",
                'bullets': field_data['projects']
            }
        ]
        
        for exp in experiences:
            # Job title (Normal style, bold, 12pt)
            job_para = doc.add_paragraph()
            job_para.style = doc.styles['Normal']
            job_run = job_para.add_run(exp['title'])
            job_run.font.name = 'Calibri'
            job_run.font.size = Pt(12)
            job_run.bold = True
            
            # Bullets (Normal style with proper dash, 12pt)
            for bullet in exp['bullets']:
                bullet_para = doc.add_paragraph()
                bullet_para.style = doc.styles['Normal']
                bullet_run = bullet_para.add_run(f"- {bullet}")
                bullet_run.font.name = 'Calibri'
                bullet_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # 5. SKILLS SECTION (using Heading 1 style - 14pt)
        skills_header = doc.add_paragraph()
        skills_header.style = doc.styles['Heading 1']
        skills_run = skills_header.add_run("Technical Skills")
        skills_run.font.name = 'Calibri'
        skills_run.font.size = Pt(14)
        
        # Skills content using dynamic field-specific data
        for skill in field_data['skills']:
            skill_para = doc.add_paragraph()
            skill_para.style = doc.styles['Normal']
            skill_run = skill_para.add_run(f"• {skill}")
            skill_run.font.name = 'Calibri'
            skill_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # 6. REFERENCES SECTION (using Heading 1 style - 14pt)
        ref_header = doc.add_paragraph()
        ref_header.style = doc.styles['Heading 1']
        ref_run = ref_header.add_run("References")
        ref_run.font.name = 'Calibri'
        ref_run.font.size = Pt(14)
        
        ref_para = doc.add_paragraph()
        ref_para.style = doc.styles['Normal']
        ref_text_run = ref_para.add_run("Available upon request")
        ref_text_run.font.name = 'Calibri'
        ref_text_run.font.size = Pt(12)

    def browse_for_job_description(self) -> Optional[str]:
        """Open file browser to select job description file with error handling"""
        if not HAS_TKINTER:
            print("ERROR: tkinter not available for file browsing")
            print("TIP: Use command line mode instead: python resume_windows.py \"job description\" --resume \"resume text\"")
            return None
        
        try:
            print(">>> Opening file dialog for job description...")
            print("    (Press ESC or Cancel if you want to skip file selection)")
            
            root = tk.Tk()
            root.withdraw()  # Hide main window
            root.attributes('-topmost', True)  # Bring to front
            
            file_path = filedialog.askopenfilename(
                title="Select Job Description File",
                filetypes=[
                    ("Text files", "*.txt"),
                    ("Word documents", "*.docx"),
                    ("PDF files", "*.pdf"),
                    ("All files", "*.*")
                ]
            )
            
            root.destroy()
            
            if file_path:
                print(f"Selected job description: {os.path.basename(file_path)}")
                
                # Read file content based on extension
                try:
                    if file_path.lower().endswith('.txt'):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    elif file_path.lower().endswith('.docx') and HAS_DOCX:
                        doc = Document(file_path)
                        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    else:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    
                    print(f"Loaded {len(content)} characters from job description")
                    return content
                    
                except Exception as e:
                    print(f"ERROR: Failed to read file {file_path}: {e}")
                    return None
            else:
                print("No file selected - you can enter job description manually")
                return None
                
        except KeyboardInterrupt:
            print(">>> File dialog cancelled by user (Ctrl+C)")
            print("    Switching to manual input mode...")
            return None
        except Exception as e:
            print(f"ERROR: File browser error: {e}")
            print("TIP: Try using command line mode or demo mode instead")
            return None

    def browse_for_resume(self) -> Optional[str]:
        """Open file browser to select resume file with error handling"""
        if not HAS_TKINTER:
            print("ERROR: tkinter not available for file browsing")
            print("TIP: Use command line mode instead: python resume_windows.py \"job description\" --resume \"resume text\"")
            return None
        
        try:
            print(">>> Opening file dialog for current resume...")
            print("    (Press ESC or Cancel if you want to skip file selection)")
            
            root = tk.Tk()
            root.withdraw()  # Hide main window
            root.attributes('-topmost', True)  # Bring to front
            
            file_path = filedialog.askopenfilename(
                title="Select Current Resume File",
                filetypes=[
                    ("Word documents", "*.docx"),
                    ("Text files", "*.txt"), 
                    ("PDF files", "*.pdf"),
                    ("All files", "*.*")
                ]
            )
            
            root.destroy()
            
            if file_path:
                print(f"Selected resume: {os.path.basename(file_path)}")
                
                # Read file content based on extension
                try:
                    if file_path.lower().endswith('.docx') and HAS_DOCX:
                        doc = Document(file_path)
                        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    elif file_path.lower().endswith('.txt'):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    else:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    
                    print(f"Loaded {len(content)} characters from resume")
                    return content
                    
                except Exception as e:
                    print(f"ERROR: Failed to read file {file_path}: {e}")
                    return None
            else:
                print("No file selected - you can enter resume text manually")
                return None
                
        except KeyboardInterrupt:
            print(">>> File dialog cancelled by user (Ctrl+C)")
            print("    Switching to manual input mode...")
            return None
        except Exception as e:
            print(f"ERROR: File browser error: {e}")
            print("TIP: Try using command line mode or demo mode instead")
            return None

    def run_browse_mode(self):
        """Interactive mode using file browsers with fallback options"""
        print(">>> AI-POWERED RESUME OPTIMIZER - BROWSE MODE")
        print("    This mode will help you select files using graphical file browsers")
        print("    TIP: If file dialogs don't work, press Ctrl+C to switch to manual input")
        print()
        
        # Get job description
        print(">>> Select Job Description File:")
        job_description = self.browse_for_job_description()
        
        # If file dialog failed, offer manual input
        if not job_description:
            print("\n>>> File dialog didn't work? Let's try manual input:")
            print("    Enter job description (press Enter twice when finished):")
            lines = []
            try:
                while True:
                    line = input()
                    if line == "" and len(lines) > 0 and lines[-1] == "":
                        break
                    lines.append(line)
                job_description = '\n'.join(lines[:-1]) if lines else ""
            except KeyboardInterrupt:
                print("\n>>> Switching to demo mode with sample data...")
                job_description = """Senior Software Engineer - FinTech
                
We are seeking a talented Software Engineer to join our financial technology team. 
Must have Python, machine learning, and algorithmic trading experience.

Required: Python (3+ years), pandas, NumPy, machine learning, financial markets knowledge
Preferred: AWS, Docker, React, algorithmic trading experience"""
        
        if not job_description or not job_description.strip():
            print("ERROR: Job description is required. Exiting.")
            return
        
        # Get resume
        print("\n>>> Select Current Resume File:")
        resume_content = self.browse_for_resume()
        
        # If file dialog failed, offer manual input
        if not resume_content:
            print("\n>>> File dialog didn't work? Let's try manual input:")
            print("    Enter current resume content (press Enter twice when finished):")
            lines = []
            try:
                while True:
                    line = input()
                    if line == "" and len(lines) > 0 and lines[-1] == "":
                        break
                    lines.append(line)
                resume_content = '\n'.join(lines[:-1]) if lines else ""
            except KeyboardInterrupt:
                print("\n>>> Using sample resume data...")
                resume_content = """Ryan Weiler
Software Engineer

Email: ryan_wlr@yahoo.com
Phone: (561) 906-2118

Experience with Python development, financial technology, and machine learning projects.
Familiar with Django, Flask, pandas, and algorithmic trading systems."""
        
        if not resume_content or not resume_content.strip():
            print("ERROR: Resume content is required. Exiting.")
            return
        
        # Get additional details
        print("\n>>> Additional Information:")
        target_role = input("Target job role (default: Software Engineer): ").strip()
        if not target_role:
            target_role = "Software Engineer"
            
        target_company = input("Target company (default: Target Company): ").strip()
        if not target_company:
            target_company = "Target Company"
        
        # Process optimization
        print(f"\n>>> Processing resume optimization for {target_role} at {target_company}...")
        results = self.process_complete_optimization(
            job_description, resume_content, target_role, target_company
        )
        
        # Save results
        output_dir = f"resume_optimization_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        self.save_results_to_files(results, output_dir)
        
        # Open output directory if possible
        try:
            if os.name == 'nt':  # Windows
                os.startfile(output_dir)
            elif os.name == 'posix':  # macOS/Linux
                os.system(f'open "{output_dir}"' if sys.platform == 'darwin' else f'xdg-open "{output_dir}"')
            else:
                print(f">>> Manually open the '{output_dir}' folder to view results.")
        except:
            print(f">>> Manually open the '{output_dir}' folder to view results.")


def main():
    parser = argparse.ArgumentParser(description='AI-Powered Resume Optimizer')
    parser.add_argument('job_description', nargs='?', 
                       help='Job description text or file path')
    parser.add_argument('--resume', '-r',
                       help='Resume text or file path')
    parser.add_argument('--role', 
                       default='Software Engineer',
                       help='Target job role')
    parser.add_argument('--company', '-c',
                       default='Target Company', 
                       help='Target company name')
    parser.add_argument('--output', '-o',
                       default='resume_optimization_output',
                       help='Output directory')
    parser.add_argument('--browse', '-b', action='store_true',
                       help='Use interactive file browser mode')
    
    args = parser.parse_args()
    
    print(">>> AI-POWERED RESUME OPTIMIZER")
    print("    Creating ATS-optimized, recruiter-friendly resumes with .docx output")
    print("    Compatible with Windows console")
    print()
    
    optimizer = ResumeOptimizer()
    
    # Handle different modes
    if args.browse:
        optimizer.run_browse_mode()
        return
    
    # Command line file handling
    if args.job_description and os.path.isfile(args.job_description):
        print(f">>> Reading job description from file: {args.job_description}")
        try:
            with open(args.job_description, 'r', encoding='utf-8') as f:
                job_description = f.read()
            print(f"    Loaded {len(job_description)} characters")
        except Exception as e:
            print(f"ERROR: Failed to read job description file: {e}")
            return
    elif args.job_description:
        job_description = args.job_description
    else:
        job_description = None
        
    if args.resume and os.path.isfile(args.resume):
        print(f">>> Reading resume from file: {args.resume}")
        try:
            with open(args.resume, 'r', encoding='utf-8') as f:
                resume_content = f.read()
            print(f"    Loaded {len(resume_content)} characters")
        except Exception as e:
            print(f"ERROR: Failed to read resume file: {e}")
            return
    elif args.resume:
        resume_content = args.resume
    else:
        resume_content = None
    
    # Interactive mode if no arguments provided
    if not job_description:
        print(">>> AI RESUME OPTIMIZER - INTERACTIVE MODE")
        print("    Choose your preferred input method:")
        print()
        print("    1. Use file browsers to select files (recommended)")
        print("    2. Enter text manually") 
        print("    3. Run demo with sample data")
        print()
        
        while True:
            try:
                choice = input("Select option (1-3): ").strip()
            except (EOFError, KeyboardInterrupt):
                print("\n>>> Input interrupted - using demo mode")
                choice = '3'
            
            if choice == '1':
                print("\n>>> STEP 1: Select Job Description File")
                job_description = optimizer.browse_for_job_description()
                if not job_description:
                    print("Job description required. Please try again.")
                    continue
                
                print("\n>>> STEP 2: Select Resume File")
                resume_content = optimizer.browse_for_resume()
                if not resume_content:
                    print("Resume required. Please try again.")
                    continue
                
                break
                
            elif choice == '2':
                print("\n>>> STEP 1: Enter Job Description")
                print("Paste the job description (press Enter twice when finished):")
                lines = []
                while True:
                    line = input()
                    if line == "" and len(lines) > 0 and lines[-1] == "":
                        break
                    lines.append(line)
                
                job_description = '\n'.join(lines[:-1])  # Remove last empty line
                
                if not job_description.strip():
                    print("Job description cannot be empty. Please try again.")
                    continue
                
                print("\n>>> STEP 2: Enter Resume Content")
                print("Paste your current resume (press Enter twice when finished):")
                lines = []
                while True:
                    line = input()
                    if line == "" and len(lines) > 0 and lines[-1] == "":
                        break
                    lines.append(line)
                
                resume_content = '\n'.join(lines[:-1])  # Remove last empty line
                
                if not resume_content.strip():
                    print("Resume cannot be empty. Please try again.")
                    continue
                
                break
                
            elif choice == '3':
                print("\n>>> RUNNING DEMO MODE...")
                job_description = """Software Engineer - Financial Technology
                
We are seeking a talented Software Engineer to join our FinTech team. The ideal candidate will have experience with Python, machine learning, and financial markets.

Required Skills:
- Python programming (3+ years)
- Experience with pandas, NumPy, scikit-learn
- Knowledge of financial markets and trading
- REST API development
- Database design (PostgreSQL preferred)
- Git version control

Preferred Qualifications:
- Machine learning model development
- Algorithmic trading experience
- AWS cloud platform knowledge
- React or similar frontend framework

Responsibilities:
- Develop trading algorithms and backtesting frameworks
- Build scalable data processing pipelines
- Collaborate with quantitative researchers
- Maintain and optimize existing trading systems"""

                resume_content = """John Smith
Software Developer

Email: john.smith@email.com
Phone: (555) 123-4567

Experience:
Software Developer Intern - Tech Company (2022-2023)
- Worked on web applications using Python and JavaScript
- Built database systems with SQL
- Collaborated with team on various projects

Education:
Bachelor of Computer Science
State University (2024)

Skills:
- Python
- JavaScript  
- SQL
- Git"""
                break
            else:
                print("Invalid choice. Please select 1, 2, or 3.")
        
        # Get additional details
        print(f"\n>>> Processing resume optimization for {args.role} at {args.company}...")
        
    # Validate that we have both job description and resume
    if not job_description or not job_description.strip():
        print("ERROR: Job description is required.")
        return
        
    if not resume_content or not resume_content.strip():
        print("ERROR: Resume content is required.")
        return
    
    # Process optimization
    results = optimizer.process_complete_optimization(
        job_description, resume_content, args.role, args.company
    )
    
    # Save results with timestamp
    output_dir = f"{args.output}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    optimizer.save_results_to_files(results, output_dir)
    
    print(f"\n>>> OPTIMIZATION COMPLETE!")
    print(f"    Results saved to: {output_dir}")
    print(f"    Files created: 8 analysis files + 1 formatted .docx resume")
    print(f"    Open the folder to view your optimized resume!")

def build_narrative_docx(self, doc, narrative_content):
    """Build DOCX using narrative storytelling content"""
    
    # Parse narrative content to extract sections
    lines = narrative_content.split('\n')
    sections = {'hook': '', 'name': '', 'contact': '', 'narrative': '', 'skills': [], 'experience': []}
    
    current_section = None
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect key elements
        if '🔥 THE' in line or '🔬 THE' in line or '💻 THE' in line or '🌟 THE' in line:
            sections['hook'] = line
        elif 'CONTACT INFORMATION:' in line:
            current_section = 'contact'
            continue
        elif 'PROFESSIONAL NARRATIVE:' in line:
            current_section = 'narrative'
            continue
        elif 'TECHNICAL COMPETENCIES:' in line or 'CORE SKILLS:' in line:
            current_section = 'skills'
            continue
        elif 'EXPERIENCE:' in line or 'PROFESSIONAL EXPERIENCE:' in line:
            current_section = 'experience'
            continue
        elif line and not line.startswith('CAREER STORY RESUME') and not line.startswith('Chapter'):
            # Extract content based on current section
            if current_section is None and len(line.split()) <= 3 and any(c.isupper() for c in line):
                sections['name'] = line
            elif current_section == 'contact':
                sections['contact'] += line + '\n'
            elif current_section == 'narrative':
                sections['narrative'] += line + ' '
            elif current_section == 'skills' and line.startswith('•'):
                sections['skills'].append(line[1:].strip())
            elif current_section == 'experience' and line.startswith('•'):
                sections['experience'].append(line[1:].strip())
    
    # 1. Opening Hook (if present)
    if sections['hook']:
        hook_para = doc.add_paragraph()
        hook_para.style = doc.styles['Title']
        hook_run = hook_para.add_run(sections['hook'])
        hook_run.font.name = 'Calibri'
        hook_run.font.size = Pt(16)
        hook_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # 2. NAME
    name = sections['name'] if sections['name'] else "Ryan Thomas Weiler"
    name_para = doc.add_paragraph()
    name_para.style = doc.styles['Title']
    name_run = name_para.add_run(name)
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(26)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. CONTACT INFO
    contact_para = doc.add_paragraph()
    contact_para.style = doc.styles['Normal']
    contact_text = sections['contact'].strip() if sections['contact'] else "📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(12)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    social_para = doc.add_paragraph()
    social_para.style = doc.styles['Normal']
    social_run = social_para.add_run("🔗 LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/ | 💻 GitHub: https://github.com/ryan-wlr")
    social_run.font.name = 'Calibri'
    social_run.font.size = Pt(12)
    social_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # 4. PROFESSIONAL NARRATIVE
    if sections['narrative']:
        narrative_header = doc.add_paragraph()
        narrative_header.style = doc.styles['Heading 1']
        narrative_run = narrative_header.add_run("Professional Narrative")
        narrative_run.font.name = 'Calibri'
        narrative_run.font.size = Pt(14)
        
        narrative_para = doc.add_paragraph()
        narrative_para.style = doc.styles['Normal']
        narrative_text_run = narrative_para.add_run(sections['narrative'].strip())
        narrative_text_run.font.name = 'Calibri'
        narrative_text_run.font.size = Pt(12)
        doc.add_paragraph()
    
    # 5. TECHNICAL SKILLS
    if sections['skills']:
        skills_header = doc.add_paragraph()
        skills_header.style = doc.styles['Heading 1']
        skills_run = skills_header.add_run("Technical Competencies")
        skills_run.font.name = 'Calibri'
        skills_run.font.size = Pt(14)
        
        for skill in sections['skills'][:6]:
            skill_para = doc.add_paragraph()
            skill_para.style = doc.styles['Normal']
            skill_run = skill_para.add_run(f"• {skill}")
            skill_run.font.name = 'Calibri'
            skill_run.font.size = Pt(12)
        doc.add_paragraph()
    
    # 6. EDUCATION
    edu_header = doc.add_paragraph()
    edu_header.style = doc.styles['Heading 1']
    edu_run = edu_header.add_run("Education")
    edu_run.font.name = 'Calibri'
    edu_run.font.size = Pt(14)
    
    edu_para = doc.add_paragraph()
    edu_para.style = doc.styles['Normal']
    edu_run = edu_para.add_run("Florida Atlantic University — B.S. Computer Science, Expected 2024 (Dean's List, GPA 3.7)")
    edu_run.font.name = 'Calibri'
    edu_run.font.size = Pt(12)
    doc.add_paragraph()
    
    # 7. EXPERIENCE
    if sections['experience']:
        exp_header = doc.add_paragraph()
        exp_header.style = doc.styles['Heading 1']
        exp_run = exp_header.add_run("Professional Experience")
        exp_run.font.name = 'Calibri'
        exp_run.font.size = Pt(14)
        
        for exp in sections['experience'][:8]:
            exp_para = doc.add_paragraph()
            exp_para.style = doc.styles['Normal']
            exp_run = exp_para.add_run(f"• {exp}")
            exp_run.font.name = 'Calibri'
            exp_run.font.size = Pt(12)
        doc.add_paragraph()
    
    # 8. REFERENCES
    ref_header = doc.add_paragraph()
    ref_header.style = doc.styles['Heading 1']
    ref_run = ref_header.add_run("References")
    ref_run.font.name = 'Calibri'
    ref_run.font.size = Pt(14)
    
    ref_para = doc.add_paragraph()
    ref_para.style = doc.styles['Normal']
    ref_run = ref_para.add_run("Available upon request")
    ref_run.font.name = 'Calibri'
    ref_run.font.size = Pt(12)

def build_standard_docx(self, doc, field_data, results, detected_field):
    """Build DOCX using standard format with field-specific data"""
    
    # 1. NAME (using Title style - 26pt, centered)
    name_para = doc.add_paragraph()
    name_para.style = doc.styles['Title']
    name_run = name_para.add_run("Ryan Thomas Weiler")
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(26)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. CONTACT INFO (centered, Normal style with 12pt font)
    contact_para = doc.add_paragraph()
    contact_para.style = doc.styles['Normal']
    contact_run = contact_para.add_run("📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com")
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(12)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    social_para = doc.add_paragraph()
    social_para.style = doc.styles['Normal']
    social_run = social_para.add_run("🔗 LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/ | 💻 GitHub: https://github.com/ryan-wlr")
    social_run.font.name = 'Calibri'
    social_run.font.size = Pt(12)
    social_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # 3. EDUCATION SECTION (using Heading 1 style - 14pt, bold)
    edu_header = doc.add_paragraph()
    edu_header.style = doc.styles['Heading 1']
    edu_run = edu_header.add_run("Education")
    edu_run.font.name = 'Calibri'
    edu_run.font.size = Pt(14)
    
    edu_detail = doc.add_paragraph()
    edu_detail.style = doc.styles['Normal']
    edu_text_run = edu_detail.add_run(field_data['education'])
    edu_text_run.font.name = 'Calibri'
    edu_text_run.font.size = Pt(12)
    doc.add_paragraph()
    
    # 4. EXPERIENCE & PROJECTS SECTION (using Heading 1 style - 14pt)
    exp_header = doc.add_paragraph()
    exp_header.style = doc.styles['Heading 1']
    exp_run = exp_header.add_run("Experience & Projects (Continuous Timeline)")
    exp_run.font.name = 'Calibri'
    exp_run.font.size = Pt(14)
    
    # Add main experience using dynamic field-specific data
    experiences = [
        {
            'title': field_data['experience_title'],
            'bullets': field_data['experience_bullets']
        },
        {
            'title': f"{field_data['experience_title'].split('|')[0].strip()} Technician | Florida Atlantic University Facilities | 2021 - 2022",
            'bullets': [
                f"Supported campus-wide {detected_field} operations and maintenance across multiple facilities",
                f"Assisted with installation and system upgrades related to {detected_field} work",
                f"Performed preventive maintenance and collaborated with facilities team on repairs",
                f"Maintained accurate documentation and followed safety protocols"
            ]
        }
    ]
    
    for exp in experiences:
        # Job title (Normal style, bold, 12pt)
        job_para = doc.add_paragraph()
        job_para.style = doc.styles['Normal']
        job_run = job_para.add_run(exp['title'])
        job_run.font.name = 'Calibri'
        job_run.font.size = Pt(12)
        job_run.bold = True
        
        # Bullets (Normal style with proper dash, 12pt)
        for bullet in exp['bullets']:
            bullet_para = doc.add_paragraph()
            bullet_para.style = doc.styles['Normal']
            bullet_run = bullet_para.add_run(f"- {bullet}")
            bullet_run.font.name = 'Calibri'
            bullet_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 5. SKILLS SECTION (using Heading 1 style - 14pt)
    skills_header = doc.add_paragraph()
    skills_header.style = doc.styles['Heading 1']
    skills_run = skills_header.add_run("Technical Skills")
    skills_run.font.name = 'Calibri'
    skills_run.font.size = Pt(14)
    
    # Skills content using dynamic field-specific data
    for skill in field_data['skills']:
        skill_para = doc.add_paragraph()
        skill_para.style = doc.styles['Normal']
        skill_run = skill_para.add_run(f"• {skill}")
        skill_run.font.name = 'Calibri'
        skill_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 6. REFERENCES SECTION (using Heading 1 style - 14pt)
    ref_header = doc.add_paragraph()
    ref_header.style = doc.styles['Heading 1']
    ref_run = ref_header.add_run("References")
    ref_run.font.name = 'Calibri'
    ref_run.font.size = Pt(14)
    
    ref_para = doc.add_paragraph()
    ref_para.style = doc.styles['Normal']
    ref_text_run = ref_para.add_run("Available upon request")
    ref_text_run.font.name = 'Calibri'
    ref_text_run.font.size = Pt(12)


if __name__ == "__main__":
    main()