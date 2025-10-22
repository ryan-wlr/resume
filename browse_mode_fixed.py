#!/usr/bin/env python3
"""
Complete Browse Mode Resume Optimizer - Fixed Version
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime

# Add current directory to path to import resume_windows
sys.path.insert(0, os.getcwd())

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def validate_docx_file(filepath):
    """Check if a DOCX file can be opened by python-docx"""
    try:
        # Try to open and read the first paragraph
        doc = Document(filepath)
        # Just try to access the paragraphs to trigger any errors
        paragraphs = [p.text for p in doc.paragraphs[:1]]
        return True, "Valid DOCX file"
    except Exception as e:
        return False, str(e)

def read_file_safely(filepath):
    """Read a file safely, handling both .txt and .docx files"""
    try:
        # First check if file exists
        if not os.path.exists(filepath):
            return f"Error: File not found at '{filepath}'"
            
        # Check file size (avoid empty files or very large files)
        file_size = os.path.getsize(filepath)
        if file_size == 0:
            return f"Error: File is empty: '{filepath}'"
        if file_size > 50 * 1024 * 1024:  # 50MB limit
            return f"Error: File too large (>{file_size/1024/1024:.1f}MB): '{filepath}'"
            
        print(f"📄 Reading file: {os.path.basename(filepath)} ({file_size:,} bytes)")
        
        if filepath.lower().endswith('.docx'):
            if not HAS_DOCX:
                return f"Error: python-docx not installed. Run: pip install python-docx"
            
            try:
                # Read .docx file with improved error handling
                doc = Document(filepath)
                text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Only add non-empty paragraphs
                        text.append(paragraph.text)
                
                if not text:
                    return f"Error: DOCX file appears to be empty or contains no readable text"
                    
                content = '\n'.join(text)
                print(f"✅ Successfully read {len(content)} characters from DOCX file")
                return content
                
            except Exception as docx_error:
                # Try to provide more specific error information
                error_msg = str(docx_error)
                if "Package" in error_msg:
                    return f"Error: DOCX file may be corrupted, password-protected, or in use by another program. Close Microsoft Word if open and try again. Details: {error_msg}"
                else:
                    return f"Error reading DOCX file: {error_msg}"
        else:
            # Read text file with proper encoding detection
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        content = f.read()
                        if content.strip():  # Make sure we got actual content
                            print(f"✅ Successfully read {len(content)} characters from text file")
                            return content
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    continue
                    
            # If all encodings fail, try reading as binary and decode
            try:
                with open(filepath, 'rb') as f:
                    content = f.read()
                    decoded = content.decode('utf-8', errors='ignore')
                    if decoded.strip():
                        print(f"✅ Successfully read {len(decoded)} characters (binary mode)")
                        return decoded
            except Exception as e:
                pass
                
            return f"Error: Could not read file with any supported encoding: '{filepath}'"
            
    except Exception as e:
        return f"Error reading file '{filepath}': {e}"

def browse_mode_optimizer():
    """Complete browse mode with file dialogs and optimization"""
    print("🚀 AI-POWERED RESUME OPTIMIZER - BROWSE MODE")
    print("=" * 60)
    print()
    print("📋 This will help you:")
    print("   1. Select your JOB DESCRIPTION file (.txt or .docx)")
    print("   2. Select your CURRENT RESUME file (.txt or .docx)")
    print("   3. Generate an optimized resume matching the job")
    print()
    print("💡 TIP: Make sure you select DIFFERENT files for job and resume!")
    print()
    
    try:
        # Create root window
        root = tk.Tk()
        root.withdraw()  # Hide main window
        root.attributes('-topmost', True)  # Keep dialogs on top
        
        # Step 1: Get job description
        print("📁 STEP 1: Select JOB DESCRIPTION file")
        print("   (This should contain the job posting/requirements)")
        input("   Press Enter to open file dialog...")
        
        job_file = filedialog.askopenfilename(
            title="Select Job Description File",
            initialdir=os.getcwd(),
            filetypes=[
                ("Text files", "*.txt"),
                ("Word documents", "*.docx"),
                ("All files", "*.*")
            ]
        )
        
        if not job_file:
            print("❌ No job file selected. Exiting.")
            root.destroy()
            return
            
        print(f"✅ Selected: {os.path.basename(job_file)}")
        job_content = read_file_safely(job_file)
        
        if job_content.startswith("Error"):
            print(f"❌ {job_content}")
            print()
            print("🔧 TROUBLESHOOTING TIPS:")
            if "DOCX" in job_content and ("Package" in job_content or "Permission" in job_content):
                print("   • The DOCX file may be corrupted, password-protected, or locked by OneDrive")
                print("   • Try closing Microsoft Word if it's open")
                print("   • If it's in OneDrive, try copying the file to your local Desktop first")
                print("   • Or save the file as a .txt file instead (File > Save As > Plain Text)")
                print("   • Or copy-paste the content into Notepad and save as .txt")
                print()
                
                # Offer to use our sample files as an example
                print("🆘 QUICK FIX OPTIONS:")
                print("   1. Use our sample brain surgeon job description")
                print("   2. Try to fix the DOCX file issue")
                print("   3. Exit and fix the file manually")
                
                choice = input("   Choose option (1-3): ").strip()
                
                if choice == '1' and os.path.exists('sample_brain_surgeon_job.txt'):
                    job_content = read_file_safely('sample_brain_surgeon_job.txt')
                    if not job_content.startswith("Error"):
                        print("✅ Using sample brain surgeon job description")
                    else:
                        print("❌ Error with sample file")
                        root.destroy()
                        return
                else:
                    print("❌ Exiting - please fix the file issue and try again")
                    print("\nNEXT STEPS:")
                    print("1. Copy the DOCX file from OneDrive to your Desktop")
                    print("2. Open it in Word and save as .txt file")
                    print("3. Run this program again and select the .txt file")
                    root.destroy()
                    return
            else:
                root.destroy()
                return
            
        print(f"📄 Loaded job description: {len(job_content)} characters")
        print(f"Preview: {job_content[:150]}...")
        print()
        
        # Step 2: Get resume
        print("📁 STEP 2: Select your CURRENT RESUME file")
        print("   (This should be your existing resume to optimize)")
        input("   Press Enter to open file dialog...")
        
        resume_file = filedialog.askopenfilename(
            title="Select Current Resume File",
            initialdir=os.getcwd(),
            filetypes=[
                ("Text files", "*.txt"),
                ("Word documents", "*.docx"),
                ("All files", "*.*")
            ]
        )
        
        if not resume_file:
            print("❌ No resume file selected. Exiting.")
            root.destroy()
            return
            
        print(f"✅ Selected: {os.path.basename(resume_file)}")
        resume_content = read_file_safely(resume_file)
        
        if resume_content.startswith("Error"):
            print(f"❌ {resume_content}")
            print()
            print("🔧 RESUME FILE TROUBLESHOOTING:")
            if "DOCX" in resume_content and ("Package" in resume_content or "Permission" in resume_content):
                print("   • The resume DOCX file has the same issue as the job description")
                print("   • Try copying it from OneDrive to your Desktop first")
                print("   • Or save as .txt file (File > Save As > Plain Text)")
                print()
                
                print("🆘 QUICK FIX: Use our sample resume?")
                choice = input("   Use sample resume? (y/n): ").strip().lower()
                
                if choice == 'y' and os.path.exists('sample_current_resume.txt'):
                    resume_content = read_file_safely('sample_current_resume.txt')
                    if not resume_content.startswith("Error"):
                        print("✅ Using sample current resume")
                    else:
                        print("❌ Error with sample resume file")
                        root.destroy()
                        return
                else:
                    print("❌ Exiting - please fix the resume file issue and try again")
                    root.destroy()
                    return
            else:
                root.destroy()
                return
            
        print(f"📄 Loaded resume: {len(resume_content)} characters")
        print(f"Preview: {resume_content[:150]}...")
        print()
        
        root.destroy()
        
        # Check if same file selected
        if job_file == resume_file:
            print("⚠️  WARNING: You selected the same file for both job and resume!")
            print("   This won't work properly. Please run again and select different files.")
            return
            
        # Step 3: Get additional info
        print("ℹ️  STEP 3: Additional Information")
        role = input("Target role (e.g., Plumber, Electrician, Software Engineer): ").strip()
        if not role:
            role = "Professional"
            
        company = input("Target company name (optional): ").strip()
        if not company:
            company = "Target Company"
            
        print()
        print("📋 SUMMARY:")
        print(f"   Job file: {os.path.basename(job_file)}")
        print(f"   Resume file: {os.path.basename(resume_file)}")
        print(f"   Target role: {role}")
        print(f"   Target company: {company}")
        print()
        
        # Step 4: Run optimization
        proceed = input("🚀 Proceed with optimization? (y/n): ").strip().lower()
        
        if proceed != 'y':
            print("Cancelled by user.")
            return
            
        print()
        print("⚙️  Running AI optimization...")
        
        # Import and run the optimizer
        try:
            from resume_windows import ResumeOptimizer
            
            optimizer = ResumeOptimizer()
            results = optimizer.process_complete_optimization(
                job_content, resume_content, role, company
            )
            
            # Save results
            output_dir = f"browse_mode_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            optimizer.save_results_to_files(results, output_dir)
            
            print(f"✅ SUCCESS! Results saved to: {output_dir}")
            print(f"📂 Check the optimized_resume.docx file in that folder!")
            
            # Show brief results
            if 'final_resume' in results:
                print(f"\n📊 Generated resume preview:")
                print(results['final_resume'][:300] + "...")
                
        except Exception as e:
            print(f"❌ Optimization error: {e}")
            import traceback
            traceback.print_exc()
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    try:
        browse_mode_optimizer()
    except KeyboardInterrupt:
        print("\n👋 Goodbye!")