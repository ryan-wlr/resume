#!/usr/bin/env python3
"""Test the duplicate link fix"""

import sys
import os
import tempfile
import shutil

# Add the current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from resume_windows import ResumeOptimizer

def test_link_fix():
    """Test that duplicate links are removed and links look clickable"""
    print("🔍 Testing duplicate link fix...")
    
    # Create temp output directory
    temp_dir = tempfile.mkdtemp()
    print(f"Using temp directory: {temp_dir}")
    
    try:
        # Initialize optimizer
        optimizer = ResumeOptimizer()
        
        # Mock detected field
        optimizer.current_detected_field = 'optical_engineer'
        
        # Create mock results with narrative content that has contact info with links
        mock_results = {
            '7_narrative_resume': """
CAREER STORY RESUME - OPTICAL ENGINEER

🔬 THE LIGHT ARCHITECT: Engineering the future through precision optics and photonic innovation

RYAN THOMAS WEILER
Optical Engineer Professional

CONTACT INFORMATION:
📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com
🔗 LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/ | 💻 GitHub: https://github.com/ryan-wlr

PROFESSIONAL NARRATIVE:
Light has always been my medium of choice for solving complex engineering challenges.

TECHNICAL COMPETENCIES:
• Advanced Optical Design & Modeling (Zemax, Code V, LightTools)

PROFESSIONAL EXPERIENCE:
• Designed revolutionary laser systems improving efficiency by 40%
"""
        }
        
        # Create the DOCX
        optimizer.create_formatted_docx_resume(mock_results, temp_dir)
        
        # Check if DOCX was created
        docx_path = os.path.join(temp_dir, 'optimized_resume.docx')
        if os.path.exists(docx_path):
            print("✅ DOCX created successfully")
            
            # Try to read the DOCX to check for link duplication
            try:
                from docx import Document
                doc = Document(docx_path)
                
                # Get all paragraphs and their text
                content_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                print(f"\n📄 All content lines ({len(content_lines)} total):")
                for i, line in enumerate(content_lines):
                    print(f"  {i+1:2d}. {line}")
                
                # Count GitHub/LinkedIn mentions
                github_count = sum(1 for line in content_lines if 'github' in line.lower())
                linkedin_count = sum(1 for line in content_lines if 'linkedin' in line.lower())
                
                print(f"\n📊 Link count:")
                print(f"   GitHub mentions: {github_count}")
                print(f"   LinkedIn mentions: {linkedin_count}")
                
                if github_count <= 1 and linkedin_count <= 1:
                    print("✅ SUCCESS: No duplicate links!")
                else:
                    print("❌ Still have duplicate links")
                
                # Check for blue/underlined formatting (simulated hyperlinks)
                print(f"\n🔗 Checking for link styling...")
                link_styled_count = 0
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run.font.color and run.font.color.rgb and 'https://' in run.text:
                            print(f"   Found styled link: {run.text}")
                            link_styled_count += 1
                
                if link_styled_count > 0:
                    print(f"✅ Found {link_styled_count} styled links")
                else:
                    print("⚠️  No styled links found")
                    
            except ImportError:
                print("⚠️  python-docx not available for content verification")
                
        else:
            print("❌ DOCX file was not created")
            
    except Exception as e:
        print(f"❌ Error during testing: {e}")
        import traceback
        traceback.print_exc()
        
    finally:
        # Clean up
        try:
            shutil.rmtree(temp_dir)
            print(f"\n🗑️  Cleaned up temp directory: {temp_dir}")
        except:
            pass

if __name__ == "__main__":
    test_link_fix()