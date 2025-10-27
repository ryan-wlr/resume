#!/usr/bin/env python3
"""Test clickable hyperlinks in DOCX"""

import sys
import os
import tempfile
import shutil

# Add the current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from resume_windows import ResumeOptimizer

def test_clickable_links():
    """Test that links in DOCX are actually clickable hyperlinks"""
    print("🔍 Testing clickable hyperlinks in DOCX...")
    
    # Create temp output directory
    temp_dir = tempfile.mkdtemp()
    print(f"Using temp directory: {temp_dir}")
    
    try:
        # Initialize optimizer
        optimizer = ResumeOptimizer()
        
        # Mock detected field
        optimizer.current_detected_field = 'optical_engineer'
        
        # Create mock results with narrative content that has contact info with links
        mock_results = {
            '7_narrative_resume': """
CAREER STORY RESUME - OPTICAL ENGINEER

🔬 THE LIGHT ARCHITECT: Engineering the future through precision optics and photonic innovation

RYAN THOMAS WEILER
Optical Engineer Professional

CONTACT INFORMATION:
📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com
🔗 LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/ | 💻 GitHub: https://github.com/ryan-wlr

PROFESSIONAL NARRATIVE:
Light has always been my medium of choice for solving complex engineering challenges.

TECHNICAL COMPETENCIES:
• Advanced Optical Design & Modeling (Zemax, Code V, LightTools)

PROFESSIONAL EXPERIENCE:
• Designed revolutionary laser systems improving efficiency by 40%
"""
        }
        
        # Create the DOCX
        optimizer.create_formatted_docx_resume(mock_results, temp_dir)
        
        # Check if DOCX was created
        docx_path = os.path.join(temp_dir, 'optimized_resume.docx')
        if os.path.exists(docx_path):
            print("✅ DOCX created successfully")
            print(f"📂 DOCX saved to: {docx_path}")
            
            # Try to read the DOCX to check for hyperlinks
            try:
                from docx import Document
                doc = Document(docx_path)
                
                # Get all paragraphs and their text
                content_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                print(f"\n📄 All content lines ({len(content_lines)} total):")
                for i, line in enumerate(content_lines):
                    print(f"  {i+1:2d}. {line}")
                
                # Check for hyperlinks in the document XML
                print(f"\n🔗 Checking for hyperlinks...")
                hyperlink_count = 0
                
                # Look for hyperlink relationships
                for rel in doc.part.rels.values():
                    if "hyperlink" in rel.reltype:
                        print(f"   Found hyperlink relationship: {rel.target_ref}")
                        hyperlink_count += 1
                
                # Also check paragraph XML for hyperlink elements
                xml_hyperlinks = 0
                for paragraph in doc.paragraphs:
                    paragraph_xml = paragraph._p.xml
                    if 'w:hyperlink' in paragraph_xml:
                        print(f"   Found hyperlink in paragraph: {paragraph.text[:50]}...")
                        xml_hyperlinks += 1
                
                print(f"\n📊 Hyperlink analysis:")
                print(f"   Hyperlink relationships: {hyperlink_count}")
                print(f"   XML hyperlink elements: {xml_hyperlinks}")
                
                if hyperlink_count > 0 or xml_hyperlinks > 0:
                    print("✅ SUCCESS: Found clickable hyperlinks!")
                    print("💡 Try opening the DOCX file and clicking the links to test")
                else:
                    print("❌ No hyperlinks found - links may be plain text")
                    
                # Show the file path for manual testing
                abs_path = os.path.abspath(docx_path)
                print(f"\n🧪 MANUAL TEST:")
                print(f"   1. Open this file: {abs_path}")
                print(f"   2. Click on the LinkedIn and GitHub links")
                print(f"   3. Verify they open in your browser")
                
            except ImportError:
                print("⚠️  python-docx not available for content verification")
            except Exception as e:
                print(f"❌ Error checking hyperlinks: {e}")
                import traceback
                traceback.print_exc()
                
        else:
            print("❌ DOCX file was not created")
            
    except Exception as e:
        print(f"❌ Error during testing: {e}")
        import traceback
        traceback.print_exc()
        
    finally:
        # Don't clean up so we can manually test the file
        print(f"\n📁 DOCX file left at: {temp_dir}")
        print("   (You can manually delete this folder when done testing)")

if __name__ == "__main__":
    test_clickable_links()