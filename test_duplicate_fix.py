#!/usr/bin/env python3
"""Test that we fixed the duplicate content issue in DOCX creation"""

import sys
import os
import tempfile
import shutil

# Add the current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from resume_windows import ResumeOptimizer

def test_no_duplicates():
    """Test that storytelling DOCX doesn't have duplicate name/contact/education"""
    print("🔍 Testing for duplicate content in storytelling DOCX...")
    
    # Create temp output directory
    temp_dir = tempfile.mkdtemp()
    print(f"Using temp directory: {temp_dir}")
    
    try:
        # Initialize optimizer
        optimizer = ResumeOptimizer()
        
        # Mock detected field
        optimizer.current_detected_field = 'ai_researcher'
        
        # Create mock results with narrative content
        mock_results = {
            '7_narrative_resume': """
CAREER STORY RESUME FOR AI RESEARCHER

🔬 THE AI VISIONARY WHO TRANSFORMS DATA INTO INTELLIGENT SOLUTIONS

RYAN WEILER

CONTACT INFORMATION:
📞 (561) 906-2118 | ✉️ ryan_wlr@yahoo.com
🔗 LinkedIn: https://www.linkedin.com/in/ryan-weiler-7a3119190/
💻 GitHub: https://github.com/ryan-wlr

PROFESSIONAL NARRATIVE:
As an innovative AI researcher with expertise in machine learning algorithms and neural networks, I bring a unique blend of theoretical knowledge and practical implementation skills to drive breakthrough solutions in artificial intelligence.

TECHNICAL COMPETENCIES:
• Advanced Machine Learning & Deep Learning Frameworks (TensorFlow, PyTorch, Scikit-learn)
• Neural Network Architectures (CNNs, RNNs, Transformers, GANs)
• Natural Language Processing & Computer Vision Systems
• Data Science Pipeline Development (Python, R, SQL)
• Research Methodology & Statistical Analysis
• Cloud Computing Platforms (AWS, Google Cloud, Azure)

PROFESSIONAL EXPERIENCE:
• Led development of novel neural network architectures for image classification
• Implemented state-of-the-art NLP models for sentiment analysis
• Published research papers on reinforcement learning applications
• Collaborated on interdisciplinary AI projects with academic institutions
• Designed and deployed machine learning models in production environments
• Conducted extensive experiments with hyperparameter optimization
• Developed data preprocessing pipelines for large-scale datasets
• Mentored junior researchers in machine learning methodologies
"""
        }
        
        # Create the DOCX
        optimizer.create_formatted_docx_resume(mock_results, temp_dir)
        
        # Check if DOCX was created
        docx_path = os.path.join(temp_dir, 'optimized_resume.docx')
        if os.path.exists(docx_path):
            print("✅ DOCX created successfully")
            print(f"✅ File size: {os.path.getsize(docx_path)} bytes")
            
            # Try to read the DOCX to check for basic content structure
            try:
                from docx import Document
                doc = Document(docx_path)
                
                # Count paragraphs
                total_paragraphs = len(doc.paragraphs)
                print(f"✅ Total paragraphs: {total_paragraphs}")
                
                # Look for potential duplicate patterns
                content_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                # Check for duplicate names
                name_count = sum(1 for line in content_lines if 'Ryan' in line and 'Weiler' in line)
                print(f"✅ Name occurrences: {name_count} (should be 1)")
                
                # Check for duplicate contact info
                contact_count = sum(1 for line in content_lines if '(561) 906-2118' in line)
                print(f"✅ Contact info occurrences: {contact_count} (should be 1)")
                
                # Check for duplicate education headers
                edu_header_count = sum(1 for line in content_lines if line == 'Education')
                print(f"✅ Education header occurrences: {edu_header_count} (should be 1)")
                
                # Print first few content lines for verification
                print("\n📄 First 10 content lines:")
                for i, line in enumerate(content_lines[:10]):
                    print(f"  {i+1}. {line}")
                
                if name_count <= 1 and contact_count <= 1 and edu_header_count <= 1:
                    print("\n🎉 SUCCESS: No duplicate content detected!")
                else:
                    print("\n❌ WARNING: Potential duplicate content found")
                    
            except ImportError:
                print("⚠️  python-docx not available for content verification")
                
        else:
            print("❌ DOCX file was not created")
            
    except Exception as e:
        print(f"❌ Error during testing: {e}")
        import traceback
        traceback.print_exc()
        
    finally:
        # Clean up
        try:
            shutil.rmtree(temp_dir)
            print(f"🗑️  Cleaned up temp directory: {temp_dir}")
        except:
            pass

if __name__ == "__main__":
    test_no_duplicates()